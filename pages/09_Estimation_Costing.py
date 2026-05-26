import streamlit as st
from st_supabase_connection import SupabaseConnection
from database_utils import fetch_all_master_data
import json, math, io
from datetime import date, datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False

st.set_page_config(
    page_title="Estimation & Costing | BGEngg ERP",
    page_icon="📐",
    layout="wide",
)

# ─────────────────────────────────────────────────────────────────────────────
# COMPANY CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
BG_NAME    = "B&G Engineering Industries"
BG_TAGLINE = "Evaporation | Mixing | Drying"
BG_GSTIN   = "36AAIFB3357M1Z5"
BG_PAN     = "AAIFB3357M"
BG_ADDRESS = "Plot No.207/B & 208/A, Phase-III Industrial Park, Pashamylaram, Patancheru Mandal, Sangareddy Dist, Hyderabad – 502307"
BG_PHONE   = "+91 7995565800 / +91 9154971801"
BG_EMAIL   = "info@bgengineeringind.com"
BG_WEB     = "www.bgengineeringind.com"

conn = st.connection("supabase", type=SupabaseConnection)
if "master_data" not in st.session_state:
    st.session_state.master_data = fetch_all_master_data(conn)

# ─────────────────────────────────────────────────────────────────────────────
# LOGO — fetch from Supabase storage bucket once per session
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(ttl=3600)
def _fetch_logo_bytes():
    BUCKET = "progress-photos"
    LOGO_FILE = "logo.png"
    try:
        data = conn.client.storage.from_(BUCKET).download(LOGO_FILE)
        if data:
            return data, LOGO_FILE
    except Exception:
        pass
    return None, None

_LOGO_BYTES, _LOGO_FNAME = _fetch_logo_bytes()

# ─────────────────────────────────────────────────────────────────────────────
# SUPABASE HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def sb_fetch(table, select="*", order=None, filters=None):
    try:
        q = conn.table(table).select(select)
        if order:
            q = q.order(order)
        if filters:
            for col, val in filters.items():
                q = q.eq(col, val)
        return q.execute().data or []
    except Exception as e:
        st.error(f"DB read error ({table}): {e}")
        return []

def sb_insert(table, row):
    try:
        conn.table(table).insert(row).execute()
        return True
    except Exception as e:
        st.error(f"DB insert error ({table}): {e}")
        return False

def sb_update(table, row, match_col, match_val):
    try:
        conn.table(table).update(row).eq(match_col, match_val).execute()
        return True
    except Exception as e:
        st.error(f"DB update error ({table}): {e}")
        return False

# ─────────────────────────────────────────────────────────────────────────────
# GEOMETRY ENGINE
# ─────────────────────────────────────────────────────────────────────────────
PI = math.pi
DENSITY = {"SS316L": 8000, "SS304": 8000, "MS": 7850, "EN8": 7800, "Ti": 4500, "C22": 8690, "Hastelloy": 8890}
# ─────────────────────────────────────────────────────────────────────────────
# OH BUCKET MAPPING (Option B cost structure)
# Maps each oh_type from the OH master into one of two buckets:
#   FACTORY_OH = scales with production volume (shop floor)
#   ADMIN_OH   = mostly fixed (office, documentation, misc)
# Edit this dict any time to reclassify — no DB change needed.
# ─────────────────────────────────────────────────────────────────────────────
OH_BUCKET = {
    "LABOUR":         "FACTORY_OH",
    "LABOUR_BUFF":    "FACTORY_OH",
    "CONSUMABLES":    "FACTORY_OH",
    "TESTING":        "FACTORY_OH",
    "ELECTRO_POLISH": "FACTORY_OH",
    "DOCS":           "ADMIN_OH",
    "MISC":           "ADMIN_OH",
}
# Default bucket for any oh_type not listed above (defensive default)
OH_BUCKET_DEFAULT = "FACTORY_OH"
def _m(mm):
    return mm / 1000.0

def geom_cylindrical_shell(id_mm, ht_mm, thk_mm, density=8000, scrap=0.05):
    return PI * _m(id_mm) * _m(ht_mm) * _m(thk_mm) * density * (1 + scrap)

def geom_dish_end(shell_id_mm, thk_mm, density=8000, scrap=0.15):
    r = _m(shell_id_mm * 1.167) / 2
    return 1.09 * PI * r * r * _m(thk_mm) * density * (1 + scrap)

def geom_annular_plate(od_mm, id_mm, thk_mm, density=8000, scrap=0.05):
    return (PI / 4.0) * (_m(od_mm) ** 2 - _m(id_mm) ** 2) * _m(thk_mm) * density * (1 + scrap)

def geom_solid_round(dia_mm, length_mm, density=8000, scrap=0.15):
    return (PI / 4.0) * _m(dia_mm) ** 2 * _m(length_mm) * density * (1 + scrap)

def geom_flat_rect(w_mm, h_mm, thk_mm, density=8000, scrap=0.05):
    return _m(w_mm) * _m(h_mm) * _m(thk_mm) * density * (1 + scrap)

def geom_stiffener_ring(shell_id_mm, shell_thk_mm, shell_ht_mm,
                        pitch_mm, bar_w_mm, bar_thk_mm, density=8000, scrap=0.05):
    shell_od_mm = shell_id_mm + 2.0 * shell_thk_mm
    circ_m = PI * _m(shell_od_mm)
    n_rings = _m(shell_ht_mm) / _m(pitch_mm) if pitch_mm > 0 else 0.0
    wt_per_ring = circ_m * _m(bar_w_mm) * _m(bar_thk_mm) * density * (1 + scrap)
    return wt_per_ring, n_rings, wt_per_ring * n_rings

def geom_cone(large_id_mm, small_id_mm, ht_mm, thk_mm, density=8000, scrap=0.05):
    R1 = _m(large_id_mm) / 2
    R2 = _m(small_id_mm) / 2
    slant = math.sqrt(_m(ht_mm) ** 2 + (R1 - R2) ** 2)
    return PI * (R1 + R2) * slant * _m(thk_mm) * density * (1 + scrap)

def geom_rect_plate(length_mm, width_mm, thk_mm, density=8000, scrap=0.05):
    return _m(length_mm) * _m(width_mm) * _m(thk_mm) * density * (1 + scrap)

def geom_tube_bundle(tube_od_mm, tube_thk_mm, tube_length_mm, n_tubes, density=8000, scrap=0.05):
    mid_r = (_m(tube_od_mm) / 2) - (_m(tube_thk_mm) / 2)
    return PI * 2 * mid_r * _m(tube_length_mm) * _m(tube_thk_mm) * density * n_tubes * (1 + scrap)

def geom_limpet_coil(shell_id_mm, shell_thk_mm, shell_ht_mm,
                     pipe_od_mm, pipe_thk_mm, pitch_mm,
                     cover_bottom_dish=False, density=8000, scrap=0.10):
    """
    Limpet coil (half-pipe coil) welded on shell OD and optionally bottom dish.
    Returns: (wt_per_m, total_wt_kg, total_length_m)
    """
    shell_od_mm    = shell_id_mm + 2.0 * shell_thk_mm
    coil_mean_d_m  = _m(shell_od_mm + pipe_od_mm)
    coil_circ_m    = PI * coil_mean_d_m
    n_turns        = _m(shell_ht_mm) / _m(pitch_mm) if pitch_mm > 0 else 0
    shell_coil_m   = n_turns * coil_circ_m

    dish_coil_m = 0.0
    if cover_bottom_dish:
        dish_r_m    = _m(shell_od_mm) / 2.0 * 0.7
        dish_coil_m = PI * dish_r_m ** 2 / _m(pitch_mm) if pitch_mm > 0 else 0

    total_length_m = shell_coil_m + dish_coil_m

    mid_r_m     = (_m(pipe_od_mm) - _m(pipe_thk_mm)) / 2.0
    wt_per_m    = PI * 2 * mid_r_m * _m(pipe_thk_mm) * density
    total_wt    = wt_per_m * total_length_m * (1 + scrap)

    return round(wt_per_m, 4), round(total_wt, 3), round(total_length_m, 3)

# Default scrap % per part type
DEFAULT_SCRAP = {
    "shell": 5.0, "dish": 15.0, "annular": 5.0, "solid": 15.0,
    "flat": 5.0, "stiff": 5.0, "cone": 5.0, "rect": 5.0, "tube": 5.0,
    "limpet": 10.0,
}

PART_TYPES = {
    "Cylindrical shell": {
        "fields": [("id_mm", "Shell ID (mm)"), ("ht_mm", "Height (mm)"), ("thk_mm", "Thickness (mm)")],
        "fn": "shell",
    },
    "Dish end (torispherical)": {
        "fields": [("shell_id_mm", "Shell ID (mm)"), ("thk_mm", "Thickness (mm)")],
        "fn": "dish",
    },
    "Annular plate / flange": {
        "fields": [("od_mm", "Outer Dia OD (mm)"), ("id_mm", "Inner Dia ID (mm)"), ("thk_mm", "Thickness (mm)")],
        "fn": "annular",
    },
    "Solid round (shaft / bush)": {
        "fields": [("dia_mm", "Diameter (mm)"), ("length_mm", "Length (mm)")],
        "fn": "solid",
    },
    "Flat rectangle (blade / pad / gusset)": {
        "fields": [("w_mm", "Width (mm)"), ("h_mm", "Height (mm)"), ("thk_mm", "Thickness (mm)")],
        "fn": "flat",
    },
    "Stiffener rings (flat bar on shell OD)": {
        "fields": [
            ("shell_id_mm", "Shell ID (mm)"), ("shell_thk_mm", "Shell Thickness (mm)"),
            ("shell_ht_mm", "Shell Height (mm)"), ("pitch_mm", "Ring Pitch (mm)"),
            ("bar_w_mm", "Bar Width (mm)"), ("thk_mm", "Bar Thickness (mm)"),
        ],
        "fn": "stiff",
        "qty_derived": True,
    },
    "Cone / reducer": {
        "fields": [
            ("large_id_mm", "Large End ID (mm)"), ("small_id_mm", "Small End ID (mm)"),
            ("ht_mm", "Height (mm)"), ("thk_mm", "Thickness (mm)"),
        ],
        "fn": "cone",
    },
    "Rectangular plate": {
        "fields": [("length_mm", "Length (mm)"), ("width_mm", "Width (mm)"), ("thk_mm", "Thickness (mm)")],
        "fn": "rect",
    },
    "Tube bundle": {
        "fields": [
            ("tube_od_mm", "Tube OD (mm)"), ("tube_thk_mm", "Tube Thickness (mm)"),
            ("tube_length_mm", "Tube Length (mm)"), ("n_tubes", "Number of Tubes"),
        ],
        "fn": "tube",
    },
    "Limpet coil (half-pipe on shell)": {
        "fields": [
            ("shell_id_mm",   "Shell ID (mm)"),
            ("shell_thk_mm",  "Shell Thickness (mm)"),
            ("shell_ht_mm",   "Shell Height (mm)"),
            ("pipe_od_mm",    "Half-pipe OD (mm)"),
            ("pipe_thk_mm",   "Half-pipe Thickness (mm)"),
            ("pitch_mm",      "Coil Pitch (mm)"),
        ],
        "fn": "limpet",
        "has_checkbox": True,
    },
}

def calc_weight(fn, dims, density, qty, scrap_pct=None):
    d = dims
    used_qty = qty
    wt = 0.0
    sc = scrap_pct / 100.0 if scrap_pct is not None else None
    try:
        if fn == "shell":
            wt = geom_cylindrical_shell(d["id_mm"], d["ht_mm"], d["thk_mm"], density,
                                         scrap=sc if sc is not None else 0.05)
        elif fn == "dish":
            wt = geom_dish_end(d["shell_id_mm"], d["thk_mm"], density,
                                scrap=sc if sc is not None else 0.15)
        elif fn == "annular":
            wt = geom_annular_plate(d["od_mm"], d["id_mm"], d["thk_mm"], density,
                                     scrap=sc if sc is not None else 0.05)
        elif fn == "solid":
            wt = geom_solid_round(d["dia_mm"], d["length_mm"], density,
                                   scrap=sc if sc is not None else 0.15)
        elif fn == "flat":
            wt = geom_flat_rect(d["w_mm"], d["h_mm"], d["thk_mm"], density,
                                 scrap=sc if sc is not None else 0.05)
        elif fn == "stiff":
            sc2 = sc if sc is not None else 0.05
            wt, used_qty, total = geom_stiffener_ring(
                d.get("shell_id_mm", 0), d.get("shell_thk_mm", 0), d.get("shell_ht_mm", 0),
                d.get("pitch_mm", 100), d.get("bar_w_mm", 0), d.get("thk_mm", 0), density,
                scrap=sc2,
            )
            return round(wt, 3), round(total, 3), round(used_qty, 2)
        elif fn == "cone":
            wt = geom_cone(d["large_id_mm"], d["small_id_mm"], d["ht_mm"], d["thk_mm"], density,
                           scrap=sc if sc is not None else 0.05)
        elif fn == "rect":
            wt = geom_rect_plate(d["length_mm"], d["width_mm"], d["thk_mm"], density,
                                  scrap=sc if sc is not None else 0.05)
        elif fn == "tube":
            wt = geom_tube_bundle(d["tube_od_mm"], d["tube_thk_mm"], d["tube_length_mm"], d["n_tubes"], density,
                                   scrap=sc if sc is not None else 0.05)
        elif fn == "limpet":
            _, wt_total, length_m = geom_limpet_coil(
                d.get("shell_id_mm", 0), d.get("shell_thk_mm", 0), d.get("shell_ht_mm", 0),
                d.get("pipe_od_mm", 0), d.get("pipe_thk_mm", 0), d.get("pitch_mm", 80),
                cover_bottom_dish=bool(d.get("cover_bottom_dish", False)),
                density=density,
                scrap=sc if sc is not None else 0.10,
            )
            wt = wt_total / max(qty, 1)
            return round(wt, 3), round(wt_total, 3), qty
        else:
            wt = 0.0
    except Exception as e:
        st.warning(f"Weight calc error ({fn}): {e} | dims: {dims}")
        wt = 0.0
    return round(wt, 3), round(wt * used_qty, 3), used_qty

# ─────────────────────────────────────────────────────────────────────────────
# STRUCTURAL STEEL — Angles, Channels, Beams (IS standard unit weights kg/m)
# ─────────────────────────────────────────────────────────────────────────────
STRUCTURAL_SECTIONS = {
    # Equal Angles (size_mm × thk_mm) → kg/m
    "Angle 25x25x3":   {"type": "Equal Angle", "unit_wt": 1.11},
    "Angle 25x25x5":   {"type": "Equal Angle", "unit_wt": 1.80},
    "Angle 35x35x3":   {"type": "Equal Angle", "unit_wt": 1.60},
    "Angle 35x35x5":   {"type": "Equal Angle", "unit_wt": 2.60},
    "Angle 40x40x3":   {"type": "Equal Angle", "unit_wt": 1.84},
    "Angle 40x40x5":   {"type": "Equal Angle", "unit_wt": 2.95},
    "Angle 40x40x6":   {"type": "Equal Angle", "unit_wt": 3.50},
    "Angle 50x50x5":   {"type": "Equal Angle", "unit_wt": 3.77},
    "Angle 50x50x6":   {"type": "Equal Angle", "unit_wt": 4.47},
    "Angle 50x50x8":   {"type": "Equal Angle", "unit_wt": 5.80},
    "Angle 60x60x5":   {"type": "Equal Angle", "unit_wt": 4.57},
    "Angle 60x60x6":   {"type": "Equal Angle", "unit_wt": 5.42},
    "Angle 60x60x8":   {"type": "Equal Angle", "unit_wt": 7.09},
    "Angle 65x65x6":   {"type": "Equal Angle", "unit_wt": 5.80},
    "Angle 65x65x8":   {"type": "Equal Angle", "unit_wt": 7.70},
    "Angle 65x65x10":  {"type": "Equal Angle", "unit_wt": 9.42},
    "Angle 75x75x6":   {"type": "Equal Angle", "unit_wt": 6.85},
    "Angle 75x75x8":   {"type": "Equal Angle", "unit_wt": 8.99},
    "Angle 75x75x10":  {"type": "Equal Angle", "unit_wt": 11.00},
    "Angle 90x90x6":   {"type": "Equal Angle", "unit_wt": 8.20},
    "Angle 90x90x8":   {"type": "Equal Angle", "unit_wt": 10.80},
    "Angle 90x90x10":  {"type": "Equal Angle", "unit_wt": 13.40},
    "Angle 100x100x6": {"type": "Equal Angle", "unit_wt": 9.20},
    "Angle 100x100x8": {"type": "Equal Angle", "unit_wt": 12.10},
    "Angle 100x100x10":{"type": "Equal Angle", "unit_wt": 14.90},
    "Angle 100x100x12":{"type": "Equal Angle", "unit_wt": 17.70},

    # ISMC Channels → kg/m
    "ISMC 75":  {"type": "Channel", "unit_wt": 7.14},
    "ISMC 100": {"type": "Channel", "unit_wt": 9.56},
    "ISMC 125": {"type": "Channel", "unit_wt": 13.10},
    "ISMC 150": {"type": "Channel", "unit_wt": 16.40},
    "ISMC 175": {"type": "Channel", "unit_wt": 19.10},
    "ISMC 200": {"type": "Channel", "unit_wt": 22.10},

    # ISMB Beams → kg/m
    "ISMB 100": {"type": "Beam", "unit_wt": 11.50},
    "ISMB 125": {"type": "Beam", "unit_wt": 13.00},
    "ISMB 150": {"type": "Beam", "unit_wt": 14.90},
    "ISMB 175": {"type": "Beam", "unit_wt": 19.30},
    "ISMB 200": {"type": "Beam", "unit_wt": 25.40},
    "ISMB 225": {"type": "Beam", "unit_wt": 31.20},
    "ISMB 250": {"type": "Beam", "unit_wt": 37.30},
    "ISMB 300": {"type": "Beam", "unit_wt": 44.20},
    "ISMB 350": {"type": "Beam", "unit_wt": 52.40},
    "ISMB 400": {"type": "Beam", "unit_wt": 61.60},
    "ISMB 450": {"type": "Beam", "unit_wt": 72.40},
    "ISMB 500": {"type": "Beam", "unit_wt": 86.90},
}

STRUCT_DEFAULT_RATES = {
    "MS": 75.0,
    "SS": 280.0,
}

# ─────────────────────────────────────────────────────────────────────────────
# FABRICATION SERVICES ENGINE
# ─────────────────────────────────────────────────────────────────────────────
FAB_DEFAULTS = {
    "cutting_pct_on_plates": 2.0,
    "rolling_rate_per_m2": 800.0,
    "tig_weld_rate_per_m": 1200.0,
    "arc_weld_rate_per_m": 600.0,
    "int_grind_rate_per_m2": 350.0,
    "ext_buff_rate_per_m2": 250.0,
    "ep_rate_per_m2": 0.0,
    "hydro_test_lumpsum": 5000.0,
    "dp_test_rate_per_m2": 150.0,
    "assembly_fitting_hrs": 40.0,
    "assembly_rate_per_hr": 350.0,
    "qa_doc_lumpsum": 8000.0,
}

def calc_weld_metres(shell_id_mm, shell_ht_mm, n_nozzles=8, avg_nozzle_od_mm=100,
                     n_dish_ends=2, has_jacket=True, has_agitator=True):
    id_m = _m(shell_id_mm)
    ht_m = _m(shell_ht_mm)
    long_weld = ht_m
    n_courses = max(1, round(ht_m / 1.5))
    circ_weld = PI * id_m * (n_dish_ends + max(0, n_courses - 1))
    nozzle_weld = PI * _m(avg_nozzle_od_mm) * n_nozzles
    jacket_weld = PI * (id_m + 0.02) * 2.0 if has_jacket else 0.0
    agit_weld = PI * 0.25 * 2.0 if has_agitator else 0.0
    return round(long_weld + circ_weld + nozzle_weld + jacket_weld + agit_weld, 2)

def calc_surface_areas(shell_id_mm, shell_ht_mm, n_dish_ends=2):
    id_m = _m(shell_id_mm)
    ht_m = _m(shell_ht_mm)
    shell_area = PI * id_m * ht_m
    dish_area_each = 1.09 * PI * (_m(shell_id_mm * 1.167) / 2) ** 2
    dish_area_total = dish_area_each * n_dish_ends
    internal_area = shell_area + dish_area_total
    external_area = internal_area * 1.05
    return round(shell_area, 3), round(dish_area_total, 3), round(internal_area, 3), round(external_area, 3)

def auto_fab_services(h, fab_rates, parts):
    lines = []
    dia = float(h.get("shell_dia_mm", 0))
    ht = float(h.get("shell_ht_mm", 0))
    has_j = bool(h.get("jacket_type", ""))
    has_a = bool(h.get("agitator_type", ""))
    moc = h.get("moc_shell", "SS316L")
    if dia <= 0 or ht <= 0:
        return []
    shell_a, dish_a, int_a, ext_a = calc_surface_areas(dia, ht)
    weld_m = calc_weld_metres(dia, ht, has_jacket=has_j, has_agitator=has_a)
    plate_rm = sum(p.get("amount", 0) for p in parts)

    cr = fab_rates.get("cutting_pct_on_plates", FAB_DEFAULTS["cutting_pct_on_plates"])
    lines.append({"service": "Plate Cutting & Profiling", "basis": f"{cr}% on plate RM ₹{plate_rm:,.0f}", "qty": 1, "uom": "LS", "rate": plate_rm * cr / 100, "amount": round(plate_rm * cr / 100, 2)})

    rr = fab_rates.get("rolling_rate_per_m2", FAB_DEFAULTS["rolling_rate_per_m2"])
    lines.append({"service": "Plate Rolling / Shell Forming", "basis": f"Shell area {shell_a:.3f} m² × ₹{rr}/m²", "qty": round(shell_a, 3), "uom": "m²", "rate": rr, "amount": round(shell_a * rr, 2)})

    if moc in ("SS316L", "Ti", "C22", "Hastelloy"):
        wr = fab_rates.get("tig_weld_rate_per_m", FAB_DEFAULTS["tig_weld_rate_per_m"])
        wl = "TIG Welding (SS316L)"
    else:
        wr = fab_rates.get("arc_weld_rate_per_m", FAB_DEFAULTS["arc_weld_rate_per_m"])
        wl = "ARC / MIG Welding"
    lines.append({"service": wl, "basis": f"Est. weld {weld_m:.2f} m × ₹{wr}/m", "qty": weld_m, "uom": "m", "rate": wr, "amount": round(weld_m * wr, 2)})

    gr = fab_rates.get("int_grind_rate_per_m2", FAB_DEFAULTS["int_grind_rate_per_m2"])
    lines.append({"service": "Internal Grinding & Buffing (Ra 0.8)", "basis": f"Internal area {int_a:.3f} m² × ₹{gr}/m²", "qty": round(int_a, 3), "uom": "m²", "rate": gr, "amount": round(int_a * gr, 2)})

    er = fab_rates.get("ext_buff_rate_per_m2", FAB_DEFAULTS["ext_buff_rate_per_m2"])
    lines.append({"service": "External Buffing & Finishing", "basis": f"External area {ext_a:.3f} m² × ₹{er}/m²", "qty": round(ext_a, 3), "uom": "m²", "rate": er, "amount": round(ext_a * er, 2)})

    ah = fab_rates.get("assembly_fitting_hrs", FAB_DEFAULTS["assembly_fitting_hrs"])
    ar = fab_rates.get("assembly_rate_per_hr", FAB_DEFAULTS["assembly_rate_per_hr"])
    lines.append({"service": "Assembly, Fitting & Erection", "basis": f"{ah} hrs × ₹{ar}/hr", "qty": ah, "uom": "Hr", "rate": ar, "amount": round(ah * ar, 2)})

    hl = fab_rates.get("hydro_test_lumpsum", FAB_DEFAULTS["hydro_test_lumpsum"])
    lines.append({"service": "Hydrostatic Pressure Testing", "basis": "Lumpsum per ASME", "qty": 1, "uom": "LS", "rate": hl, "amount": round(hl, 2)})

    dr = fab_rates.get("dp_test_rate_per_m2", FAB_DEFAULTS["dp_test_rate_per_m2"])
    lines.append({"service": "Dye Penetration (DP) Testing", "basis": f"Area {ext_a:.3f} m² × ₹{dr}/m²", "qty": round(ext_a, 3), "uom": "m²", "rate": dr, "amount": round(ext_a * dr, 2)})

    ql = fab_rates.get("qa_doc_lumpsum", FAB_DEFAULTS["qa_doc_lumpsum"])
    lines.append({"service": "QA Dossier, MTC, Test Reports & Documentation", "basis": "Lumpsum", "qty": 1, "uom": "LS", "rate": ql, "amount": round(ql, 2)})

    epr = fab_rates.get("ep_rate_per_m2", 0)
    if epr > 0:
        lines.append({"service": "Electropolishing (Ra 0.4)", "basis": f"Internal area {int_a:.3f} m² × ₹{epr}/m²", "qty": round(int_a, 3), "uom": "m²", "rate": epr, "amount": round(int_a * epr, 2)})

    return lines

# ─────────────────────────────────────────────────────────────────────────────
# EQUIPMENT TYPES
# ─────────────────────────────────────────────────────────────────────────────
EQUIPMENT_TYPES = {
    "SSR — Stainless Steel Reactor":        {"icon": "⚗️",  "category": "Reactor",        "margin_hint": (12, 18), "labour_norm": "High",      "description": "SS316L reactor with jacket, agitator, mechanical seal"},
    "HAR — Hastelloy Reactor":              {"icon": "⚗️",  "category": "Reactor",        "margin_hint": (12, 18), "labour_norm": "High",      "description": "Hastelloy reactor with jacket, agitator, mechanical seal"},
    "RCVD — Rotary Cone Vacuum Dryer":      {"icon": "🌀",  "category": "Dryer",          "margin_hint": (14, 20), "labour_norm": "High",      "description": "Rotary cone vacuum dryer — jacketed, vacuum operation, integrated condenser"},
    "RD — Receiver / Decanter":            {"icon": "🧪",  "category": "Vessel",         "margin_hint": (10, 15), "labour_norm": "Medium",    "description": "SS316L receiver cum decanter vessel, no agitator"},
    "ANFD — Agitated Nutsche Filter Dryer": {"icon": "🔩",  "category": "Filter",         "margin_hint": (14, 20), "labour_norm": "Very High", "description": "ANFD with filter plate, agitator, jacket"},
    "VST — Vertical Storage Tank":          {"icon": "🛢️",  "category": "Storage",        "margin_hint": (10, 15), "labour_norm": "Low",       "description": "Plain vertical storage tank"},
    "HST — Horizontal Storage Tank":        {"icon": "🛢️",  "category": "Storage",        "margin_hint": (10, 15), "labour_norm": "Low",       "description": "Horizontal tank with saddle supports"},
    "PNF — Plain Nutsche Filter":           {"icon": "🔲",  "category": "Filter",         "margin_hint": (10, 15), "labour_norm": "Medium",    "description": "Plain nutsche filter, no agitator"},
    "Leaf Filter":                          {"icon": "🍃",  "category": "Filter",         "margin_hint": (12, 18), "labour_norm": "High",      "description": "Leaf filter vessel with filter leaves"},
    "Condenser":                            {"icon": "❄️",  "category": "Heat Exchanger", "margin_hint": (12, 18), "labour_norm": "High",      "description": "Shell and tube condenser"},
    "Reboiler":                             {"icon": "♨️",  "category": "Heat Exchanger", "margin_hint": (12, 18), "labour_norm": "High",      "description": "Kettle / thermosyphon reboiler"},
    "Tray Dryer":                           {"icon": "📦",  "category": "Dryer",          "margin_hint": (14, 20), "labour_norm": "High",      "description": "SS316L tray dryer with trays"},
    "Octagonal Blender":                    {"icon": "🔷",  "category": "Mixer",          "margin_hint": (15, 22), "labour_norm": "Very High", "description": "Octagonal blender / V-blender"},
    "Multi Miller":                         {"icon": "⚙️",  "category": "Powder",         "margin_hint": (15, 22), "labour_norm": "Very High", "description": "Multi mill / sifter"},
    "Distillation Column":                  {"icon": "🏛️",  "category": "Column",         "margin_hint": (14, 20), "labour_norm": "High",      "description": "Distillation column with trays or packing"},
    "ATFD — Agitated Thin Film Dryer":      {"icon": "🌀",  "category": "Dryer",          "margin_hint": (18, 25), "labour_norm": "Very High", "description": "ATFD / thin film evaporator"},
    "MEE — Multiple Effect Evaporator":     {"icon": "💧",  "category": "Evaporator",     "margin_hint": (16, 22), "labour_norm": "Very High", "description": "Multiple effect evaporator package"},
    "Rectangular Tank":                     {"icon": "📐",  "category": "Storage",        "margin_hint": (10, 15), "labour_norm": "Low",       "description": "Rectangular / sump tank"},
    "Skid / Package":                       {"icon": "🏗️",  "category": "Package",        "margin_hint": (14, 20), "labour_norm": "High",      "description": "Multi-equipment skid package"},
    "Custom Equipment":                     {"icon": "🔧",  "category": "Custom",         "margin_hint": (10, 20), "labour_norm": "Medium",    "description": "User-defined — all fields manual"},
}
EQUIPMENT_NAMES = list(EQUIPMENT_TYPES.keys())

# ─────────────────────────────────────────────────────────────────────────────
# COST ENGINE
# ─────────────────────────────────────────────────────────────────────────────
def calc_shell_area(dia_mm, ht_mm):
    return PI * _m(dia_mm) * _m(ht_mm)

def calc_dish_area(shell_id_mm):
    r = _m(shell_id_mm * 1.167) / 2
    return 1.09 * PI * r * r

def calc_shell_volume_ltrs(dia_mm, ht_mm):
    return PI * (_m(dia_mm) / 2) ** 2 * _m(ht_mm) * 1000

def calc_totals(parts, pipes, flanges, struct, fab_services, bo_items, oh_items,
                profit_pct, contingency_pct, packing, freight, gst_pct, engg_design,
                discount_pct=0.0):
    """
    Compute all cost totals for an estimation.

    Option B cost structure (internal):
      - Splits OH into Factory OH (shop floor) and Admin OH (office/docs)
      - Adds Discount on Ex-Works price
      - Adds Gross Profit line
      - Returns drill-down line lists for Tab 6 expandable view

    Backward compatibility: all original keys are preserved so existing UI,
    DOCX generator and Excel fact sheet continue to work without changes.
    """
    # ── Direct Material totals (unchanged) ───────────────────────────────────
    tot_plates  = sum(p.get("amount", 0) for p in parts)
    tot_pipes   = sum(p.get("amount", 0) for p in pipes)
    tot_flanges = sum(p.get("amount", 0) for p in flanges)
    tot_struct  = sum(p.get("amount", 0) for p in struct)
    tot_rm      = tot_plates + tot_pipes + tot_flanges + tot_struct

    # ── Fabrication services & Bought-out (unchanged) ────────────────────────
    tot_fab = sum(f.get("amount", 0) for f in fab_services)
    tot_bo  = sum(p.get("amount", 0) for p in bo_items)

    # ── Overhead split: Factory vs Admin (NEW for Option B) ──────────────────
    # Walk each OH line, look up its oh_type in OH_BUCKET, classify, and
    # collect both the totals AND the line items themselves for drill-down.
    oh_lines_factory = []
    oh_lines_admin   = []
    for o in oh_items:
        bucket = OH_BUCKET.get(o.get("oh_type", ""), OH_BUCKET_DEFAULT)
        if bucket == "ADMIN_OH":
            oh_lines_admin.append(o)
        else:
            oh_lines_factory.append(o)

    tot_factory_oh = sum(o.get("amount", 0) for o in oh_lines_factory)
    tot_admin_oh   = sum(o.get("amount", 0) for o in oh_lines_admin)
    tot_oh         = tot_factory_oh + tot_admin_oh   # keep for back-compat

    # ── Legacy sub-totals (kept so old display code still works) ─────────────
    # These were used by the old margin_issues() check and by the Excel sheet.
    tot_lab   = sum(o.get("amount", 0) for o in oh_items
                    if o.get("oh_type") in ("LABOUR", "LABOUR_BUFF"))
    tot_cons  = sum(o.get("amount", 0) for o in oh_items
                    if o.get("oh_type") == "CONSUMABLES")
    tot_other = sum(o.get("amount", 0) for o in oh_items
                    if o.get("oh_type") not in ("LABOUR", "LABOUR_BUFF", "CONSUMABLES"))

    # ── Manufacturing cost build-up ──────────────────────────────────────────
    tot_mfg    = tot_rm + tot_fab + tot_bo + tot_oh + engg_design
    cont_amt   = tot_mfg * contingency_pct / 100
    cbm        = tot_mfg + cont_amt   # Cost Base for Margin

    # ── Profit & Ex-Works ────────────────────────────────────────────────────
    profit_amt = cbm * profit_pct / 100
    ex_works   = cbm + profit_amt + packing + freight

    # ── NEW: Discount applied on Ex-Works ────────────────────────────────────
    discount_amt    = ex_works * discount_pct / 100
    net_realisation = ex_works - discount_amt   # what B&G actually keeps

    # ── NEW: Gross Profit ────────────────────────────────────────────────────
    # Definition: Net Realisation − Total Mfg Cost − Contingency
    # i.e. what's left after all manufacturing costs, before tax
    gross_profit = net_realisation - tot_mfg - cont_amt

    # ── GST & FOR ────────────────────────────────────────────────────────────
    # GST is charged on Net Realisation (after discount), per standard practice
    gst_amt   = net_realisation * gst_pct / 100
    for_price = net_realisation + gst_amt

    # ── Margin health percentages ────────────────────────────────────────────
    safe = ex_works if ex_works else 1
    return dict(
        # Direct Material
        tot_plates=tot_plates, tot_pipes=tot_pipes, tot_flanges=tot_flanges,
        tot_struct=tot_struct, tot_rm=tot_rm,
        # Conversion / Fab
        tot_fab=tot_fab,
        # Bought-out
        tot_bo=tot_bo,
        # Overheads — both old and new keys present
        tot_lab=tot_lab, tot_cons=tot_cons, tot_other=tot_other,
        tot_oh=tot_oh,
        tot_factory_oh=tot_factory_oh,                # NEW
        tot_admin_oh=tot_admin_oh,                    # NEW
        oh_lines_factory=oh_lines_factory,            # NEW — for drill-down
        oh_lines_admin=oh_lines_admin,                # NEW — for drill-down
        # Engineering
        engg_design=engg_design,
        # Manufacturing build-up
        tot_mfg=tot_mfg, cont_amt=cont_amt, cbm=cbm,
        profit_amt=profit_amt, packing=packing, freight=freight,
        ex_works=ex_works,
        # NEW: Discount & Net Realisation
        discount_pct=discount_pct,                    # NEW
        discount_amt=discount_amt,                    # NEW
        net_realisation=net_realisation,              # NEW
        gross_profit=gross_profit,                    # NEW
        # GST & FOR
        gst_amt=gst_amt, for_price=for_price,
        # Margin percentages (against Ex-Works as before)
        rm_pct=tot_rm / safe * 100,
        fab_pct=tot_fab / safe * 100,
        lab_pct=tot_lab / safe * 100,
        oh_pct=(tot_cons + tot_other) / safe * 100,
        profit_pct_actual=profit_amt / safe * 100,
        # NEW: percentages for the two new buckets
        factory_oh_pct=tot_factory_oh / safe * 100,
        admin_oh_pct=tot_admin_oh / safe * 100,
        gross_profit_pct=gross_profit / safe * 100,
    )
def margin_issues(t):
    out = []
    if not (45 <= t["rm_pct"]  <= 60): out.append(f"RM {t['rm_pct']:.1f}% — target 45–60%")
    if not (15 <= t["fab_pct"] <= 25): out.append(f"Fabrication {t['fab_pct']:.1f}% — target 15–25%")
    if not (8  <= t["oh_pct"]  <= 15): out.append(f"OH {t['oh_pct']:.1f}% — target 8–15%")
    if t["profit_pct_actual"] < 12:    out.append(f"Profit {t['profit_pct_actual']:.1f}% — min 12%")
    return out

# ─────────────────────────────────────────────────────────────────────────────
# DOCX GENERATOR
# ─────────────────────────────────────────────────────────────────────────────
def _shd(cell, hex_color):
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(s)

def _run(para, text, bold=False, size=9, color=None):
    r = para.add_run(text)
    r.font.size = Pt(size); r.font.name = "Arial"; r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r

def _kv_table(doc, rows):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        row = t.add_row()
        for j, txt in enumerate([k, v]):
            c = row.cells[j]; c.text = ""
            _run(c.paragraphs[0], str(txt), bold=(j == 0), size=9)
            _shd(c, ("D6E4F7" if i % 2 == 0 else "F2F2F2") if j == 0 else ("FFFFFF" if i % 2 == 0 else "EFF5FB"))
def _spec_sub_header(doc, text):
    """Render a small bold sub-header within the Spec Sheet."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    _run(p, text, bold=True, size=10, color=(27, 58, 107))

def _spec_sheet_block(doc, est):
    """
    Render the Equipment Specification Sheet — replaces the old flat
    key-value Section 2. Lays out 6 sub-blocks; rows with empty values
    are skipped automatically so e.g. storage tanks don't show empty
    Agitator / Seal rows.
    """
    def _add(rows, label, value, suffix=""):
        """Add a row to the buffer only if value is non-empty/non-zero."""
        if value is None:
            return
        s = str(value).strip()
        if not s or s in ("0", "0.0"):
            return
        rows.append((label, f"{s}{suffix}"))

    # ── Block 1: General ─────────────────────────────────────────────────────
    _spec_sub_header(doc, "General")
    rows = []
    _add(rows, "Equipment Description", est.get("equipment_desc", ""))
    _add(rows, "Equipment Type",        est.get("equipment_type", ""))
    _add(rows, "Tag Number",            est.get("tag_number", ""))
    _add(rows, "Quantity",              "1 No.")
    _add(rows, "Service / Process",     est.get("service_fluid", ""))
    _add(rows, "Design Code",           est.get("design_code", "ASME Sec VIII Div 1"))
    if rows:
        _kv_table(doc, rows)

    # ── Block 2: Vessel Dimensions ───────────────────────────────────────────
    _spec_sub_header(doc, "Vessel Dimensions")
    rows = []
    _add(rows, "Gross Capacity",   est.get("capacity_ltrs", ""), " Litres")
    # Working volume — defaults to 80% of capacity if not explicitly set
    wv = est.get("working_vol_ltrs", 0) or 0
    if wv == 0:
        cap = float(est.get("capacity_ltrs", 0) or 0)
        if cap > 0:
            wv = round(cap * 0.8, 0)
    if wv:
        rows.append(("Working Volume", f"{wv:.0f} Litres (approx)"))
    _add(rows, "Shell Internal Diameter", est.get("shell_dia_mm", ""), " mm")
    _add(rows, "Shell Height (T/T)",      est.get("shell_ht_mm", ""), " mm")
    _add(rows, "Shell Thickness",         est.get("shell_thk_mm", ""), " mm")
    _add(rows, "Dish End Thickness",      est.get("dish_thk_mm", ""), " mm")
    _add(rows, "Dish End Type",           "Torispherical (Crown radius = ID)")
    if rows:
        _kv_table(doc, rows)

    # ── Block 3: Design Conditions ───────────────────────────────────────────
    _spec_sub_header(doc, "Design Conditions")
    rows = []
    # Prefer new split field; fall back to old combined field for safety
    dps = est.get("design_pressure_shell") or est.get("design_pressure", "")
    _add(rows, "Design Pressure — Shell",  dps)
    _add(rows, "Design Pressure — Jacket", est.get("design_pressure_jacket", ""))
    _add(rows, "Design Temperature",       est.get("design_temp", ""))
    _add(rows, "Hydrotest Pressure",       est.get("hydrotest_pressure", ""))
    ca = est.get("corrosion_allowance_mm", None)
    if ca is not None and float(ca) > 0:
        rows.append(("Corrosion Allowance", f"{float(ca):.1f} mm"))
    je = est.get("joint_efficiency", None)
    if je is not None and float(je) > 0:
        rows.append(("Joint Efficiency", f"{float(je):.2f}"))
    if rows:
        _kv_table(doc, rows)

    # ── Block 4: Materials of Construction ───────────────────────────────────
    _spec_sub_header(doc, "Materials of Construction")
    rows = []
    _add(rows, "Shell & Dish Ends",     est.get("moc_shell", "SS316L"))
    _add(rows, "Jacket / Limpet",       est.get("moc_jacket", "SS304"))
    _add(rows, "Nozzles & Manholes",    est.get("moc_shell", "SS316L"))
    _add(rows, "Flanges & Forgings",    f"{est.get('moc_shell', 'SS316L')} (Forged)")
    _add(rows, "Fasteners",             "SS304 / SS316 (as per design)")
    _add(rows, "Gaskets",               "PTFE / Spiral wound (per service)")
    if rows:
        _kv_table(doc, rows)

    # ── Block 5: Heating / Cooling System ────────────────────────────────────
    # Only print this block if a jacket type is specified
    if (est.get("jacket_type") or "").strip():
        _spec_sub_header(doc, "Heating / Cooling System")
        rows = []
        _add(rows, "Jacket / Heating Type", est.get("jacket_type", ""))
        _add(rows, "Heating Medium",        est.get("heating_medium", ""))
        _add(rows, "Cooling Medium",        est.get("cooling_medium", ""))
        if rows:
            _kv_table(doc, rows)

    # ── Block 6: Agitator & Drive Assembly ───────────────────────────────────
    # Only print this block if an agitator type is specified
    if (est.get("agitator_type") or "").strip():
        _spec_sub_header(doc, "Agitator & Drive Assembly")
        rows = []
        _add(rows, "Agitator Type",       est.get("agitator_type", ""))
        _add(rows, "Agitator Speed",      est.get("agitator_rpm", ""))
        _add(rows, "Motor Rating",        est.get("motor_hp", ""))
        _add(rows, "Motor Make",          est.get("motor_make", ""))
        _add(rows, "Gearbox Make",        est.get("gearbox_make", ""))
        _add(rows, "Gearbox Ratio",       est.get("gearbox_ratio", ""))
        _add(rows, "Mechanical Seal",     est.get("seal_type", ""))
        _add(rows, "Seal Make",           est.get("seal_make", ""))
        if rows:
            _kv_table(doc, rows)

    # ── Block 7: Surface Finish & Testing ────────────────────────────────────
    _spec_sub_header(doc, "Surface Finish & Testing")
    rows = []
    _add(rows, "Surface Finish",   est.get("surface_finish", ""))
    _add(rows, "Hydrotest",        "Conducted per ASME Sec VIII Div 1")
    _add(rows, "DP Testing",       "All critical welds")
    _add(rows, "PMI",              "100% on pressure parts")
    _add(rows, "MTC",              "EN 10204 Type 3.1 for pressure parts")
    if rows:
        _kv_table(doc, rows)
def generate_docx(est, customer, totals, fab_services, show_breakup=False):
    """
    Generate customer-facing quotation DOCX.
    show_breakup=False  →  Clean quote: one price, Ex-Works + GST + FOR only.
    show_breakup=True   →  Scope-based breakup: groups costs into 4 customer-friendly
                           scope heads (NOT internal cost breakdown).
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

    def _add_header_to_all_pages():
        import io as _bio2
        from docx.oxml.ns import qn as _qn2
        from docx.oxml import OxmlElement as _OE2

        for section in doc.sections:
            section.different_first_page_header_footer = False
            header = section.header
            header.is_linked_to_previous = False

            for p in header.paragraphs:
                p.clear()

            # Width matches the page content width (21cm − 2cm margin × 2 = 17cm)
            # so the header aligns flush with the body text below it.
            hdr_tbl = header.add_table(rows=1, cols=2 if _LOGO_BYTES else 1,
                                        width=Cm(17.0))
            hdr_tbl.style = "Table Grid"
            from docx.oxml.ns import qn as _qn4
            from docx.oxml import OxmlElement as _OE4
            tblPr = hdr_tbl._tbl.tblPr
            if tblPr is None:
                tblPr = _OE4("w:tblPr"); hdr_tbl._tbl.insert(0, tblPr)
            tblBdr = _OE4("w:tblBorders")
            for side in ["top","bottom","left","right","insideH","insideV"]:
                el = _OE4(f"w:{side}")
                el.set(_qn4("w:val"), "none")
                el.set(_qn4("w:sz"), "0")
                el.set(_qn4("w:space"), "0")
                el.set(_qn4("w:color"), "auto")
                tblBdr.append(el)
            tblPr.append(tblBdr)

            if _LOGO_BYTES and len(hdr_tbl.columns) >= 2:
                from docx.oxml.ns import qn as _qn3
                from docx.oxml import OxmlElement as _OE3
                from docx.shared import Cm as _Cm3
                tblGrid = hdr_tbl._tbl.find(_qn3("w:tblGrid"))
                if tblGrid is None:
                    tblGrid = _OE3("w:tblGrid"); hdr_tbl._tbl.insert(0, tblGrid)
                for col_el in tblGrid.findall(_qn3("w:gridCol")):
                    tblGrid.remove(col_el)
                # Logo cell 2.8 cm + Address cell 14.2 cm = 17.0 cm total
                # (matches the hdr_tbl width above, which matches page body)
                for w_twips in [int(_Cm3(2.8).twips), int(_Cm3(14.2).twips)]:
                    gc = _OE3("w:gridCol"); gc.set(_qn3("w:w"), str(w_twips)); tblGrid.append(gc)

            lhc = hdr_tbl.rows[0].cells[0]
            _shd(lhc, "FFFFFF" if _LOGO_BYTES else "1B3A6B")
            lhc.paragraphs[0].clear()
            lhp = lhc.paragraphs[0]
            lhp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lhp.paragraph_format.space_before = Pt(2)
            lhp.paragraph_format.space_after  = Pt(2)

            if _LOGO_BYTES:
                try:
                    run = lhp.add_run()
                    run.add_picture(_bio2.BytesIO(_LOGO_BYTES), width=Cm(2.5))
                except Exception:
                    _run(lhp, "B&G", bold=True, size=10, color=(27,58,107))

                rhc = hdr_tbl.rows[0].cells[1]
                _shd(rhc, "1B3A6B")
                rhc.paragraphs[0].clear()
                rp = rhc.paragraphs[0]
            else:
                rp = lhp

            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rp.paragraph_format.space_before = Pt(4)
            rp.paragraph_format.space_after  = Pt(2)
            _run(rp, f"{BG_NAME}  |  {BG_TAGLINE}  |  {BG_PHONE}\n",
                 bold=True, size=9, color=(255,255,255))
            _run(rp, f"{BG_EMAIL}  |  {BG_WEB}  |  GSTIN: {BG_GSTIN}",
                 size=8, color=(180,210,255))

            for para in header.paragraphs:
                pPr = para._p.get_or_add_pPr()
                pBdr = _OE2("w:pBdr")
                for side in ["top","bottom","left","right"]:
                    el = _OE2(f"w:{side}")
                    el.set(_qn2("w:val"), "none")
                    el.set(_qn2("w:sz"), "0")
                    el.set(_qn2("w:space"), "0")
                    el.set(_qn2("w:color"), "auto")
                    pBdr.append(el)
                pPr.append(pBdr)

    def banner():
        _add_header_to_all_pages()

    def footer_block():
        doc.add_paragraph()
        t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
        c = t.rows[0].cells[0]; _shd(c, "1B3A6B"); c.paragraphs[0].clear()
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        _run(p, f"{BG_NAME}  |  {BG_TAGLINE}\n", bold=True, size=9, color=(255,255,255))
        _run(p, f"{BG_ADDRESS}\n", size=8, color=(180,210,255))
        _run(p, f"Ph: {BG_PHONE}  |  {BG_EMAIL}  |  {BG_WEB}\n", size=8, color=(180,210,255))
        _run(p, f"GSTIN: {BG_GSTIN}  |  PAN: {BG_PAN}", bold=True, size=8, color=(255,255,255))

    def sec_head(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
        _run(p, text, bold=True, size=12, color=(27,58,107))
        pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
        bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"2E75B6")
        pBdr.append(bot); p._p.get_or_add_pPr().append(pBdr)

    def body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        _run(p, text, size=9, color=(68,68,68))

    def bullet(text):
        p = doc.add_paragraph(text, style="List Bullet")
        for r in p.runs: r.font.size = Pt(9); r.font.name = "Arial"

    def price_table(rows_data):
        t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
        for i, (label, value, highlight) in enumerate(rows_data):
            row = t.add_row()
            for j, txt in enumerate([label, value]):
                c = row.cells[j]; c.text = ""
                _run(c.paragraphs[0], txt, bold=highlight, size=10 if highlight else 9,
                     color=(255,255,255) if highlight else (26,26,26))
                _shd(c, "1B3A6B" if highlight else ("F7FBFF" if i%2==0 else "FFFFFF"))
                if j == 1:
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    fmt  = lambda n: f"₹{n:,.0f}"
    cust = customer or {}

    # ── SECTION 1: Offer & Customer ───────────────────────────────────────────
    banner(); doc.add_paragraph()
    sec_head("SECTION 1 — OFFER & CUSTOMER DETAILS")
    _kv_table(doc, [
        ("Offer Reference",  est.get("qtn_number","")),
        ("Revision",         est.get("revision","R0")),
        ("Date",             date.today().strftime("%d %B %Y")),
        ("Equipment Type",   est.get("equipment_type","")),
        ("Prepared By",      est.get("prepared_by","")),
        ("Checked By",       est.get("checked_by","")),
        ("",""),
        ("Customer Name",    cust.get("name","")),
        ("Customer Address", cust.get("address","")),
        ("Customer GSTIN",   cust.get("gstin","")),
        ("Contact Person",   cust.get("contact_person","")),
        ("Phone / Email",    f"{cust.get('phone','') or ''} / {cust.get('email','') or ''}"),
        ("",""),
        ("Supplier",         BG_NAME),
        ("Supplier Address", BG_ADDRESS),
        ("Supplier GSTIN",   BG_GSTIN),
        ("Phone",            BG_PHONE),
        ("Email / Web",      f"{BG_EMAIL}  |  {BG_WEB}"),
    ])

    # ── SECTION 2: Equipment Specification Sheet ──────────────────────────────
    # Replaces the old flat key-value list with a proper engineering data sheet
    # organised into 7 sub-blocks (General, Dimensions, Design Conditions,
    # MOC, Heating/Cooling, Agitator/Drive, Surface Finish & Testing).
    doc.add_paragraph()
    sec_head("SECTION 2 — EQUIPMENT SPECIFICATION SHEET")
    _spec_sheet_block(doc, est)

    def bullets_from_text(text_block):
        for line in (text_block or "").split("\n"):
            line = line.strip()
            if line:
                bullet(line)

    # ── SECTION 3: Scope ──────────────────────────────────────────────────────
    doc.add_paragraph()
    sec_head("SECTION 3 — SCOPE OF SUPPLY")
    body("Supply of one (1) complete fabricated equipment per the technical basis above:")
    bullets_from_text(est.get("scope_items",
        "Pressure vessel / equipment fabricated as per approved GA drawing\n"
        "All nozzles, manholes and process connections per nozzle schedule\n"
        "Jacket / limpet coil as specified\n"
        "Agitator complete with gearbox, motor and mechanical seal (where applicable)\n"
        "Support structure — lugs, legs or saddles as applicable\n"
        "Internal grinding and buffing to specified Ra surface finish\n"
        "Equipment nameplate with tag number and serial number"
    ))
    excl = (est.get("scope_exclusions","") or "").strip()
    if excl:
        doc.add_paragraph()
        body("Not in scope:")
        bullets_from_text(excl)

    # ── SECTION 4: Quality ────────────────────────────────────────────────────
    doc.add_paragraph()
    sec_head("SECTION 4 — MANUFACTURING & QUALITY ASSURANCE")
    body(est.get("quality_intro",
        "B&G Engineering Industries operates as an engineering-led manufacturer. "
        "Every project is built to ASME Section VIII Division 1 requirements "
        "with full documentation and traceability."
    ))
    bullets_from_text(est.get("quality_points",
        "Raw material procurement with original Mill Test Certificates (MTC) for all pressure parts\n"
        "100% Positive Material Identification (PMI) verification before cutting\n"
        "Heat number and cast number traceability maintained throughout fabrication\n"
        "Qualified welders — WPS / PQR compliant, TIG welding for all SS316L pressure joints\n"
        "Precision plasma / laser cutting and CNC-controlled plate rolling\n"
        "Pharma-grade internal grinding to Ra ≤ 0.8 μm and external mechanical buffing\n"
        "Dimensional inspection against approved GA drawings at each stage\n"
        "Hydrostatic / pneumatic / vacuum leak test as per ASME Sec VIII Div 1\n"
        "Dye Penetrant (DP) testing on all critical welds\n"
        "Complete QA dossier prepared and delivered with equipment\n"
        "Factory Acceptance Test (FAT) support at works on request"
    ))

    # ── SECTION 5: Documentation ──────────────────────────────────────────────
    doc.add_paragraph()
    sec_head("SECTION 5 — DOCUMENTATION DELIVERABLES")
    body("The following documents are included in scope and delivered with the equipment:")
    bullets_from_text(est.get("doc_deliverables",
        "General Arrangement (GA) Drawing — IFC revision\n"
        "Nozzle orientation and schedule drawing\n"
        "Bill of Materials (BOM)\n"
        "Mill Test Certificates (MTC) for all pressure parts\n"
        "PMI verification records\n"
        "Weld map and weld log\n"
        "DP / RT inspection reports\n"
        "Dimensional inspection report\n"
        "Hydrostatic / leak test certificate\n"
        "Surface finish inspection record (Ra measurement)\n"
        "Inspection and Release Note (IRN)\n"
        "Equipment nameplate photograph"
    ))

    # ── SECTION 6: Commercial ─────────────────────────────────────────────────
    doc.add_paragraph()
    sec_head("SECTION 6 — COMMERCIAL OFFER")

    if show_breakup:
        ex = totals["ex_works"]
        vessel_pct  = 0.68
        drive_pct   = 0.18
        testing_pct = 0.08
        engg_pct    = 0.06

        vessel_amt  = round(ex * vessel_pct,  0)
        drive_amt   = round(ex * drive_pct,   0)
        testing_amt = round(ex * testing_pct, 0)
        engg_amt    = round(ex - vessel_amt - drive_amt - testing_amt, 0)

        body("Scope-based price breakup (for your reference):")
        doc.add_paragraph()
        scope_rows = [
            ("A. Pressure Vessel, Jacket, Nozzles & Structural",
             fmt(vessel_amt), False),
            ("B. Mechanical Drive, Seal, Gearbox & Motor Assembly",
             fmt(drive_amt), False),
            ("C. Surface Finishing, Testing & Third-Party Inspection",
             fmt(testing_amt), False),
            ("D. Engineering, Drawings, Documentation & QA Dossier",
             fmt(engg_amt), False),
            ("Ex-Works Price — Hyderabad (A+B+C+D)",
             fmt(totals["ex_works"]), True),
            (f"GST @ {est.get('gst_pct',18):.0f}%",
             fmt(totals["gst_amt"]), False),
            ("FINAL FOR PRICE",
             fmt(totals["for_price"]), True),
        ]
        price_table(scope_rows)
        doc.add_paragraph()
        body("Note: The above scope breakup is provided for procurement allocation purposes. "
             "B&G Engineering supplies the complete equipment as a single integrated scope — "
             "partial scope orders are not accepted.")
    else:
        clean_rows = [
            ("Equipment Description",
             est.get("equipment_desc",""), False),
            ("Quantity",
             "1 No.", False),
            ("Ex-Works Price — Hyderabad",
             fmt(totals["ex_works"]), True),
            (f"GST @ {est.get('gst_pct',18):.0f}% (if applicable)",
             fmt(totals["gst_amt"]), False),
            ("FINAL FOR PRICE (inclusive of GST)",
             fmt(totals["for_price"]), True),
        ]
        price_table(clean_rows)

    doc.add_paragraph()

    terms_rows = []
    def _add(label, key, fallback):
        val = (est.get(key,"") or "").strip() or fallback
        if val:
            terms_rows.append((label, val))

    _add("Price Basis",          "price_basis",
         "Ex-Works, Pashamylaram, Hyderabad — 502307. Packing in MS crate included. Freight, insurance and unloading at site excluded.")
    _add("GST & Statutory Levies","gst_clause",
         "GST @ 18% (HSN 8419) as applicable at time of invoicing. Any new statutory levy introduced after offer date will be charged additionally.")
    _add("Payment Terms",         "payment_terms",
         "40% advance along with Purchase Order  |  50% against Pro-forma invoice on readiness for dispatch  |  10% on delivery")

    dw = (est.get("delivery_weeks","") or "12–16").strip()
    dn = (est.get("delivery_note","") or "Subject to availability of raw material at time of order.").strip()
    terms_rows.append(("Delivery Period", f"{dw} weeks from date of Purchase Order + advance payment. {dn}"))

    _add("Offer Validity",        "offer_validity",
         "This offer is valid for 7 calendar days from the date above. Prices subject to change if raw material rates move by more than 3%.")
    _add("Warranty",              "warranty_clause",
         "12 months from date of commissioning or 18 months from date of dispatch, whichever is earlier.")
    _add("Inspection Rights",     "inspection_clause",
         "Customer may depute inspector for stage and final inspection at our works. TPI charges, if any, are in customer scope.")

    excl2 = (est.get("scope_exclusions","") or "").strip()
    if excl2:
        excl_inline = "  |  ".join([l.strip() for l in excl2.split("\n") if l.strip()])
        terms_rows.append(("Not in Scope", excl_inline))

    _kv_table(doc, terms_rows)

    sn = (est.get("special_notes","") or "").strip()
    if sn:
        doc.add_paragraph()
        body(f"Additional Notes: {sn}")

    # ── SECTION 7: Sign-off ───────────────────────────────────────────────────
    doc.add_paragraph()
    sec_head("SECTION 7 — ACCEPTANCE & SIGN-OFF")
    body("We thank you for the opportunity to offer our services and look forward to your valued order. "
         "Please feel free to contact us for any technical or commercial clarifications.")
    doc.add_paragraph()

    t_sign = doc.add_table(rows=3, cols=2); t_sign.style = "Table Grid"
    for j, lbl in enumerate(["For B&G Engineering Industries", "Customer Acceptance"]):
        c = t_sign.rows[0].cells[j]; c.text = ""
        _run(c.paragraphs[0], lbl, bold=True, size=9, color=(255,255,255)); _shd(c, "1B3A6B")
    for j, name in enumerate([est.get("prepared_by",""), ""]):
        c = t_sign.rows[1].cells[j]; c.text = ""
        _run(c.paragraphs[0], f"{'Authorised Signatory: ' if j==0 else 'Authorised Signatory: '}{name}", size=9)
    for j, txt in enumerate([f"Date: {date.today().strftime('%d %B %Y')}", "Date & Company Stamp:"]):
        c = t_sign.rows[2].cells[j]; c.text = ""
        _run(c.paragraphs[0], txt, size=9, color=(100,100,100))

    footer_block()
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────────────────
# ESTIMATION FACT SHEET — XLSX
# ─────────────────────────────────────────────────────────────────────────────
def generate_fact_sheet_xlsx(est, parts, pipes, flanges, struct_items, fab_services,
                              bo_items, oh_items, fab_rates, totals):
    if not OPENPYXL_OK:
        raise ImportError("openpyxl not installed. Add 'openpyxl' to requirements.txt and redeploy.")
    import io as _io

    wb = openpyxl.Workbook()

    DARK_BLUE  = "1B3A6B"
    MID_BLUE   = "2E75B6"
    LIGHT_BLUE = "D6E4F7"
    ALT_ROW    = "EFF5FB"
    GREEN_BG   = "E2EFDA"
    AMBER_BG   = "FFF2CC"
    RED_BG     = "FFE0E0"
    WHITE      = "FFFFFF"

    def hdr_style(cell, bg=DARK_BLUE, fg="FFFFFF", sz=10, bold=True):
        cell.font      = Font(name="Arial", bold=bold, size=sz, color=fg)
        cell.fill      = PatternFill("solid", fgColor=bg)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border    = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"),  bottom=Side(style="thin"),
        )

    def data_style(cell, bg=WHITE, bold=False, align="left", sz=9, color="000000"):
        cell.font      = Font(name="Arial", size=sz, bold=bold, color=color)
        cell.fill      = PatternFill("solid", fgColor=bg)
        cell.alignment = Alignment(horizontal=align, vertical="center")
        cell.border    = Border(
            left=Side(style="hair"), right=Side(style="hair"),
            top=Side(style="hair"),  bottom=Side(style="hair"),
        )

    def section_title(ws, row, text, ncols=10):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
        c = ws.cell(row=row, column=1, value=text)
        hdr_style(c, bg=MID_BLUE, sz=11)
        ws.row_dimensions[row].height = 20

    def kv_row(ws, row, label, value, bg=WHITE):
        lc = ws.cell(row=row, column=1, value=label)
        vc = ws.cell(row=row, column=2, value=value)
        data_style(lc, bg=LIGHT_BLUE, bold=True)
        data_style(vc, bg=bg)
        ws.row_dimensions[row].height = 15

    fmt_inr = '#,##0.00'
    fmt_kg  = '#,##0.000'

    # SHEET 1 — SUMMARY
    ws1 = wb.active; ws1.title = "Summary"
    ws1.column_dimensions["A"].width = 30
    ws1.column_dimensions["B"].width = 40

    r = 1
    ws1.merge_cells(start_row=r, start_column=1, end_row=r, end_column=8)
    c = ws1.cell(row=r, column=1, value="B&G ENGINEERING INDUSTRIES — ESTIMATION FACT SHEET")
    hdr_style(c, sz=14); ws1.row_dimensions[r].height = 28; r += 1

    ws1.merge_cells(start_row=r, start_column=1, end_row=r, end_column=8)
    c = ws1.cell(row=r, column=1, value="INTERNAL DOCUMENT — NOT FOR CUSTOMER DISTRIBUTION")
    hdr_style(c, bg=RED_BG, fg="CC0000", sz=9, bold=True); r += 2

    section_title(ws1, r, "OFFER DETAILS", 8); r += 1
    for label, value in [
        ("Quotation Number",  est.get("qtn_number","")),
        ("Revision",          est.get("revision","R0")),
        ("Date",              date.today().strftime("%d-%b-%Y")),
        ("Customer",          est.get("customer_name","")),
        ("Equipment Type",    est.get("equipment_type","")),
        ("Description",       est.get("equipment_desc","")),
        ("Tag Number",        est.get("tag_number","")),
        ("Status",            est.get("status","Draft")),
        ("Prepared By",       est.get("prepared_by","")),
        ("Checked By",        est.get("checked_by","")),
    ]:
        kv_row(ws1, r, label, value); r += 1

    r += 1
    section_title(ws1, r, "EQUIPMENT PARAMETERS", 8); r += 1
    for label, value in [
        ("Shell ID (mm)",        est.get("shell_dia_mm","")),
        ("Shell Height (mm)",    est.get("shell_ht_mm","")),
        ("Shell Thickness (mm)", est.get("shell_thk_mm","")),
        ("Dish Thickness (mm)",  est.get("dish_thk_mm","")),
        ("Capacity (Ltrs)",      est.get("capacity_ltrs","")),
        ("MOC — Shell",          est.get("moc_shell","SS316L")),
        ("MOC — Jacket",         est.get("moc_jacket","SS304")),
        ("Jacket Type",          est.get("jacket_type","")),
        ("Agitator Type",        est.get("agitator_type","")),
        ("Design Code",          est.get("design_code","ASME Sec VIII Div 1")),
        ("Design Pressure",      est.get("design_pressure","")),
        ("Design Temperature",   est.get("design_temp","")),
    ]:
        kv_row(ws1, r, label, value); r += 1

    r += 1
    section_title(ws1, r, "COST SUMMARY", 8); r += 1
    summary_rows = [
        ("Plates & Parts (RM)",            totals["tot_plates"],   WHITE),
        ("Pipes (RM)",                      totals["tot_pipes"],    WHITE),
        ("Flanges (RM)",                    totals["tot_flanges"],  WHITE),
        ("Structural Steel (RM)",           totals["tot_struct"],   WHITE),
        ("▶ Total Raw Material",            totals["tot_rm"],       LIGHT_BLUE),
        ("Fabrication Services",            totals["tot_fab"],      WHITE),
        ("Bought-Out Items",                totals["tot_bo"],       WHITE),
        ("Additional Overheads",            totals["tot_oh"],       WHITE),
        ("Engineering & ASME Design",       totals["engg_design"],  WHITE),
        ("▶ Total Manufacturing Cost",      totals["tot_mfg"],      LIGHT_BLUE),
        ("Contingency",                     totals["cont_amt"],     WHITE),
        ("Profit / Margin",                 totals["profit_amt"],   WHITE),
        ("Packing",                         totals["packing"],      WHITE),
        ("Freight",                         totals["freight"],      WHITE),
        ("▶ Ex-Works Price",               totals["ex_works"],     GREEN_BG),
        ("GST",                             totals["gst_amt"],      WHITE),
        ("▶ FOR Price",                    totals["for_price"],    GREEN_BG),
    ]
    for label, value, bg in summary_rows:
        lc = ws1.cell(row=r, column=1, value=label)
        vc = ws1.cell(row=r, column=2, value=value)
        bold = "▶" in label
        data_style(lc, bg=LIGHT_BLUE if bold else WHITE, bold=bold)
        data_style(vc, bg=bg, bold=bold, align="right")
        vc.number_format = fmt_inr; r += 1

    r += 1
    section_title(ws1, r, "MARGIN ANALYSIS", 8); r += 1
    safe = totals["ex_works"] if totals["ex_works"] else 1
    for label, val in [
        ("RM %",      totals["tot_rm"]/safe*100),
        ("Fab Svc %", totals["tot_fab"]/safe*100),
        ("OH %",      totals["tot_oh"]/safe*100),
        ("Profit %",  totals["profit_amt"]/safe*100),
    ]:
        lc = ws1.cell(row=r, column=1, value=label)
        vc = ws1.cell(row=r, column=2, value=round(val,2))
        data_style(lc, bg=LIGHT_BLUE, bold=True)
        data_style(vc, bg=AMBER_BG if (val<8 or val>60) else WHITE, align="right")
        vc.number_format = "0.00"; r += 1
        vc.value = str(round(val,2)) + "%"

    # SHEET 2 — PLATES & PARTS
    ws2 = wb.create_sheet("Plates & Parts")
    headers = [
        "Part Name","Group","Part Type","Material","Density\n(kg/m³)",
        "Formula Used","Key Dimensions","Qty","Net Wt/Unit\n(kg)",
        "Total Wt\n(kg)","Rate\n(₹/kg)","Amount\n(₹)","Scrap %"
    ]
    col_widths = [20,14,22,10,10,38,36,6,12,12,10,14,8]
    for i,(h_txt,w) in enumerate(zip(headers,col_widths),1):
        c = ws2.cell(row=1, column=i, value=h_txt)
        hdr_style(c, sz=9)
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.row_dimensions[1].height = 32

    FORMULA_DESC = {
        "shell":   "π × ID × Ht × Thk × ρ × (1+scrap)",
        "dish":    "1.09 × π × (ID×1.167/2)² × Thk × ρ × (1+scrap)",
        "annular": "(π/4) × (OD²−ID²) × Thk × ρ × (1+scrap)",
        "solid":   "(π/4) × D² × L × ρ × (1+scrap)",
        "flat":    "W × H × Thk × ρ × (1+scrap)",
        "stiff":   "π × (ID+2×ShThk) × BarW × BarThk × ρ × (1+scrap) × (ShHt/Pitch)",
        "cone":    "π × (R1+R2) × slant × Thk × ρ × (1+scrap)",
        "rect":    "L × W × Thk × ρ × (1+scrap)",
        "tube":    "π × (OD−Thk) × Thk × TubeL × ρ × N × (1+scrap)",
        "limpet":  "CoilMeanD=ShOD+PipeOD | Turns=ShHt/Pitch | Wt=π×(PipeOD−PipeThk)×PipeThk×TotalL×ρ×(1+scrap)",
    }

    def dims_str(fn, dims):
        d = dims or {}
        if fn=="shell":  return f"ID={d.get('id_mm','')} Ht={d.get('ht_mm','')} Thk={d.get('thk_mm','')} mm"
        if fn=="dish":   return f"ShID={d.get('shell_id_mm','')} Thk={d.get('thk_mm','')} mm"
        if fn=="annular":return f"OD={d.get('od_mm','')} ID={d.get('id_mm','')} Thk={d.get('thk_mm','')} mm"
        if fn=="solid":  return f"D={d.get('dia_mm','')} L={d.get('length_mm','')} mm"
        if fn=="flat":   return f"W={d.get('w_mm','')} H={d.get('h_mm','')} Thk={d.get('thk_mm','')} mm"
        if fn=="stiff":  return (f"ShID={d.get('shell_id_mm','')} ShThk={d.get('shell_thk_mm','')} "
                                  f"ShHt={d.get('shell_ht_mm','')} Pitch={d.get('pitch_mm','')} "
                                  f"BarW={d.get('bar_w_mm','')} BarThk={d.get('thk_mm','')} mm")
        if fn=="cone":   return f"LgID={d.get('large_id_mm','')} SmID={d.get('small_id_mm','')} Ht={d.get('ht_mm','')} Thk={d.get('thk_mm','')} mm"
        if fn=="rect":   return f"L={d.get('length_mm','')} W={d.get('width_mm','')} Thk={d.get('thk_mm','')} mm"
        if fn=="tube":   return f"OD={d.get('tube_od_mm','')} Thk={d.get('tube_thk_mm','')} L={d.get('tube_length_mm','')} N={d.get('n_tubes','')} mm"
        if fn=="limpet": return (f"ShID={d.get('shell_id_mm','')} ShThk={d.get('shell_thk_mm','')} ShHt={d.get('shell_ht_mm','')} mm | "
                                 f"PipeOD={d.get('pipe_od_mm','')} PipeThk={d.get('pipe_thk_mm','')} Pitch={d.get('pitch_mm','')} mm | "
                                 f"BottomDish={'Yes' if d.get('cover_bottom_dish') else 'No'}")
        return str(dims)

    row_n = 2
    tot_wt = 0; tot_amt = 0
    for i, p in enumerate(parts):
        fn = PART_TYPES.get(p.get("part_type",""),{}).get("fn","")
        bg = WHITE if i%2==0 else ALT_ROW
        vals = [
            p.get("name",""),
            p.get("group",""),
            p.get("part_type",""),
            p.get("material",""),
            DENSITY.get(p.get("material","SS316L"),8000),
            FORMULA_DESC.get(fn,"Manual entry"),
            dims_str(fn, p.get("dims",{})),
            p.get("qty",1),
            p.get("net_wt_kg",0),
            p.get("total_wt_kg",0),
            p.get("rate",0),
            p.get("amount",0),
            f"{p.get('scrap_pct', DEFAULT_SCRAP.get(fn, 5)):.0f}%",
        ]
        for col, val in enumerate(vals, 1):
            c = ws2.cell(row=row_n, column=col, value=val)
            align = "right" if col in (5,8,9,10,11,12) else "left"
            data_style(c, bg=bg, align=align)
            if col in (9,10):  c.number_format = fmt_kg
            if col in (11,12): c.number_format = fmt_inr
        tot_wt  += p.get("total_wt_kg",0)
        tot_amt += p.get("amount",0)
        row_n += 1

    for col in range(1,14):
        c = ws2.cell(row=row_n, column=col)
        if col==1: c.value = "TOTAL"
        if col==10: c.value = round(tot_wt,3)
        if col==12: c.value = round(tot_amt,2)
        hdr_style(c, bg=MID_BLUE, sz=9)
        if col in (10,): c.number_format = fmt_kg
        if col in (12,): c.number_format = fmt_inr
    ws2.freeze_panes = "A2"

    # SHEET 3 — PIPES & FLANGES
    ws3 = wb.create_sheet("Pipes & Flanges")
    for i, hdr in enumerate(["Description","Item Code","Type","Length(m)","Qty",
                               "Wt/m (kg/m)","Formula","Total Wt (kg)","Rate (₹/kg)","Amount (₹)"],1):
        c = ws3.cell(row=1, column=i, value=hdr); hdr_style(c, sz=9)
    for w, col in zip([22,24,10,10,6,10,32,12,12,14],[1,2,3,4,5,6,7,8,9,10]):
        ws3.column_dimensions[get_column_letter(col)].width = w

    row_n = 2
    for i, p in enumerate(pipes + flanges):
        is_flange = i >= len(pipes)
        bg = WHITE if i%2==0 else ALT_ROW
        formula = "Wt/m × Length × 1.05 (5% fitting allowance) × Qty" if not is_flange else "Wt/m × 1.15 (15% allowance) × Qty"
        vals = [
            p.get("name",""),
            p.get("item_code",""),
            "Flange" if is_flange else "Pipe",
            p.get("length_m","") if not is_flange else "—",
            p.get("qty",1),
            p.get("wt_per_m",0) if not is_flange else p.get("total_wt_kg",0)/max(p.get("qty",1),1),
            formula,
            p.get("total_wt_kg",0),
            p.get("rate",0),
            p.get("amount",0),
        ]
        for col, val in enumerate(vals, 1):
            c = ws3.cell(row=row_n, column=col, value=val)
            data_style(c, bg=bg, align="right" if col in (4,5,6,8,9,10) else "left")
            if col in (8,): c.number_format = fmt_kg
            if col in (9,10): c.number_format = fmt_inr
        row_n += 1
    ws3.freeze_panes = "A2"

    # SHEET 4 — STRUCTURAL STEEL
    ws_struct = wb.create_sheet("Structural Steel")
    for i, hdr in enumerate(["Description","Section","Type","Material","Group",
                              "Unit Wt (kg/m)","Length (m)","Qty",
                              "Total Wt (kg)","Rate (₹/kg)","Amount (₹)"],1):
        c = ws_struct.cell(row=1, column=i, value=hdr); hdr_style(c, sz=9)
    for w, col in zip([22,16,12,10,12,12,10,6,12,12,14],[1,2,3,4,5,6,7,8,9,10,11]):
        ws_struct.column_dimensions[get_column_letter(col)].width = w

    row_n = 2
    tot_s_wt = 0; tot_s_amt = 0
    for i, s in enumerate(struct_items or []):
        bg = WHITE if i%2==0 else ALT_ROW
        vals = [
            s.get("name",""),
            s.get("section",""),
            s.get("type",""),
            s.get("material",""),
            s.get("group",""),
            s.get("unit_wt_kg_per_m",0),
            s.get("length_m",0),
            s.get("qty",1),
            s.get("total_wt_kg",0),
            s.get("rate",0),
            s.get("amount",0),
        ]
        for col, val in enumerate(vals, 1):
            c = ws_struct.cell(row=row_n, column=col, value=val)
            data_style(c, bg=bg, align="right" if col in (6,7,8,9,10,11) else "left")
            if col in (9,): c.number_format = fmt_kg
            if col in (10,11): c.number_format = fmt_inr
        tot_s_wt += s.get("total_wt_kg",0)
        tot_s_amt += s.get("amount",0)
        row_n += 1

    if struct_items:
        for col in range(1,12):
            c = ws_struct.cell(row=row_n, column=col)
            if col==1: c.value = "TOTAL"
            if col==9: c.value = round(tot_s_wt,3)
            if col==11: c.value = round(tot_s_amt,2)
            hdr_style(c, bg=MID_BLUE, sz=9)
            if col == 9: c.number_format = fmt_kg
            if col == 11: c.number_format = fmt_inr
    ws_struct.freeze_panes = "A2"

    # SHEET 5 — FABRICATION SERVICES
    ws4 = wb.create_sheet("Fabrication Services")
    for i, hdr in enumerate(["Service","Basis / Formula","Qty","UOM","Rate","Amount (₹)"],1):
        c = ws4.cell(row=1, column=i, value=hdr); hdr_style(c, sz=9)
    for w, col in zip([30,50,10,8,14,16],[1,2,3,4,5,6]):
        ws4.column_dimensions[get_column_letter(col)].width = w

    row_n = 2
    for i, fs in enumerate(fab_services):
        bg = WHITE if i%2==0 else ALT_ROW
        for col, val in enumerate([
            fs.get("service",""), fs.get("basis",""),
            fs.get("qty",""), fs.get("uom",""),
            fs.get("rate",0), fs.get("amount",0),
        ], 1):
            c = ws4.cell(row=row_n, column=col, value=val)
            data_style(c, bg=bg, align="right" if col in (3,5,6) else "left")
            if col in (5,6): c.number_format = fmt_inr
        row_n += 1

    row_n += 2
    ws4.cell(row=row_n, column=1, value="FABRICATION RATES USED")
    hdr_style(ws4.cell(row=row_n, column=1), bg=MID_BLUE); row_n += 1
    for key, val in fab_rates.items():
        lc = ws4.cell(row=row_n, column=1, value=key.replace("_"," ").title())
        vc = ws4.cell(row=row_n, column=2, value=val)
        data_style(lc, bg=LIGHT_BLUE, bold=True)
        data_style(vc, bg=WHITE, align="right")
        row_n += 1
    ws4.freeze_panes = "A2"

    # SHEET 6 — BO & OH
    ws5 = wb.create_sheet("Bought-Out & OH")
    for i, hdr in enumerate(["Description","Item Code","Group","Qty","Rate","Amount (₹)"],1):
        c = ws5.cell(row=1, column=i, value=hdr); hdr_style(c, sz=9)
    for w, col in zip([28,24,14,8,14,16],[1,2,3,4,5,6]):
        ws5.column_dimensions[get_column_letter(col)].width = w

    row_n = 2
    ws5.cell(row=row_n, column=1, value="── Bought-Out Items ──")
    hdr_style(ws5.cell(row=row_n, column=1), bg=MID_BLUE, sz=9); row_n += 1
    for i, b in enumerate(bo_items):
        bg = WHITE if i%2==0 else ALT_ROW
        for col, val in enumerate([b.get("name",""),b.get("item_code",""),b.get("group",""),b.get("qty",1),b.get("rate",0),b.get("amount",0)],1):
            c = ws5.cell(row=row_n, column=col, value=val)
            data_style(c, bg=bg, align="right" if col in (4,5,6) else "left")
            if col in (5,6): c.number_format = fmt_inr
        row_n += 1

    row_n += 1
    ws5.cell(row=row_n, column=1, value="── Additional Overheads ──")
    hdr_style(ws5.cell(row=row_n, column=1), bg=MID_BLUE, sz=9); row_n += 1
    for i, o in enumerate(oh_items):
        bg = WHITE if i%2==0 else ALT_ROW
        for col, val in enumerate([o.get("description",""),o.get("oh_code",""),o.get("oh_type",""),o.get("qty",1),o.get("rate",0),o.get("amount",0)],1):
            c = ws5.cell(row=row_n, column=col, value=val)
            data_style(c, bg=bg, align="right" if col in (4,5,6) else "left")
            if col in (5,6): c.number_format = fmt_inr
        row_n += 1
    ws5.freeze_panes = "A2"

    # SHEET 7 — FORMULA REFERENCE
    ws6 = wb.create_sheet("Formula Reference")
    ws6.column_dimensions["A"].width = 28
    ws6.column_dimensions["B"].width = 55
    ws6.column_dimensions["C"].width = 20
    ws6.column_dimensions["D"].width = 20

    r = 1
    ws6.merge_cells(start_row=r, start_column=1, end_row=r, end_column=4)
    c = ws6.cell(row=r, column=1, value="B&G ENGINEERING INDUSTRIES — GEOMETRY FORMULA REFERENCE")
    hdr_style(c, sz=12); ws6.row_dimensions[r].height = 24; r += 2

    formulas = [
        ("CYLINDRICAL SHELL",
         "W = π × ID(m) × Ht(m) × Thk(m) × ρ(kg/m³) × (1 + scrap)",
         "scrap = 5%", "density per MOC"),
        ("DISH END (Torispherical)",
         "Blank dia = ID × 1.167 | Area = 1.09 × π × (blank_dia/2)² | W = Area × Thk(m) × ρ × (1 + scrap)",
         "scrap = 15%", "Crown radius = ID"),
        ("ANNULAR PLATE / FLANGE",
         "W = (π/4) × (OD² − ID²) × Thk(m) × ρ × (1 + scrap)",
         "scrap = 5%", "All dims in metres"),
        ("SOLID ROUND (Shaft / Bush)",
         "W = (π/4) × D(m)² × L(m) × ρ × (1 + scrap)",
         "scrap = 15%", ""),
        ("FLAT RECTANGLE",
         "W = W(m) × H(m) × Thk(m) × ρ × (1 + scrap)",
         "scrap = 5%", ""),
        ("STIFFENER RINGS",
         "Shell OD = ID + 2×Thk | Circ = π×OD | N = ShHt/Pitch | Wt/ring = Circ×BarW×BarThk×ρ×(1+scrap) | Total = Wt/ring×N",
         "scrap = 5%", "N rings = derived qty"),
        ("CONE / REDUCER",
         "R1=LargeID/2, R2=SmallID/2 | Slant=√(Ht²+(R1−R2)²) | W=π×(R1+R2)×Slant×Thk×ρ×(1+scrap)",
         "scrap = 5%", ""),
        ("RECTANGULAR PLATE",
         "W = L(m) × W(m) × Thk(m) × ρ × (1 + scrap)",
         "scrap = 5%", ""),
        ("TUBE BUNDLE",
         "Mid_r=(OD−Thk)/2 | Wt/tube=π×2×Mid_r×L×Thk×ρ | Total=Wt/tube×N_tubes×(1+scrap)",
         "scrap = 5%", ""),
        ("STRUCTURAL (Angle / Channel / Beam)",
         "W = Unit Wt (kg/m, per IS 808/2062) × Length(m) × Qty × 1.05",
         "5% cutting allowance", "IS-standard unit weights"),
    ]

    for i, hdr in enumerate(["Part Type","Formula","Notes","Reference"],1):
        c = ws6.cell(row=r, column=i, value=hdr); hdr_style(c, sz=9)
    r += 1

    for idx, (pt, formula, notes, ref) in enumerate(formulas):
        bg = WHITE if idx%2==0 else ALT_ROW
        for col, val in enumerate([pt, formula, notes, ref], 1):
            c = ws6.cell(row=r, column=col, value=val)
            data_style(c, bg=LIGHT_BLUE if col==1 else bg, bold=(col==1))
            c.alignment = Alignment(wrap_text=True, vertical="top")
        ws6.row_dimensions[r].height = max(15, formula.count("\n")*14 + 14)
        r += 1

    r += 2
    ws6.merge_cells(start_row=r, start_column=1, end_row=r, end_column=4)
    c = ws6.cell(row=r, column=1, value="MATERIAL DENSITY TABLE (kg/m³)")
    hdr_style(c, bg=MID_BLUE, sz=10); r += 1
    for mat, den in DENSITY.items():
        lc = ws6.cell(row=r, column=1, value=mat)
        vc = ws6.cell(row=r, column=2, value=den)
        data_style(lc, bg=LIGHT_BLUE, bold=True)
        data_style(vc, bg=WHITE, align="right")
        r += 1

    buf = _io.BytesIO()
    wb.save(buf); buf.seek(0)
    return buf
# ─────────────────────────────────────────────────────────────────────────────
# MASTER LOADERS
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(ttl=300)
def load_rm_master():
    rows = sb_fetch("est_rm_master", order="category")
    for r in rows:
        r["rate"] = float(r["rate"]) if r.get("rate") else 0.0
        r["unit_wt_kg_per_m"] = float(r["unit_wt_kg_per_m"]) if r.get("unit_wt_kg_per_m") else None
    return {r["ref_code"]: r for r in rows}

@st.cache_data(ttl=300)
def load_oh_master():
    rows = sb_fetch("est_oh_master", order="oh_type")
    for r in rows:
        r["rate"] = float(r["rate"]) if r.get("rate") else 0.0
    return {r["oh_code"]: r for r in rows}

@st.cache_data(ttl=60)
def load_clients_full():
    return sb_fetch("master_clients", order="name")

@st.cache_data(ttl=60)
def load_anchor_qtns():
    return sb_fetch("anchor_projects", select="quote_ref,project_description,client_name", order="created_at")

@st.cache_data(ttl=30)
def load_all_estimations():
    return sb_fetch("estimations", order="updated_at")

# ─────────────────────────────────────────────────────────────────────────────
# SESSION STATE HELPERS
# ─────────────────────────────────────────────────────────────────────────────

# Fields that get packed into the spec_json column on save instead of
# becoming individual DB columns. To add more spec fields in future,
# just add to _blank_hdr() AND append the field name here — no DB change.
SPEC_FIELDS = [
    # Option B cost-structure additions
    "discount_pct",
    # Detailed spec sheet fields (Part A of Spec Sheet feature)
    "service_fluid", "working_vol_ltrs",
    "design_pressure_shell", "design_pressure_jacket",
    "corrosion_allowance_mm", "joint_efficiency", "hydrotest_pressure",
    "heating_medium", "cooling_medium",
    "agitator_rpm", "motor_hp", "motor_make",
    "gearbox_make", "gearbox_ratio",
    "seal_type", "seal_make",
]

def _blank_hdr():
    return dict(
        qtn_number="", revision="R0", customer_name="",
        equipment_type=EQUIPMENT_NAMES[0], equipment_desc="", tag_number="",
        capacity_ltrs=2000.0, shell_dia_mm=1300.0, shell_ht_mm=1500.0,
        shell_thk_mm=8.0, dish_thk_mm=10.0,
        jacket_type="SS304 Half-pipe Jacket with Insulation Jacket",
        agitator_type="Anchor", design_code="ASME Sec VIII Div 1",
        design_pressure="FV to 4.5 Bar", design_temp="-50 to 250°C",
        moc_shell="SS316L", moc_jacket="SS304", surface_finish="Internal: Ra ≤ 0.8 μm  |  External: Buffed",
        status="Draft", prepared_by="", checked_by="",
        profit_margin_pct=10.0, contingency_pct=0.0,
        packing_amt=5000.0, freight_amt=10000.0,
        gst_pct=18.0, engg_design_amt=25000.0, notes="",
        discount_pct=0.0,  # Discount applied on Ex-Works price (Option B cost structure)
        # ── Detailed Specification fields (for customer Spec Sheet in Section 2) ──
        # These print in the customer-facing quotation, organised into 6 sub-blocks.
        # Leave empty to skip a row in the printed spec sheet.

        # General
        service_fluid="",                                  # e.g. "API intermediate / Solvent"

        # Vessel dimensions
        working_vol_ltrs=0.0,                              # 0 = auto-default to 80% of capacity

        # Design conditions
        design_pressure_shell="FV to 4.5 Bar",             # Internal/External pressure for shell
        design_pressure_jacket="6.0 Bar / FV",             # Jacket / limpet design pressure
        corrosion_allowance_mm=1.5,
        joint_efficiency=0.85,
        hydrotest_pressure="1.5× Design Pressure",

        # Heating / Cooling
        heating_medium="Steam / Hot Water / Thermic Fluid",
        cooling_medium="Chilled Water / Brine",

        # Agitator & Drive
        agitator_rpm="",                                   # e.g. "60 RPM"
        motor_hp="",                                       # e.g. "7.5 HP / 5.5 kW"
        motor_make="ABB / Crompton / Bharat Bijlee",
        gearbox_make="Bonfiglioli / Premium / Elecon",
        gearbox_ratio="",                                  # e.g. "1:25"
        seal_type="Single mechanical seal",
        seal_make="EagleBurgmann / John Crane / Flowserve",
        scope_items=(
            "Pressure vessel / equipment fabricated as per approved GA drawing\n"
            "All nozzles, manholes, handholes and process connections per nozzle schedule\n"
            "Jacket / limpet coil / half-pipe with insulation jacket (where applicable)\n"
            "Agitator complete with gearbox, motor and mechanical seal (where applicable)\n"
            "Support structure — lugs, legs or saddles as applicable\n"
            "Internal grinding and buffing to specified Ra surface finish\n"
            "Equipment nameplate with tag number and serial number"
        ),
        scope_exclusions=(
            "Civil / structural works\n"
            "Electrical & Instrumentation\n"
            "Erection & commissioning at site\n"
            "DQ / IQ / OQ / PQ validation\n"
            "Freight, insurance and unloading at site\n"
            "Import duties if applicable"
        ),

        quality_intro=(
            "B&G Engineering Industries operates as an engineering-led manufacturer. "
            "Every project is built to ASME Section VIII Division 1 requirements "
            "with full documentation and traceability."
        ),
        quality_points=(
            "Raw material procurement with original Mill Test Certificates (MTC) for all pressure parts\n"
            "100% Positive Material Identification (PMI) verification before cutting\n"
            "Heat number and cast number traceability maintained throughout fabrication\n"
            "Qualified welders — WPS / PQR compliant, TIG welding for all SS316L pressure joints\n"
            "Precision plasma / laser cutting and CNC-controlled plate rolling\n"
            "Pharma-grade internal grinding to Ra ≤ 0.8 μm and external mechanical buffing\n"
            "Dimensional inspection against approved GA drawings at each stage\n"
            "Hydrostatic / pneumatic / vacuum leak test as per ASME Sec VIII Div 1\n"
            "Dye Penetrant (DP) testing on all critical welds\n"
            "Complete QA dossier prepared and delivered with equipment\n"
            "Factory Acceptance Test (FAT) support at works on request"
        ),

        doc_deliverables=(
            "General Arrangement (GA) Drawing — IFC revision\n"
            "Nozzle orientation and schedule drawing\n"
            "Bill of Materials (BOM)\n"
            "Mill Test Certificates (MTC) for all pressure parts\n"
            "PMI verification records\n"
            "Weld map and weld log\n"
            "DP / RT inspection reports\n"
            "Dimensional inspection report\n"
            "Hydrostatic / leak test certificate\n"
            "Surface finish inspection record (Ra measurement)\n"
            "Inspection and Release Note (IRN)\n"
            "Equipment nameplate photograph"
        ),

        price_basis="Ex-Works, Pashamylaram, Hyderabad — 502307. Packing in MS crate included. Freight, insurance and unloading at site excluded.",
        gst_clause="GST @ 18% (HSN 8419) as applicable at time of invoicing. Any new statutory levy introduced after offer date will be charged additionally.",
        payment_terms="40% advance along with Purchase Order  |  50% against Pro-forma invoice on readiness for dispatch  |  10% on delivery",
        delivery_weeks="12–16",
        delivery_note="Subject to availability of raw material at time of order.",
        offer_validity="This offer is valid for 7 calendar days from the date above. Prices subject to change if raw material rates move by more than 3%.",
        warranty_clause="12 months from date of commissioning or 18 months from date of dispatch, whichever is earlier. Covers manufacturing defects under normal operating conditions as per design basis.",
        inspection_clause="Customer may depute inspector for stage and final inspection at our works. Third-party inspection (TPI) agency charges, if any, are in customer scope.",
        special_notes="",
    )

def _reset_form():
    for k in ["est_hdr", "est_parts", "est_pipes", "est_flanges", "est_struct",
              "est_fab", "est_bo", "est_oh", "est_edit_id", "edit_part_idx", "fab_rates"]:
        st.session_state.pop(k, None)

def _load_est_into_form(est):
    h = _blank_hdr()

    # Load regular header fields from individual DB columns
    for k in h:
        if k in est and est[k] is not None:
            h[k] = est[k]

    # Unpack spec_json blob back into header fields. Any SPEC_FIELDS missing
    # from spec_json (e.g. fields added after the row was saved) keep their
    # default value from _blank_hdr() — so old rows don't break.
    try:
        spec_dict = json.loads(est.get("spec_json") or "{}")
    except (json.JSONDecodeError, TypeError):
        spec_dict = {}
    for k in SPEC_FIELDS:
        if k in spec_dict and spec_dict[k] is not None:
            h[k] = spec_dict[k]

    st.session_state.est_hdr     = h
    st.session_state.est_parts   = json.loads(est.get("parts_json")     or "[]")
    st.session_state.est_pipes   = json.loads(est.get("pipes_json")     or "[]")
    st.session_state.est_flanges = json.loads(est.get("flanges_json")   or "[]")
    st.session_state.est_struct  = json.loads(est.get("struct_json")    or "[]")
    st.session_state.est_fab     = json.loads(est.get("fab_json")       or "[]")
    st.session_state.est_bo      = json.loads(est.get("bo_json")        or "[]")
    st.session_state.est_oh      = json.loads(est.get("oh_json")        or "[]")
    st.session_state.fab_rates   = json.loads(est.get("fab_rates_json") or json.dumps(FAB_DEFAULTS))
    st.session_state.est_edit_id = est.get("id")

def _do_save(reset_after=False):
    """
    Core save function — always reads from st.session_state directly.
    If reset_after=True, clears the form after saving.
    """
    h = st.session_state.est_hdr
    edit_id = st.session_state.est_edit_id
    qtn = h.get("qtn_number", "").strip()

    if not qtn:
        st.warning("⚠️ Quotation Number is empty. Enter it in Tab 1️⃣ Header first.")
        return False

    existing = sb_fetch("estimations", select="id", filters={"qtn_number": qtn})
    if existing and not edit_id:
        st.error(f"❌ QTN **{qtn}** already exists. Load it from the search panel to edit it.")
        return False
    if existing and edit_id:
        if any(str(e.get("id")) != str(edit_id) for e in existing):
            st.error(f"❌ QTN **{qtn}** belongs to a different estimation.")
            return False

    # Fields that don't go to the DB at all (UI-only or computed elsewhere)
    skip = {"customer_id"}

    # Extract SPEC_FIELDS into a separate JSON blob — they live in spec_json
    # column, not as individual DB columns. This lets us add new spec fields
    # later without ALTER TABLE.
    spec_dict = {k: h.get(k) for k in SPEC_FIELDS if h.get(k) is not None}

    # Build the clean header for DB write: skip the UI-only fields AND the
    # spec fields (those are already captured in spec_dict above).
    clean_h = {
        k: v for k, v in h.items()
        if k not in skip and k not in SPEC_FIELDS and v is not None
    }

    row = {
        **clean_h,
        "parts_json":     json.dumps(st.session_state.est_parts),
        "pipes_json":     json.dumps(st.session_state.est_pipes),
        "flanges_json":   json.dumps(st.session_state.est_flanges),
        "struct_json":    json.dumps(st.session_state.est_struct),
        "fab_json":       json.dumps(st.session_state.est_fab),
        "bo_json":        json.dumps(st.session_state.est_bo),
        "oh_json":        json.dumps(st.session_state.est_oh),
        "fab_rates_json": json.dumps(st.session_state.fab_rates),
        "spec_json":      json.dumps(spec_dict),
        "updated_at":     datetime.now().isoformat(),
    }

    if edit_id:
        ok = sb_update("estimations", row, "id", edit_id)
        msg = f"Updated {qtn}"
    else:
        row["created_at"] = datetime.now().isoformat()
        ok = sb_insert("estimations", row)
        msg = f"Saved {qtn}"
        if ok:
            saved = sb_fetch("estimations", select="id", filters={"qtn_number": qtn})
            if saved:
                st.session_state.est_edit_id = saved[0]["id"]

    if ok:
        n_p = len(st.session_state.est_parts)
        rm  = sum(p.get("amount", 0) for p in st.session_state.est_parts)
        fab = sum(f.get("amount", 0) for f in st.session_state.est_fab)
        st.success(f"✅ {msg}  |  {n_p} parts  |  RM ₹{rm:,.0f}  |  Fab ₹{fab:,.0f}")
        st.cache_data.clear()
        if reset_after:
            _reset_form()
        return True
    return False

def _save_draft_bar(tab_key):
    h = st.session_state.est_hdr
    st.divider()
    sb1, sb2, sb3 = st.columns([4, 1, 1])
    qtn  = h.get("qtn_number", "") or "—"
    n_p  = len(st.session_state.est_parts)
    n_pi = len(st.session_state.est_pipes)
    n_st = len(st.session_state.est_struct)
    n_f  = len(st.session_state.est_fab)
    n_b  = len(st.session_state.est_bo)
    sb1.caption(f"💾 **{qtn}**  |  {n_p} parts  |  {n_pi} pipes  |  {n_st} structural  |  {n_f} fab lines  |  {n_b} BO items")
    if sb2.button("💾 Save Draft", use_container_width=True, type="primary", key=f"sd_{tab_key}"):
        _do_save(reset_after=False)
    if sb3.button("🗑️ Reset / New", use_container_width=True, key=f"rst_{tab_key}"):
        _reset_form()
        st.rerun()

for key, default in [
    ("est_hdr",       _blank_hdr()),
    ("est_parts",     []),
    ("est_pipes",     []),
    ("est_flanges",   []),
    ("est_struct",    []),
    ("est_fab",       []),
    ("est_bo",        []),
    ("est_oh",        []),
    ("est_edit_id",   None),
    ("edit_part_idx", None),
    ("fab_rates",     dict(FAB_DEFAULTS)),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ─────────────────────────────────────────────────────────────────────────────
# PAGE HEADER
# ─────────────────────────────────────────────────────────────────────────────
st.title("📐 Estimation & Costing")
if _LOGO_BYTES:
    st.caption(f"🖼️ Logo loaded from Supabase: `{_LOGO_FNAME}` — will appear in all quotations.")
else:
    st.caption("🖼️ No logo found in `progress-photos` bucket. Upload as `logo.png` to include in quotations.")
_qtn_now = st.session_state.est_hdr.get("qtn_number", "")
_eid_now  = st.session_state.est_edit_id
if _eid_now and _qtn_now:
    st.info(f"✏️ Editing: **{_qtn_now}**  |  {len(st.session_state.est_parts)} parts  |  RM ₹{sum(p.get('amount',0) for p in st.session_state.est_parts):,.0f}")
elif _qtn_now:
    st.info(f"📝 New estimation in progress: **{_qtn_now}**")
st.markdown("---")

TAB_LIST, TAB_NEW, TAB_QUOTE, TAB_SIMILAR, TAB_MASTERS = st.tabs([
    "📋 Register", "➕ New / Edit", "✍️ Quote Editor", "🔍 Similar Equipment", "📊 Masters",
])

# ══════════════════════════════════════════════════════════════════════════════
# TAB: REGISTER
# ══════════════════════════════════════════════════════════════════════════════
with TAB_LIST:
    st.subheader("Estimations Register")
    col_f1, col_f2, col_f3, col_f4 = st.columns(4)
    f_status   = col_f1.selectbox("Status",    ["All", "Draft", "Issued", "Won", "Lost", "On Hold"])
    f_equip    = col_f2.selectbox("Equipment", ["All"] + EQUIPMENT_NAMES)
    f_customer = col_f3.text_input("Customer")
    f_search   = col_f4.text_input("QTN / Tag")

    all_est = load_all_estimations()
    if f_status != "All":   all_est = [e for e in all_est if e.get("status") == f_status]
    if f_equip != "All":    all_est = [e for e in all_est if e.get("equipment_type") == f_equip]
    if f_customer:          all_est = [e for e in all_est if f_customer.lower() in (e.get("customer_name", "") or "").lower()]
    if f_search:            all_est = [e for e in all_est if f_search.lower() in (e.get("qtn_number", "") or "").lower() or f_search.lower() in (e.get("tag_number", "") or "").lower()]

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Total",  len(all_est))
    m2.metric("Draft",  sum(1 for e in all_est if e.get("status") == "Draft"))
    m3.metric("Issued", sum(1 for e in all_est if e.get("status") == "Issued"))
    m4.metric("Won",    sum(1 for e in all_est if e.get("status") == "Won"))
    st.divider()

    if not all_est:
        st.info("No estimations yet. Use ➕ New / Edit tab to create one.")
    else:
        summary_rows = []
        for est in reversed(all_est):
            eq_info = EQUIPMENT_TYPES.get(est.get("equipment_type", ""), {})
            si = {"Draft": "🟡", "Issued": "🔵", "Won": "🟢", "Lost": "🔴", "On Hold": "⚪"}.get(est.get("status", ""), "🟡")
            summary_rows.append({
                "":           f"{si} {eq_info.get('icon', '🔧')}",
                "QTN No":     est.get("qtn_number", ""),
                "Customer":   est.get("customer_name", ""),
                "Equipment":  est.get("equipment_desc", ""),
                "Cap (L)":    est.get("capacity_ltrs", ""),
                "Status":     est.get("status", ""),
                "Prepared By": est.get("prepared_by", ""),
                "Updated":    str(est.get("updated_at", ""))[:10],
            })
        st.dataframe(pd.DataFrame(summary_rows), use_container_width=True, hide_index=True)

        st.markdown("#### Select a quotation to view details and actions")
        qtn_opts = ["— select —"] + [e.get("qtn_number", "") for e in reversed(all_est) if e.get("qtn_number")]
        selected_qtn = st.selectbox("QTN", qtn_opts, label_visibility="collapsed")

        if selected_qtn != "— select —":
            est = next((e for e in all_est if e.get("qtn_number") == selected_qtn), None)
            if est:
                parts   = json.loads(est.get("parts_json")   or "[]")
                pipes   = json.loads(est.get("pipes_json")   or "[]")
                flanges = json.loads(est.get("flanges_json") or "[]")
                struct_l = json.loads(est.get("struct_json")  or "[]")
                fab_s   = json.loads(est.get("fab_json")     or "[]")
                bo      = json.loads(est.get("bo_json")      or "[]")
                oh      = json.loads(est.get("oh_json")      or "[]")
                T = calc_totals(
                    parts, pipes, flanges, struct_l, fab_s, bo, oh,
                    float(est.get("profit_margin_pct") or 10), float(est.get("contingency_pct") or 0),
                    float(est.get("packing_amt") or 0), float(est.get("freight_amt") or 0),
                    float(est.get("gst_pct") or 18), float(est.get("engg_design_amt") or 0),
                )
                st.markdown("---")
                d1, d2, d3 = st.columns(3)
                d1.write(f"**Type:** {est.get('equipment_type', '')}"); d1.write(f"**Tag:** {est.get('tag_number', '-')}")
                d2.write(f"**Revision:** {est.get('revision', 'R0')}"); d2.write(f"**Prepared By:** {est.get('prepared_by', '-')}")
                d3.write(f"**Status:** {est.get('status', '')}"); d3.write(f"**Updated:** {str(est.get('updated_at', ''))[:10]}")

                k1, k2, k3, k4, k5 = st.columns(5)
                k1.metric("Raw Material", f"₹{T['tot_rm']:,.0f}")
                k2.metric("Fabrication",  f"₹{T['tot_fab']:,.0f}")
                k3.metric("Mfg Cost",     f"₹{T['tot_mfg']:,.0f}")
                k4.metric("Ex-Works",     f"₹{T['ex_works']:,.0f}")
                k5.metric("FOR Price",    f"₹{T['for_price']:,.0f}")

                st.markdown("**Actions**")
                a1, a2, a3, a4 = st.columns(4)
                if a1.button("✏️ Edit", use_container_width=True, type="primary"):
                    fresh = sb_fetch("estimations", filters={"id": est["id"]})
                    _load_est_into_form(fresh[0] if fresh else est)
                    st.success(f"Loaded — {len(st.session_state.est_parts)} parts. Click ➕ New / Edit tab.")
                if a2.button("📋 Clone to New", use_container_width=True):
                    _load_est_into_form(est)
                    st.session_state.est_hdr["qtn_number"] = ""
                    st.session_state.est_hdr["revision"]   = "R0"
                    st.session_state.est_hdr["status"]     = "Draft"
                    st.session_state.est_hdr["notes"]      = f"Cloned from {est.get('qtn_number', '')}"
                    st.session_state.est_edit_id           = None
                    st.success("Cloned — go to ➕ New / Edit and enter new QTN.")
                cust_rows = sb_fetch("master_clients", filters={"name": est.get("customer_name", "")})
                cust_data = cust_rows[0] if cust_rows else {}
                a3.download_button(
                    "📄 Standard Quote",
                    generate_docx(est, cust_data, T, fab_s, show_breakup=False),
                    file_name=f"{est.get('qtn_number','QTN')}_{est.get('revision','R0')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"dl_{est.get('id')}",
                )
                a4_dl_col = st.columns(4)[3]
                a4_dl_col.download_button(
                    "📋 With Breakup",
                    generate_docx(est, cust_data, T, fab_s, show_breakup=True),
                    file_name=f"{est.get('qtn_number','QTN')}_{est.get('revision','R0')}_breakup.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"dl_bk_{est.get('id')}",
                )
                new_status = a4.selectbox(
                    "Status",
                    ["Draft", "Issued", "Won", "Lost", "On Hold"],
                    index=["Draft", "Issued", "Won", "Lost", "On Hold"].index(est.get("status", "Draft")),
                    key=f"st_{est.get('id')}",
                )
                if new_status != est.get("status"):
                    if a4.button("✅ Apply", key=f"ap_{est.get('id')}", use_container_width=True):
                        sb_update("estimations", {"status": new_status, "updated_at": datetime.now().isoformat()}, "id", est.get("id"))
                        st.cache_data.clear(); st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# TAB: NEW / EDIT
# ══════════════════════════════════════════════════════════════════════════════
with TAB_NEW:
    edit_id  = st.session_state.est_edit_id
    _qtn_hdr = st.session_state.est_hdr.get("qtn_number", "")
    if edit_id and _qtn_hdr:
        st.subheader(f"✏️ Editing: {_qtn_hdr}")
    elif _qtn_hdr:
        st.subheader(f"📝 New Estimation: {_qtn_hdr}")
    else:
        st.subheader("➕ New Estimation")

    rm_master   = load_rm_master()
    oh_master   = load_oh_master()
    clients     = load_clients_full()
    anchor_qtns = load_anchor_qtns()

    client_names = [c["name"] for c in clients]
    oh_codes     = list(oh_master.keys())
    plate_rm     = [k for k, v in rm_master.items() if v.get("category") == "RM"]
    pipe_rm      = [k for k, v in rm_master.items() if v.get("rm_type") == "Pipe"]
    flg_rm       = [k for k, v in rm_master.items() if v.get("rm_type") == "FLG"]
    bo_rm        = [k for k, v in rm_master.items() if v.get("category") == "BO"]
    all_groups   = sorted({"SHELL", "DISH_ENDS", "JACKET", "INS_JACKET", "AGITATOR", "BAFFLES",
                            "LUGS", "STIFFNERS", "MANHOLE", "NOZZLES", "RM_MISC", "BODY_FL",
                            "TUBE_BUNDLE", "TUBE_SHEET", "FILTER_PLATE", "TRAYS", "FRAME", "OTHER"})

    f1, f2, f3, f_struct, f4, f5, f6 = st.tabs([
        "1️⃣ Header", "2️⃣ Plates & Parts", "3️⃣ Pipes & Flanges",
        "4️⃣ Structural", "5️⃣ Fabrication Services", "6️⃣ Bought-Out & OH", "7️⃣ Summary & Save",
    ])
    h = st.session_state.est_hdr

    # ── Initialise edit-index trackers for all line-item lists ────────────────
    for _k in ["edit_pipe_idx", "edit_flg_idx", "edit_fab_idx", "edit_bo_idx", "edit_oh_idx"]:
        if _k not in st.session_state:
            st.session_state[_k] = None

    # ── F1: HEADER ─────────────────────────────────────────────────────────────
    with f1:
        all_est_list = load_all_estimations()
        all_est_valid = [e for e in all_est_list if e.get("qtn_number")]

        with st.container(border=True):
            if edit_id:
                st.success(f"✏️ **Editing:** {h.get('qtn_number', '')}  |  {h.get('customer_name', '')}  |  {len(st.session_state.est_parts)} parts loaded")
            elif h.get("qtn_number"):
                st.info(f"📝 **Working on new:** {h.get('qtn_number', '')}  |  {h.get('customer_name', '')}  |  {len(st.session_state.est_parts)} parts")
            else:
                st.warning("📄 New estimation — fill details below, or search and load a saved one.")

            st.markdown("**🔍 Search & Load a Saved Estimation**")
            sf1, sf2, sf3, sf4 = st.columns(4)
            srch_qtn  = sf1.text_input("QTN",      placeholder="e.g. B&G/MAITHRI", key="srch_qtn")
            srch_cust = sf2.text_input("Customer",  placeholder="e.g. Neuland",    key="srch_cust")
            srch_eq   = sf3.selectbox("Equipment",  ["All"] + EQUIPMENT_NAMES,     key="srch_eq")
            srch_stat = sf4.selectbox("Status",     ["All", "Draft", "Issued", "Won", "Lost", "On Hold"], key="srch_stat")

            filtered = all_est_valid
            if srch_qtn:           filtered = [e for e in filtered if srch_qtn.lower()  in (e.get("qtn_number", "") or "").lower()]
            if srch_cust:          filtered = [e for e in filtered if srch_cust.lower() in (e.get("customer_name", "") or "").lower()]
            if srch_eq != "All":   filtered = [e for e in filtered if e.get("equipment_type") == srch_eq]
            if srch_stat != "All": filtered = [e for e in filtered if e.get("status") == srch_stat]

            if filtered:
                def _lbl(e):
                    si = {"Draft": "🟡", "Issued": "🔵", "Won": "🟢", "Lost": "🔴", "On Hold": "⚪"}.get(e.get("status", ""), "🟡")
                    return f"{si} {e.get('qtn_number', '')}  |  {e.get('customer_name', '—')}  |  {e.get('equipment_desc', '—')}  |  {e.get('capacity_ltrs', '')}L  |  {e.get('status', '')}"

                opts = ["— select to load —"] + [_lbl(e) for e in filtered]
                ld1, ld2, ld3 = st.columns([5, 1, 1])
                load_sel = ld1.selectbox("Select", opts, key="f1_load_sel", label_visibility="collapsed")

                if ld2.button("📂 Load / Edit", use_container_width=True, type="primary", key="f1_load_btn"):
                    if load_sel != "— select to load —":
                        preview = filtered[opts.index(load_sel) - 1]
                        fresh = sb_fetch("estimations", filters={"id": preview["id"]})
                        match = fresh[0] if fresh else preview
                        _load_est_into_form(match)
                        st.success(f"✅ Loaded **{match.get('qtn_number','')}** — {len(st.session_state.est_parts)} parts, {len(st.session_state.est_pipes)} pipes, {len(st.session_state.est_fab)} fab lines.")
                        st.rerun()

                if ld3.button("📋 Clone", use_container_width=True, key="f1_clone_btn"):
                    if load_sel != "— select to load —":
                        preview = filtered[opts.index(load_sel) - 1]
                        fresh = sb_fetch("estimations", filters={"id": preview["id"]})
                        match = fresh[0] if fresh else preview
                        _load_est_into_form(match)
                        st.session_state.est_hdr["qtn_number"] = ""
                        st.session_state.est_hdr["revision"]   = "R0"
                        st.session_state.est_hdr["status"]     = "Draft"
                        st.session_state.est_hdr["notes"]      = f"Cloned from {match.get('qtn_number', '')}"
                        st.session_state.est_edit_id           = None
                        st.success(f"📋 Cloned from **{match.get('qtn_number', '')}** — enter new QTN number below.")
                        st.rerun()
            else:
                st.caption("No saved estimations match the filters. Fill the form below to create a new one.")

            if st.button("🗑️ Clear / Start Fresh", key="f1_clear_btn"):
                _reset_form(); st.rerun()

        st.divider()

        st.markdown("##### Equipment Type")
        prev_type = h.get("equipment_type", EQUIPMENT_NAMES[0])
        if prev_type not in EQUIPMENT_NAMES:
            prev_type = EQUIPMENT_NAMES[0]
        eq_c1, eq_c2 = st.columns([4, 1])
        h["equipment_type"] = eq_c1.selectbox(
            "Equipment Type", EQUIPMENT_NAMES,
            index=EQUIPMENT_NAMES.index(prev_type),
            format_func=lambda x: f"{EQUIPMENT_TYPES[x]['icon']}  {x}",
            label_visibility="collapsed",
        )
        eq_info = EQUIPMENT_TYPES[h["equipment_type"]]
        eq_c2.metric("Category", eq_info["category"])
        st.caption(f"_{eq_info['description']}_  •  Margin: **{eq_info['margin_hint'][0]}–{eq_info['margin_hint'][1]}%**  •  Labour: **{eq_info['labour_norm']}**")
        st.divider()

        st.markdown("##### Pull from Anchor Portal  _(optional)_")
        anc_options = ["— type QTN manually —"] + [
            f"{a.get('quote_ref', '')}  |  {a.get('client_name', '')}  |  {a.get('project_description', '')}"
            for a in anchor_qtns
        ]
        anc_sel = st.selectbox("Anchor Portal", anc_options, label_visibility="collapsed")
        if anc_sel != "— type QTN manually —":
            chosen = anchor_qtns[anc_options.index(anc_sel) - 1]
            h["qtn_number"]    = chosen.get("quote_ref", "")
            h["customer_name"] = chosen.get("client_name", "")
            st.success(f"Auto-filled — QTN: **{h['qtn_number']}**  |  Customer: **{h['customer_name']}**")
        st.divider()

        st.markdown("##### Offer Details")
        c1, c2, c3 = st.columns(3)
        h["qtn_number"] = c1.text_input("Quotation Number *", value=h["qtn_number"], placeholder="e.g. B&G/MAITHRI/2026/2922")
        h["revision"]   = c2.selectbox("Revision", ["R0", "R1", "R2", "R3", "R4", "R5"],
                                        index=["R0", "R1", "R2", "R3", "R4", "R5"].index(h.get("revision", "R0")))
        h["status"]     = c3.selectbox("Status", ["Draft", "Issued", "Won", "Lost", "On Hold"],
                                        index=["Draft", "Issued", "Won", "Lost", "On Hold"].index(h.get("status", "Draft")))
        st.divider()

        st.markdown("##### Customer")
        cust_opts = ["— select —"] + client_names
        cust_idx  = cust_opts.index(h["customer_name"]) if h["customer_name"] in cust_opts else 0
        sel_cust  = st.selectbox("Customer", cust_opts, index=cust_idx, label_visibility="collapsed")
        if sel_cust != "— select —":
            h["customer_name"] = sel_cust
            cd = next((c for c in clients if c["name"] == sel_cust), {})
            cc = st.columns(4)
            cc[0].caption(f"📍 {cd.get('address', '—')}")
            cc[1].caption(f"🏷️ GSTIN: {cd.get('gstin', '—')}")
            cc[2].caption(f"👤 {cd.get('contact_person', '—')}")
            cc[3].caption(f"📞 {cd.get('phone', '—')}")
        st.divider()

        st.markdown("##### Equipment Parameters")
        c1, c2 = st.columns(2)
        h["equipment_desc"] = c1.text_input("Description", value=h["equipment_desc"], placeholder="e.g. 2000 Ltrs SS316L Jacketed Reactor")
        h["tag_number"]     = c2.text_input("Tag Number",  value=h["tag_number"],     placeholder="e.g. R-505")

        c1, c2, c3, c4, c5 = st.columns(5)
        h["capacity_ltrs"] = c1.number_input("Capacity (Ltrs)", value=float(h["capacity_ltrs"]), min_value=0.0, step=100.0)
        h["shell_dia_mm"]  = c2.number_input("Shell ID (mm)",   value=float(h["shell_dia_mm"]),  min_value=0.0, step=50.0)
        h["shell_ht_mm"]   = c3.number_input("Shell Ht (mm)",   value=float(h["shell_ht_mm"]),   min_value=0.0, step=100.0)
        h["shell_thk_mm"]  = c4.number_input("Shell Thk (mm)",  value=float(h["shell_thk_mm"]),  min_value=0.0, step=1.0)
        h["dish_thk_mm"]   = c5.number_input("Dish Thk (mm)",   value=float(h["dish_thk_mm"]),   min_value=0.0, step=1.0)

        s_area   = calc_shell_area(h["shell_dia_mm"], h["shell_ht_mm"])
        d_area   = calc_dish_area(h["shell_dia_mm"])
        vol      = calc_shell_volume_ltrs(h["shell_dia_mm"], h["shell_ht_mm"])
        int_area = s_area + d_area * 2
        weld_m   = calc_weld_metres(h["shell_dia_mm"], h["shell_ht_mm"],
                                    has_jacket=bool(h.get("jacket_type", "")),
                                    has_agitator=bool(h.get("agitator_type", "")))
        g1, g2, g3, g4, g5 = st.columns(5)
        g1.metric("Shell Area (m²)",      f"{s_area:.3f}")
        g2.metric("Dish Area (m²)",       f"{d_area:.3f}")
        g3.metric("Total Internal (m²)",  f"{int_area:.3f}")
        g4.metric("Shell Vol (Ltrs)",     f"{vol:.0f}")
        g5.metric("Est. Weld (m)",        f"{weld_m:.1f}")

        c1, c2 = st.columns(2)
        h["jacket_type"]   = c1.text_input("Jacket Type",   value=h["jacket_type"])
        h["agitator_type"] = c2.text_input("Agitator Type", value=h["agitator_type"])

        c1, c2, c3 = st.columns(3)
        h["design_code"]     = c1.text_input("Design Code",     value=h["design_code"])
        h["design_pressure"] = c2.text_input("Design Pressure", value=h["design_pressure"])
        h["design_temp"]     = c3.text_input("Design Temp",     value=h["design_temp"])

        c1, c2 = st.columns(2)
        h["moc_shell"]  = c1.text_input("MOC – Shell",  value=h["moc_shell"])
        h["moc_jacket"] = c2.text_input("MOC – Jacket", value=h["moc_jacket"])
        st.divider()

        st.markdown("##### Prepared By / Checked By")
        staff_list = [""] + st.session_state.master_data.get("staff", [])
        c1, c2 = st.columns(2)
        pb_idx = staff_list.index(h["prepared_by"]) if h["prepared_by"] in staff_list else 0
        cb_idx = staff_list.index(h["checked_by"])  if h["checked_by"]  in staff_list else 0
        h["prepared_by"] = c1.selectbox("Prepared By", staff_list, index=pb_idx)
        h["checked_by"]  = c2.selectbox("Checked By",  staff_list, index=cb_idx)
        h["notes"] = st.text_area("Internal Notes (not printed in quote)", value=h["notes"], height=60)
        st.divider()

        st.markdown("##### 📝 Quotation Content Customisation")
        st.caption("All sections below print directly into the customer quotation. Edit per client URS / requirements.")

        with st.expander("📦 Section 3 — Scope of Supply & Exclusions", expanded=False):
            st.markdown("**Scope items** — one per line, prints as bullet list")
            h["scope_items"] = st.text_area(
                "Scope items", value=h.get("scope_items",""), height=220,
                help="One item per line. Each line becomes a bullet point in the quotation.",
                label_visibility="collapsed",
            )
            st.markdown("**Exclusions** — one per line")
            h["scope_exclusions"] = st.text_area(
                "Exclusions", value=h.get("scope_exclusions",""), height=150,
                help="One exclusion per line. Printed under 'Not in Scope' in Section 3.",
                label_visibility="collapsed",
            )

        with st.expander("🔬 Section 4 — Manufacturing & Quality", expanded=False):
            st.markdown("**Opening paragraph**")
            h["quality_intro"] = st.text_area(
                "Quality intro", value=h.get("quality_intro",""), height=80,
                label_visibility="collapsed",
            )
            st.markdown("**Quality points** — one per line, prints as bullet list")
            h["quality_points"] = st.text_area(
                "Quality points", value=h.get("quality_points",""), height=280,
                help="One point per line. Add/remove/edit based on equipment type and client URS.",
                label_visibility="collapsed",
            )

        with st.expander("📋 Section 5 — Documentation Deliverables", expanded=False):
            st.markdown("**Document list** — one per line")
            h["doc_deliverables"] = st.text_area(
                "Documentation", value=h.get("doc_deliverables",""), height=250,
                help="One document per line. Add client-specific requirements like ASME data report, 3.1 certs, etc.",
                label_visibility="collapsed",
            )

        # ── NEW: Detailed Equipment Specification ──────────────────────────────
        # Replaces the flat key-value Section 2 with a proper engineering spec sheet.
        # Fields left blank simply don't print — clean for storage tanks etc.
        with st.expander("📐 Section 2 — Detailed Equipment Specification", expanded=False):
            st.caption("This prints as the Equipment Specification Sheet in your customer quotation. Leave any field blank to skip it.")

            st.markdown("**Process / Service**")
            h["service_fluid"] = st.text_input(
                "Service fluid / Process",
                value=h.get("service_fluid", ""),
                placeholder="e.g. API intermediate / Solvent / Water",
                key="spec_service",
            )

            st.markdown("**Vessel Dimensions**")
            wv_c1, wv_c2 = st.columns(2)
            _cap = float(h.get("capacity_ltrs", 0) or 0)
            _wv_default = float(h.get("working_vol_ltrs", 0) or 0)
            if _wv_default == 0 and _cap > 0:
                # Auto-suggest 80% of capacity if not yet set
                _wv_default = round(_cap * 0.8, 0)
            h["working_vol_ltrs"] = wv_c1.number_input(
                "Working Volume (Ltrs)",
                value=_wv_default, min_value=0.0, step=50.0,
                help=f"Typically 80% of gross capacity. Suggested: {_cap * 0.8:.0f} L",
                key="spec_wv",
            )

            st.markdown("**Design Conditions**")
            dc1, dc2 = st.columns(2)
            h["design_pressure_shell"] = dc1.text_input(
                "Design Pressure — Shell",
                value=h.get("design_pressure_shell", h.get("design_pressure", "FV to 4.5 Bar")),
                placeholder="e.g. FV to 4.5 Bar(g)",
                key="spec_dps",
            )
            h["design_pressure_jacket"] = dc2.text_input(
                "Design Pressure — Jacket",
                value=h.get("design_pressure_jacket", "6.0 Bar / FV"),
                placeholder="e.g. 6.0 Bar(g) / FV",
                key="spec_dpj",
            )
            dc3, dc4, dc5 = st.columns(3)
            h["hydrotest_pressure"] = dc3.text_input(
                "Hydrotest Pressure",
                value=h.get("hydrotest_pressure", "1.5× Design Pressure"),
                key="spec_hyd",
            )
            h["corrosion_allowance_mm"] = dc4.number_input(
                "Corrosion Allowance (mm)",
                value=float(h.get("corrosion_allowance_mm", 1.5)),
                min_value=0.0, max_value=10.0, step=0.5,
                key="spec_ca",
            )
            h["joint_efficiency"] = dc5.number_input(
                "Joint Efficiency",
                value=float(h.get("joint_efficiency", 0.85)),
                min_value=0.0, max_value=1.0, step=0.05,
                help="0.85 = spot RT, 1.00 = full RT, 0.70 = no RT",
                key="spec_je",
            )

            st.markdown("**Heating / Cooling Medium**")
            hc1, hc2 = st.columns(2)
            h["heating_medium"] = hc1.text_input(
                "Heating Medium",
                value=h.get("heating_medium", "Steam / Hot Water / Thermic Fluid"),
                key="spec_hm",
            )
            h["cooling_medium"] = hc2.text_input(
                "Cooling Medium",
                value=h.get("cooling_medium", "Chilled Water / Brine"),
                key="spec_cm",
            )

            st.markdown("**Agitator & Drive Assembly**  _(leave blank for non-agitated equipment)_")
            ag1, ag2 = st.columns(2)
            h["agitator_rpm"] = ag1.text_input(
                "Agitator RPM",
                value=h.get("agitator_rpm", ""),
                placeholder="e.g. 60 RPM",
                key="spec_rpm",
            )
            h["motor_hp"] = ag2.text_input(
                "Motor Rating",
                value=h.get("motor_hp", ""),
                placeholder="e.g. 7.5 HP / 5.5 kW",
                key="spec_hp",
            )
            ag3, ag4 = st.columns(2)
            h["motor_make"] = ag3.text_input(
                "Motor Make",
                value=h.get("motor_make", "ABB / Crompton / Bharat Bijlee"),
                key="spec_mm",
            )
            h["gearbox_make"] = ag4.text_input(
                "Gearbox Make",
                value=h.get("gearbox_make", "Bonfiglioli / Premium / Elecon"),
                key="spec_gm",
            )
            ag5, ag6 = st.columns(2)
            h["gearbox_ratio"] = ag5.text_input(
                "Gearbox Ratio",
                value=h.get("gearbox_ratio", ""),
                placeholder="e.g. 1:25",
                key="spec_gr",
            )
            h["seal_type"] = ag6.text_input(
                "Mechanical Seal Type",
                value=h.get("seal_type", "Single mechanical seal"),
                placeholder="e.g. Single / Double, Cartridge",
                key="spec_st",
            )
            h["seal_make"] = st.text_input(
                "Mechanical Seal Make",
                value=h.get("seal_make", "EagleBurgmann / John Crane / Flowserve"),
                key="spec_sk",
            )

     
        with st.expander("💰 Section 6 — Commercial Terms", expanded=False):
            h["surface_finish"]    = st.text_input("Surface Finish (shown in Tech Basis)", value=h.get("surface_finish","Internal: Ra ≤ 0.8 μm  |  External: Buffed"))
            h["price_basis"]       = st.text_area("Price Basis",        value=h.get("price_basis",""),  height=60)
            h["gst_clause"]        = st.text_area("GST / Taxes Clause", value=h.get("gst_clause",""),   height=60)
            h["payment_terms"]     = st.text_area("Payment Terms",      value=h.get("payment_terms",""),height=60)
            c1, c2 = st.columns(2)
            h["delivery_weeks"]    = c1.text_input("Delivery (weeks)",   value=h.get("delivery_weeks","12–16"))
            h["delivery_note"]     = c2.text_input("Delivery note",      value=h.get("delivery_note","Subject to availability of raw material at time of order."))
            h["offer_validity"]    = st.text_area("Offer Validity",      value=h.get("offer_validity",""), height=60)
            h["warranty_clause"]   = st.text_area("Warranty",            value=h.get("warranty_clause",""),height=60)
            h["inspection_clause"] = st.text_area("Inspection Rights",   value=h.get("inspection_clause",""),height=60)
            h["special_notes"]     = st.text_area("Special Notes / Additional Conditions", value=h.get("special_notes",""), height=80)

        _save_draft_bar("f1")

    # ── F2: PLATES & PARTS ─────────────────────────────────────────────────────
    with f2:
        qtn_d = h.get("qtn_number", "") or "New"
        n_p   = len(st.session_state.est_parts)
        rm_t  = sum(p.get("amount", 0) for p in st.session_state.est_parts)
        if n_p > 0:
            st.success(f"**{qtn_d}** — {h.get('customer_name', 'no customer')}  |  {n_p} parts  |  RM ₹{rm_t:,.0f}")
        else:
            st.info(f"**{qtn_d}** — No parts yet. Load from Tab 1️⃣ or add below.")

        st.divider()
        st.markdown("##### Add / Edit Fabricated Parts")
        st.caption("Select Part Type → only the required dimension inputs appear → click Add Part.")

        edit_pidx = st.session_state.get("edit_part_idx")
        if edit_pidx is not None and isinstance(edit_pidx, int) and 0 <= edit_pidx < len(st.session_state.est_parts):
            editing_part = st.session_state.est_parts[edit_pidx]
        else:
            editing_part = None; st.session_state["edit_part_idx"] = None; edit_pidx = None

        ek = f"e{edit_pidx}_" if edit_pidx is not None else "new_"

        with st.container(border=True):
            if edit_pidx is not None and editing_part is not None:
                st.info(f"Editing row {edit_pidx + 1}: **{st.session_state.est_parts[edit_pidx].get('name', '')}** — update values then click Update Part")

            rc1, rc2, rc3 = st.columns(3)
            pt_keys = list(PART_TYPES.keys())
            def_pt  = editing_part.get("part_type", pt_keys[0]) if editing_part else pt_keys[0]
            if def_pt not in pt_keys:
                def_pt = pt_keys[0]
            p_name  = rc1.text_input("Part Name", value=editing_part.get("name", "") if editing_part else "", placeholder="e.g. Main Shell", key=f"{ek}pn")
            p_type  = rc2.selectbox("Part Type",  pt_keys, index=pt_keys.index(def_pt), key=f"{ek}pt")
            p_group = rc3.selectbox("Group",       all_groups,
                                    index=all_groups.index(editing_part.get("group", "SHELL")) if editing_part and editing_part.get("group", "SHELL") in all_groups else 0,
                                    key=f"{ek}pg")

            rc4, rc5, rc6 = st.columns(3)
            mat_list   = list(DENSITY.keys())
            def_mat    = editing_part.get("material", "SS316L") if editing_part else "SS316L"
            p_material = rc4.selectbox("Material", mat_list, index=mat_list.index(def_mat) if def_mat in mat_list else 0, key=f"{ek}pm")
            p_code     = rc5.selectbox("RM Code (for rate)", plate_rm or ["—"],
                                       index=plate_rm.index(editing_part.get("item_code", "")) if editing_part and editing_part.get("item_code", "") in plate_rm else 0,
                                       key=f"{ek}pc")
            # Pre-fill rate override with the actual stored rate when editing.
            # User can leave it (uses this rate), change it, or set to 0 to fall back to master rate.
            _def_rate_ov = float(editing_part.get("rate", 0)) if editing_part is not None else 0.0
            p_rate_ov = rc6.number_input(
                "Rate ₹/kg  (0 = use master rate)",
                value=_def_rate_ov, min_value=0.0, key=f"{ek}pr",
            )
            pt_info    = PART_TYPES[p_type]
            fn_key     = pt_info["fn"]
            is_derived = pt_info.get("qty_derived", False)

            def_scrap = float(editing_part.get("scrap_pct", DEFAULT_SCRAP.get(fn_key, 5.0))) if editing_part else DEFAULT_SCRAP.get(fn_key, 5.0)
            sc_col1, sc_col2 = st.columns([3,1])
            if is_derived:
                sc_col1.info("Qty is auto-calculated from geometry (shell height ÷ pitch).")
                p_qty = 1.0
            else:
                p_qty = sc_col1.number_input("Qty", value=float(editing_part.get("qty", 1)) if editing_part else 1.0, min_value=1.0, step=1.0, key=f"{ek}pq")
            p_scrap = sc_col2.number_input(
                "Scrap %",
                value=def_scrap,
                min_value=0.0, max_value=50.0, step=0.5,
                key=f"{ek}scrap",
                help=f"Default for {p_type}: {DEFAULT_SCRAP.get(fn_key,5)}%.",
            )

            needed   = pt_info["fields"]
            dim_cols = st.columns(min(len(needed), 6))
            dims     = {}
            for i, (field, label) in enumerate(needed):
                def_val = float(editing_part.get("dims", {}).get(field, 0.0)) if editing_part else 0.0
                dims[field] = dim_cols[i % 6].number_input(label, value=def_val, min_value=0.0, step=1.0, key=f"{ek}d_{p_type}_{field}")

            if pt_info.get("has_checkbox"):
                lc1, lc2, lc3 = st.columns(3)
                def_dish = bool(editing_part.get("dims", {}).get("cover_bottom_dish", False)) if editing_part else False
                cover_dish = lc1.checkbox(
                    "Include bottom dish limpet",
                    value=def_dish,
                    key=f"{ek}limpet_dish",
                )
                dims["cover_bottom_dish"] = cover_dish

                if all(dims.get(k, 0) > 0 for k in ["shell_id_mm","shell_thk_mm","shell_ht_mm","pipe_od_mm","pipe_thk_mm","pitch_mm"]):
                    _, prev_wt, prev_len = geom_limpet_coil(
                        dims["shell_id_mm"], dims["shell_thk_mm"], dims["shell_ht_mm"],
                        dims["pipe_od_mm"], dims["pipe_thk_mm"], dims["pitch_mm"],
                        cover_bottom_dish=cover_dish,
                        density=DENSITY.get(p_material, 8000),
                        scrap=p_scrap/100.0,
                    )
                    lc2.metric("Coil Length (m)", f"{prev_len:.2f} m")
                    lc3.metric("Est. Weight (kg)", f"{prev_wt:.2f} kg")

            btn_c1, btn_c2 = st.columns([3, 1])
            add_btn    = btn_c1.button("➕ Add Part" if not editing_part else "✅ Update Part", type="primary", use_container_width=True)
            cancel_btn = btn_c2.button("✖ Cancel", use_container_width=True) if editing_part else False

            if cancel_btn:
                st.session_state["edit_part_idx"] = None

            if add_btn:
                density  = DENSITY.get(p_material, 8000)
                fn       = pt_info["fn"]
                rm       = rm_master.get(p_code, {})
                rate     = p_rate_ov if p_rate_ov > 0 else rm.get("rate", 0)
                wt, total_wt, used_qty = calc_weight(fn, dims, density, p_qty, scrap_pct=p_scrap)
                if wt == 0:
                    st.warning("⚠️ Weight is zero — check all dimension inputs.")
                new_part = dict(
                    name=p_name, part_type=p_type, group=p_group,
                    material=p_material, item_code=p_code, dims=dims,
                    qty=used_qty, net_wt_kg=wt, total_wt_kg=total_wt,
                    rate=rate, amount=round(total_wt * rate, 2),
                    scrap_pct=p_scrap,
                )
                if editing_part is not None and edit_pidx is not None:
                    st.session_state.est_parts[edit_pidx] = new_part
                    st.session_state["edit_part_idx"] = None
                    st.success(f"✅ Updated: {p_name}")
                else:
                    st.session_state.est_parts.append(new_part)
                    st.success(f"✅ Added: {p_name}  |  {total_wt:.2f} kg  |  ₹{total_wt * rate:,.0f}")
                st.session_state["edit_part_idx"] = None

        if st.session_state.est_parts:
            st.markdown("---")
            st.markdown("**Parts list — click ✏️ to edit, 🗑️ to delete**")
            # Column headers
            hc = st.columns([0.5, 2.5, 2, 1.5, 1, 0.8, 0.8, 1.5, 2, 0.5, 0.5])
            for col, lbl in zip(hc, ["#", "Name", "Part Type", "Group", "Material",
                                      "Qty", "Scrap", "Weight", "Amount", "✏️", "🗑️"]):
                col.caption(f"**{lbl}**")

            for idx, p in enumerate(st.session_state.est_parts):
                # Highlight the row currently being edited
                _is_editing_row = (st.session_state.get("edit_part_idx") == idx)
                _prefix = "🟡 " if _is_editing_row else ""
                _bold   = "**" if _is_editing_row else ""

                c0, c1, c2, c3, c4, c5, c6, c7, c8, c9, c10 = st.columns(
                    [0.5, 2.5, 2, 1.5, 1, 0.8, 0.8, 1.5, 2, 0.5, 0.5])
                c0.write(f"{_prefix}**{idx + 1}**")
                c1.write(f"{_bold}{p.get('name', '')}{_bold}")
                c2.write(f"{_bold}{p.get('part_type', '')[:20]}{_bold}")
                c3.write(f"{_bold}{p.get('group', '')}{_bold}")
                c4.write(f"{_bold}{p.get('material', '')}{_bold}")
                c5.write(f"{_bold}{p.get('qty', 1):.1f}{_bold}")
                sc = p.get("scrap_pct", None)
                c6.write(f"{_bold}{(f'{sc:.0f}%' if sc is not None else '—')}{_bold}")
                c7.write(f"{_bold}{p.get('total_wt_kg', 0):.1f} kg{_bold}")
                c8.write(f"{_bold}₹{p.get('amount', 0):,.0f}{_bold}")
                if c9.button("✏️", key=f"ep_{idx}", help=f"Edit {p.get('name', '')}"):
                    st.session_state["edit_part_idx"] = idx
                    st.rerun()
                if c10.button("🗑️", key=f"dp_{idx}", help=f"Delete row {idx + 1}"):
                    st.session_state.est_parts.pop(idx)
                    st.session_state["edit_part_idx"] = None
                    st.rerun()
                
            tot_wt  = sum(p.get("total_wt_kg", 0) for p in st.session_state.est_parts)
            tot_amt = sum(p.get("amount", 0) for p in st.session_state.est_parts)
            st.success(f"**Total — Weight: {tot_wt:,.1f} kg  |  Amount: ₹{tot_amt:,.0f}**")

            if st.button("🗑️ Clear All Parts"):
                st.session_state.est_parts = []
                st.session_state["edit_part_idx"] = None
                st.rerun()
        _save_draft_bar("f2")
                  
    # ── F3: PIPES & FLANGES ────────────────────────────────────────────────────
    
    with f3:
        st.markdown("##### Nozzle Pipes")

        # Are we editing an existing pipe row?
        _epi = st.session_state.get("edit_pipe_idx")
        _editing_pipe = None
        if _epi is not None and 0 <= _epi < len(st.session_state.est_pipes):
            _editing_pipe = st.session_state.est_pipes[_epi]

        with st.container(border=True):
            if _editing_pipe is not None:
                st.info(f"Editing pipe row #{_epi + 1}: **{_editing_pipe.get('name', '')}**")

            # Pre-fill values from editing row
            _def_name = _editing_pipe.get("name", "") if _editing_pipe else ""
            _def_code = _editing_pipe.get("item_code", "") if _editing_pipe else ""
            _def_len  = float(_editing_pipe.get("length_m", 0.2)) if _editing_pipe else 0.2
            _def_qty  = int(_editing_pipe.get("qty", 1)) if _editing_pipe else 1
            _def_rate = float(_editing_pipe.get("rate", 0)) if _editing_pipe else 0.0

            # Key suffix changes when editing so widgets reset cleanly
            _pk = f"ep{_epi}_" if _epi is not None else "pp_new_"

            pc1, pc2, pc3, pc4, pc5 = st.columns(5)
            pp_name = pc1.text_input("Description", value=_def_name, placeholder='e.g. 2" Nozzle', key=f"{_pk}name")
            _pipe_opts = pipe_rm or ["—"]
            _code_idx = _pipe_opts.index(_def_code) if _def_code in _pipe_opts else 0
            pp_code = pc2.selectbox("Pipe Size", _pipe_opts, index=_code_idx, key=f"{_pk}code")
            pp_len  = pc3.number_input("Length (m)", value=_def_len, min_value=0.0, step=0.1, key=f"{_pk}len")
            pp_qty  = pc4.number_input("Qty", value=_def_qty, min_value=1, step=1, key=f"{_pk}qty")
            pp_rate = pc5.number_input("Rate Override (0=master)", value=_def_rate, min_value=0.0, key=f"{_pk}rate")

            if pipe_rm:
                rm = rm_master.get(pp_code, {})
                st.caption(f"Selected: {rm.get('description', '')} | {rm.get('unit_wt_kg_per_m', 0)} kg/m | Rate: ₹{rm.get('rate', 0)}/kg")

            bcol1, bcol2 = st.columns([3, 1])
            _btn_label = "✅ Update Pipe" if _editing_pipe else "➕ Add Pipe"
            if bcol1.button(_btn_label, type="primary", key=f"{_pk}btn", use_container_width=True):
                rm   = rm_master.get(pp_code, {})
                rate = pp_rate if pp_rate > 0 else rm.get("rate", 0)
                wpm  = rm.get("unit_wt_kg_per_m") or 0
                wt   = wpm * pp_len * 1.05 * pp_qty
                new_pipe = dict(name=pp_name, item_code=pp_code, length_m=pp_len, qty=pp_qty,
                                 wt_per_m=wpm, total_wt_kg=round(wt, 3), rate=rate,
                                 amount=round(wt * rate, 2))
                if _editing_pipe is not None:
                    st.session_state.est_pipes[_epi] = new_pipe
                    st.session_state["edit_pipe_idx"] = None
                    st.success(f"✅ Updated pipe row #{_epi + 1}")
                else:
                    st.session_state.est_pipes.append(new_pipe)
                    st.success(f"✅ Added: {pp_name}")
                st.rerun()
            if _editing_pipe and bcol2.button("✖ Cancel", key=f"{_pk}cancel", use_container_width=True):
                st.session_state["edit_pipe_idx"] = None
                st.rerun()

        if st.session_state.est_pipes:
            st.markdown("**Pipes list**")
            hc = st.columns([0.5, 3, 2, 1, 0.8, 1.2, 1.5, 0.5, 0.5])
            for col, lbl in zip(hc, ["#", "Description", "Code", "Len(m)", "Qty",
                                      "Wt(kg)", "Amount", "✏️", "🗑️"]):
                col.caption(f"**{lbl}**")

            for idx, p in enumerate(st.session_state.est_pipes):
                c0, c1, c2, c3, c4, c5, c6, c7, c8 = st.columns(
                    [0.5, 3, 2, 1, 0.8, 1.2, 1.5, 0.5, 0.5])
                c0.write(f"**{idx + 1}**")
                c1.write(p.get("name", ""))
                c2.write(p.get("item_code", ""))
                c3.write(f"{p.get('length_m', 0):.2f}")
                c4.write(f"{p.get('qty', 1)}")
                c5.write(f"{p.get('total_wt_kg', 0):.1f}")
                c6.write(f"₹{p.get('amount', 0):,.0f}")
                if c7.button("✏️", key=f"epipe_{idx}", help=f"Edit row {idx + 1}"):
                    st.session_state["edit_pipe_idx"] = idx
                    st.rerun()
                if c8.button("🗑️", key=f"dpipe_{idx}", help=f"Delete row {idx + 1}"):
                    st.session_state.est_pipes.pop(idx)
                    st.session_state["edit_pipe_idx"] = None
                    st.rerun()

            st.success(f"Total Pipes: ₹{sum(p['amount'] for p in st.session_state.est_pipes):,.0f}")
            if st.button("🗑️ Clear All Pipes", key="clr_pipes"):
                st.session_state.est_pipes = []
                st.session_state["edit_pipe_idx"] = None
                st.rerun()
        
        st.divider()
        st.markdown("##### Flanges & Fittings")

        _efi = st.session_state.get("edit_flg_idx")
        _editing_flg = None
        if _efi is not None and 0 <= _efi < len(st.session_state.est_flanges):
            _editing_flg = st.session_state.est_flanges[_efi]

        with st.container(border=True):
            if _editing_flg is not None:
                st.info(f"Editing flange row #{_efi + 1}: **{_editing_flg.get('name', '')}**")

            _def_name = _editing_flg.get("name", "") if _editing_flg else ""
            _def_code = _editing_flg.get("item_code", "") if _editing_flg else ""
            _def_qty  = int(_editing_flg.get("qty", 1)) if _editing_flg else 1
            _def_rate = float(_editing_flg.get("rate", 0)) if _editing_flg else 0.0

            _fk = f"ef{_efi}_" if _efi is not None else "fl_new_"

            fl1, fl2, fl3, fl4 = st.columns(4)
            fl_name = fl1.text_input("Description", value=_def_name, placeholder='e.g. 4" #150 Flange', key=f"{_fk}name")
            _flg_opts = flg_rm or ["—"]
            _flg_idx = _flg_opts.index(_def_code) if _def_code in _flg_opts else 0
            fl_code = fl2.selectbox("Flange Size", _flg_opts, index=_flg_idx, key=f"{_fk}code")
            fl_qty  = fl3.number_input("Qty", value=_def_qty, min_value=1, step=1, key=f"{_fk}qty")
            fl_rate = fl4.number_input("Rate Override (0=master)", value=_def_rate, min_value=0.0, key=f"{_fk}rate")

            bcol1, bcol2 = st.columns([3, 1])
            _btn_label = "✅ Update Flange" if _editing_flg else "➕ Add Flange"
            if bcol1.button(_btn_label, type="primary", key=f"{_fk}btn", use_container_width=True):
                rm   = rm_master.get(fl_code, {})
                rate = fl_rate if fl_rate > 0 else rm.get("rate", 0)
                wt   = ((rm.get("unit_wt_kg_per_m") or 0) * 1.15) * fl_qty
                new_flg = dict(name=fl_name, item_code=fl_code, qty=fl_qty,
                               total_wt_kg=round(wt, 3), rate=rate,
                               amount=round(wt * rate, 2))
                if _editing_flg is not None:
                    st.session_state.est_flanges[_efi] = new_flg
                    st.session_state["edit_flg_idx"] = None
                    st.success(f"✅ Updated flange row #{_efi + 1}")
                else:
                    st.session_state.est_flanges.append(new_flg)
                    st.success(f"✅ Added: {fl_name}")
                st.rerun()
            if _editing_flg and bcol2.button("✖ Cancel", key=f"{_fk}cancel", use_container_width=True):
                st.session_state["edit_flg_idx"] = None
                st.rerun()

        if st.session_state.est_flanges:
            st.markdown("**Flanges list**")
            hc = st.columns([0.5, 3, 2, 1, 1.2, 1.5, 0.5, 0.5])
            for col, lbl in zip(hc, ["#", "Description", "Code", "Qty", "Wt(kg)",
                                      "Amount", "✏️", "🗑️"]):
                col.caption(f"**{lbl}**")

            for idx, p in enumerate(st.session_state.est_flanges):
                c0, c1, c2, c3, c4, c5, c6, c7 = st.columns(
                    [0.5, 3, 2, 1, 1.2, 1.5, 0.5, 0.5])
                c0.write(f"**{idx + 1}**")
                c1.write(p.get("name", ""))
                c2.write(p.get("item_code", ""))
                c3.write(f"{p.get('qty', 1)}")
                c4.write(f"{p.get('total_wt_kg', 0):.1f}")
                c5.write(f"₹{p.get('amount', 0):,.0f}")
                if c6.button("✏️", key=f"eflg_{idx}", help=f"Edit row {idx + 1}"):
                    st.session_state["edit_flg_idx"] = idx
                    st.rerun()
                if c7.button("🗑️", key=f"dflg_{idx}", help=f"Delete row {idx + 1}"):
                    st.session_state.est_flanges.pop(idx)
                    st.session_state["edit_flg_idx"] = None
                    st.rerun()

            st.success(f"Total Flanges: ₹{sum(p['amount'] for p in st.session_state.est_flanges):,.0f}")
            if st.button("🗑️ Clear All Flanges", key="clr_flg"):
                st.session_state.est_flanges = []
                st.session_state["edit_flg_idx"] = None
                st.rerun()

        _save_draft_bar("f3")
    # ── F_STRUCT: STRUCTURAL STEEL (Angles, Channels, Beams) ───────────────────
    with f_struct:
        st.markdown("##### Structural Steel — Angles, Channels & Beams")
        st.caption(
            "For supports, lugs, saddles, lifting frames and skid structures. "
            "Unit weights from IS 808 / IS 2062. Total weight includes 5% cutting & fitting allowance."
        )

        sc1, sc2, sc3 = st.columns(3)
        type_filter = sc1.selectbox(
            "Section Type", ["All", "Equal Angle", "Channel", "Beam"],
            key="st_type_filter",
        )
        moc_pick = sc2.selectbox("Material", ["MS", "SS"], key="st_moc")

        if type_filter == "All":
            section_options = list(STRUCTURAL_SECTIONS.keys())
        else:
            section_options = [k for k, v in STRUCTURAL_SECTIONS.items()
                               if v["type"] == type_filter]
        sec_pick = sc3.selectbox("Section", section_options, key="st_sec")

        sec_info = STRUCTURAL_SECTIONS.get(sec_pick, {})
        unit_wt = sec_info.get("unit_wt", 0)
        st.caption(f"Selected: **{sec_pick}** ({sec_info.get('type','')}) — Unit weight: **{unit_wt} kg/m**")

        with st.container(border=True):
            ic1, ic2, ic3, ic4, ic5 = st.columns(5)
            st_name = ic1.text_input(
                "Description",
                placeholder="e.g. Bottom support lug",
                key="st_name",
            )
            st_len = ic2.number_input(
                "Length (m)", value=1.0, min_value=0.0, step=0.1, key="st_len",
            )
            st_qty = ic3.number_input(
                "Qty", value=1, min_value=1, step=1, key="st_qty",
            )
            default_rate = STRUCT_DEFAULT_RATES.get(moc_pick, 75.0)
            st_rate = ic4.number_input(
                f"Rate ₹/kg (default {moc_pick}: {default_rate})",
                value=float(default_rate), min_value=0.0, step=5.0, key="st_rate",
            )
            st_group = ic5.selectbox(
                "Group",
                ["SUPPORT", "LUGS", "FRAME", "SADDLE", "STIFFNERS", "OTHER"],
                key="st_group",
            )

            prev_wt  = unit_wt * st_len * st_qty * 1.05
            prev_amt = prev_wt * st_rate
            pc1, pc2 = st.columns(2)
            pc1.metric("Total Weight", f"{prev_wt:.2f} kg")
            pc2.metric("Amount", f"₹{prev_amt:,.0f}")

            if st.button("➕ Add Structural Item", type="primary", key="add_struct"):
                total_wt = unit_wt * st_len * st_qty * 1.05
                amount = total_wt * st_rate
                st.session_state.est_struct.append(dict(
                    name=st_name or sec_pick,
                    section=sec_pick,
                    type=sec_info.get("type", ""),
                    material=moc_pick,
                    group=st_group,
                    unit_wt_kg_per_m=unit_wt,
                    length_m=st_len,
                    qty=st_qty,
                    total_wt_kg=round(total_wt, 3),
                    rate=st_rate,
                    amount=round(amount, 2),
                ))
                st.success(f"✅ Added: {sec_pick}  |  {total_wt:.2f} kg  |  ₹{amount:,.0f}")

        if st.session_state.est_struct:
            st.divider()
            st.markdown("**Structural items added**")
            for idx, s in enumerate(st.session_state.est_struct):
                lc1, lc2, lc3, lc4, lc5, lc6, lc7, lc8 = st.columns(
                    [2.5, 2, 1, 1, 1, 1.2, 1.5, 0.7])
                lc1.write(s.get("name", ""))
                lc2.write(s.get("section", ""))
                lc3.write(s.get("material", ""))
                lc4.write(f"{s.get('length_m', 0):.2f} m")
                lc5.write(f"× {s.get('qty', 1)}")
                lc6.write(f"{s.get('total_wt_kg', 0):.1f} kg")
                lc7.write(f"₹{s.get('amount', 0):,.0f}")
                if lc8.button("🗑️", key=f"st_del_{idx}", help="Remove"):
                    st.session_state.est_struct.pop(idx)
                    st.rerun()

            tot_wt = sum(s.get("total_wt_kg", 0) for s in st.session_state.est_struct)
            tot_amt = sum(s.get("amount", 0) for s in st.session_state.est_struct)
            st.success(
                f"**Total Structural — Weight: {tot_wt:,.1f} kg  |  "
                f"Amount: ₹{tot_amt:,.0f}**"
            )

            if st.button("🗑️ Clear All Structural", key="clr_struct"):
                st.session_state.est_struct = []
                st.rerun()
        else:
            st.info("No structural items added yet.")

        _save_draft_bar("f_struct")

    # ── F4: FABRICATION SERVICES ───────────────────────────────────────────────
    with f4:
        st.markdown("##### Fabrication Services — Geometry-Driven Cost")
        st.caption("Adjust rates if needed, then click ⚡ Auto-Calculate to generate all line items from shell dimensions in Tab 1.")

        dia = float(h.get("shell_dia_mm", 0))
        ht  = float(h.get("shell_ht_mm", 0))
        if dia > 0 and ht > 0:
            _, _, int_a, ext_a = calc_surface_areas(dia, ht)
            wm_prev = calc_weld_metres(dia, ht, has_jacket=bool(h.get("jacket_type", "")), has_agitator=bool(h.get("agitator_type", "")))
            gp1, gp2, gp3 = st.columns(3)
            gp1.metric("Internal Surface Area", f"{int_a:.3f} m²")
            gp2.metric("External Surface Area", f"{ext_a:.3f} m²")
            gp3.metric("Estimated Weld Length", f"{wm_prev:.1f} m")
        else:
            st.warning("⚠️ Enter Shell ID and Shell Height in Tab 1️⃣ first.")

        fr = st.session_state.fab_rates
        with st.container(border=True):
            st.markdown("**Fabrication Rates**")
            rc1, rc2, rc3 = st.columns(3)
            fr["cutting_pct_on_plates"] = rc1.number_input("Cutting % on plate RM", value=float(fr["cutting_pct_on_plates"]), min_value=0.0, step=0.5)
            fr["rolling_rate_per_m2"]   = rc2.number_input("Rolling ₹/m²",          value=float(fr["rolling_rate_per_m2"]),   min_value=0.0, step=50.0)
            moc = h.get("moc_shell", "SS316L")
            if moc in ("SS316L", "Ti", "C22", "Hastelloy"):
                fr["tig_weld_rate_per_m"] = rc3.number_input("TIG Weld ₹/m", value=float(fr["tig_weld_rate_per_m"]), min_value=0.0, step=50.0)
            else:
                fr["arc_weld_rate_per_m"] = rc3.number_input("ARC Weld ₹/m", value=float(fr["arc_weld_rate_per_m"]), min_value=0.0, step=50.0)

            rc4, rc5, rc6 = st.columns(3)
            fr["int_grind_rate_per_m2"] = rc4.number_input("Int. Grinding ₹/m²", value=float(fr["int_grind_rate_per_m2"]), min_value=0.0, step=50.0)
            fr["ext_buff_rate_per_m2"]  = rc5.number_input("Ext. Buffing ₹/m²",  value=float(fr["ext_buff_rate_per_m2"]),  min_value=0.0, step=50.0)
            fr["ep_rate_per_m2"]        = rc6.number_input("Electropolish ₹/m² (0=skip)", value=float(fr.get("ep_rate_per_m2", 0)), min_value=0.0, step=50.0)

            rc7, rc8, rc9 = st.columns(3)
            fr["assembly_fitting_hrs"]  = rc7.number_input("Assembly Hours",       value=float(fr["assembly_fitting_hrs"]), min_value=0.0, step=5.0)
            fr["assembly_rate_per_hr"]  = rc8.number_input("Assembly ₹/hr",        value=float(fr["assembly_rate_per_hr"]), min_value=0.0, step=50.0)
            fr["hydro_test_lumpsum"]    = rc9.number_input("Hydro Test ₹ (lumpsum)", value=float(fr["hydro_test_lumpsum"]), min_value=0.0, step=500.0)

            rc10, rc11, _ = st.columns(3)
            fr["dp_test_rate_per_m2"]   = rc10.number_input("DP Test ₹/m²",       value=float(fr["dp_test_rate_per_m2"]), min_value=0.0, step=10.0)
            fr["qa_doc_lumpsum"]        = rc11.number_input("QA & Docs ₹ (lumpsum)", value=float(fr["qa_doc_lumpsum"]),   min_value=0.0, step=500.0)

        col_auto, col_clear = st.columns([2, 1])
        if col_auto.button("⚡ Auto-Calculate All Fabrication Services", type="primary", use_container_width=True):
            if dia > 0 and ht > 0:
                st.session_state.est_fab = auto_fab_services(h, fr, st.session_state.est_parts)
                st.success(f"✅ Generated {len(st.session_state.est_fab)} line items  |  Total: ₹{sum(f['amount'] for f in st.session_state.est_fab):,.0f}")
            else:
                st.error("Enter Shell ID and Shell Height in Tab 1️⃣ first.")
        if col_clear.button("🗑️ Clear Fabrication", use_container_width=True):
            st.session_state.est_fab = []

        if st.session_state.est_fab:
            st.markdown("---")
            st.markdown("**Fabrication services — click ✏️ to edit, 🗑️ to delete**")
            hc = st.columns([0.5, 3.5, 3.5, 1.2, 1, 1.5, 0.5, 0.5])
            for col, lbl in zip(hc, ["#", "Service", "Basis", "Qty", "UOM",
                                      "Amount", "✏️", "🗑️"]):
                col.caption(f"**{lbl}**")

            fab_total = 0
            for idx, fs in enumerate(st.session_state.est_fab):
                c0, c1, c2, c3, c4, c5, c6, c7 = st.columns(
                    [0.5, 3.5, 3.5, 1.2, 1, 1.5, 0.5, 0.5])
                c0.write(f"**{idx + 1}**")
                c1.write(fs.get("service", ""))
                c2.caption(fs.get("basis", ""))
                c3.write(f"{fs.get('qty', '')}")
                c4.write(fs.get("uom", ""))
                c5.write(f"₹{fs.get('amount', 0):,.0f}")
                if c6.button("✏️", key=f"efab_{idx}", help=f"Edit row {idx + 1}"):
                    st.session_state["edit_fab_idx"] = idx
                    st.rerun()
                if c7.button("🗑️", key=f"dfab_{idx}", help=f"Delete row {idx + 1}"):
                    st.session_state.est_fab.pop(idx)
                    st.session_state["edit_fab_idx"] = None
                    st.rerun()
                fab_total += fs.get("amount", 0)

            st.success(f"**Total Fabrication Services: ₹{fab_total:,.0f}**")

        # ── Edit / Add fabrication line ──────────────────────────────────────
        _efabi = st.session_state.get("edit_fab_idx")
        _editing_fab = None
        if _efabi is not None and 0 <= _efabi < len(st.session_state.est_fab):
            _editing_fab = st.session_state.est_fab[_efabi]

        st.markdown(f"**{'✏️ Edit fabrication line #' + str(_efabi + 1) if _editing_fab else '➕ Add custom line'}**")
        with st.container(border=True):
            _def_svc   = _editing_fab.get("service", "") if _editing_fab else ""
            _def_basis = _editing_fab.get("basis", "") if _editing_fab else ""
            _def_qty   = float(_editing_fab.get("qty", 1)) if _editing_fab else 1.0
            _def_uom   = _editing_fab.get("uom", "LS") if _editing_fab else "LS"
            _def_rate  = float(_editing_fab.get("rate", 0)) if _editing_fab else 0.0
            _def_amt   = float(_editing_fab.get("amount", 0)) if _editing_fab else 0.0

            _fak = f"efab{_efabi}_" if _efabi is not None else "fab_new_"

            ma1, ma2 = st.columns(2)
            ma_svc   = ma1.text_input("Service description", value=_def_svc, key=f"{_fak}svc")
            ma_basis = ma2.text_input("Basis / note", value=_def_basis, key=f"{_fak}basis")

            ma3, ma4, ma5, ma6 = st.columns(4)
            ma_qty   = ma3.number_input("Qty",  value=_def_qty,  min_value=0.0, step=0.1,    key=f"{_fak}qty")
            ma_uom   = ma4.text_input("UOM",    value=_def_uom,                              key=f"{_fak}uom")
            ma_rate  = ma5.number_input("Rate", value=_def_rate, min_value=0.0, step=10.0,   key=f"{_fak}rate")
            ma_amt   = ma6.number_input("Amount ₹", value=_def_amt, min_value=0.0, step=100.0, key=f"{_fak}amt")

            bcol1, bcol2 = st.columns([3, 1])
            _btn_label = "✅ Update Line" if _editing_fab else "➕ Add Line"
            if bcol1.button(_btn_label, type="primary", key=f"{_fak}btn", use_container_width=True):
                new_line = {"service": ma_svc, "basis": ma_basis, "qty": ma_qty,
                            "uom": ma_uom, "rate": ma_rate, "amount": ma_amt}
                if _editing_fab is not None:
                    st.session_state.est_fab[_efabi] = new_line
                    st.session_state["edit_fab_idx"] = None
                    st.success(f"✅ Updated fab row #{_efabi + 1}")
                else:
                    st.session_state.est_fab.append(new_line)
                    st.success(f"✅ Added: {ma_svc}")
                st.rerun()
            if _editing_fab and bcol2.button("✖ Cancel", key=f"{_fak}cancel", use_container_width=True):
                st.session_state["edit_fab_idx"] = None
                st.rerun()

        _save_draft_bar("f4")

    
    # ── F5: BOUGHT-OUT & OH ────────────────────────────────────────────────────
    with f5:
        st.markdown("##### Bought-Out Items  _(Motor, Gearbox, Seal, Fasteners, Insulation, etc.)_")

        _eboi = st.session_state.get("edit_bo_idx")
        _editing_bo = None
        if _eboi is not None and 0 <= _eboi < len(st.session_state.est_bo):
            _editing_bo = st.session_state.est_bo[_eboi]

        with st.container(border=True):
            if _editing_bo is not None:
                st.info(f"Editing BO row #{_eboi + 1}: **{_editing_bo.get('name', '')}**")

            _def_name  = _editing_bo.get("name", "") if _editing_bo else ""
            _def_code  = _editing_bo.get("item_code", "") if _editing_bo else ""
            _def_qty   = int(_editing_bo.get("qty", 1)) if _editing_bo else 1
            _def_rate  = float(_editing_bo.get("rate", 0)) if _editing_bo else 0.0
            _def_group = _editing_bo.get("group", "BO") if _editing_bo else "BO"

            _bok = f"ebo{_eboi}_" if _eboi is not None else "bo_new_"

            b1, b2, b3, b4, b5 = st.columns(5)
            bo_desc  = b1.text_input("Description", value=_def_name, placeholder="e.g. 7.5HP Motor", key=f"{_bok}d")
            _bo_opts = bo_rm or ["—"]
            _bo_idx  = _bo_opts.index(_def_code) if _def_code in _bo_opts else 0
            bo_code  = b2.selectbox("BO Code", _bo_opts, index=_bo_idx, key=f"{_bok}c")
            bo_qty   = b3.number_input("Qty", value=_def_qty, min_value=1, step=1, key=f"{_bok}q")
            bo_rate  = b4.number_input("Rate Override (0=master)", value=_def_rate, min_value=0.0, key=f"{_bok}r")
            _grp_opts = ["BO", "FASTENERS", "INSULATION", "OTHER"]
            _grp_idx  = _grp_opts.index(_def_group) if _def_group in _grp_opts else 0
            bo_group = b5.selectbox("Group", _grp_opts, index=_grp_idx, key=f"{_bok}g")

            if bo_rm:
                rm = rm_master.get(bo_code, {})
                st.caption(f"Selected: {rm.get('description', '')} | Rate: ₹{rm.get('rate', 0):,.0f} | UOM: {rm.get('uom', '')}")

            bcol1, bcol2 = st.columns([3, 1])
            _btn_label = "✅ Update BO Item" if _editing_bo else "➕ Add BO Item"
            if bcol1.button(_btn_label, type="primary", key=f"{_bok}btn", use_container_width=True):
                rm   = rm_master.get(bo_code, {})
                rate = bo_rate if bo_rate > 0 else rm.get("rate", 0)
                new_bo = dict(name=bo_desc or rm.get("description", ""),
                              item_code=bo_code, qty=bo_qty, rate=rate,
                              amount=round(rate * bo_qty, 2), group=bo_group)
                if _editing_bo is not None:
                    st.session_state.est_bo[_eboi] = new_bo
                    st.session_state["edit_bo_idx"] = None
                    st.success(f"✅ Updated BO row #{_eboi + 1}")
                else:
                    st.session_state.est_bo.append(new_bo)
                    st.success(f"✅ Added: {bo_desc or rm.get('description', '')}")
                st.rerun()
            if _editing_bo and bcol2.button("✖ Cancel", key=f"{_bok}cancel", use_container_width=True):
                st.session_state["edit_bo_idx"] = None
                st.rerun()

        if st.session_state.est_bo:
            st.markdown("**Bought-Out items list**")
            hc = st.columns([0.5, 3, 2, 0.8, 1.2, 1.5, 1.2, 0.5, 0.5])
            for col, lbl in zip(hc, ["#", "Description", "Code", "Qty", "Rate",
                                      "Amount", "Group", "✏️", "🗑️"]):
                col.caption(f"**{lbl}**")

            for idx, b in enumerate(st.session_state.est_bo):
                c0, c1, c2, c3, c4, c5, c6, c7, c8 = st.columns(
                    [0.5, 3, 2, 0.8, 1.2, 1.5, 1.2, 0.5, 0.5])
                c0.write(f"**{idx + 1}**")
                c1.write(b.get("name", ""))
                c2.write(b.get("item_code", ""))
                c3.write(f"{b.get('qty', 1)}")
                c4.write(f"₹{b.get('rate', 0):,.0f}")
                c5.write(f"₹{b.get('amount', 0):,.0f}")
                c6.write(b.get("group", ""))
                if c7.button("✏️", key=f"ebo_{idx}", help=f"Edit row {idx + 1}"):
                    st.session_state["edit_bo_idx"] = idx
                    st.rerun()
                if c8.button("🗑️", key=f"dbo_{idx}", help=f"Delete row {idx + 1}"):
                    st.session_state.est_bo.pop(idx)
                    st.session_state["edit_bo_idx"] = None
                    st.rerun()

            st.success(f"Total Bought-Out: ₹{sum(b['amount'] for b in st.session_state.est_bo):,.0f}")
            if st.button("🗑️ Clear All BO", key="clr_bo"):
                st.session_state.est_bo = []
                st.session_state["edit_bo_idx"] = None
                st.rerun()

        st.divider()
        st.markdown("##### Additional Overheads  _(any cost not covered above)_")

        _eohi = st.session_state.get("edit_oh_idx")
        _editing_oh = None
        if _eohi is not None and 0 <= _eohi < len(st.session_state.est_oh):
            _editing_oh = st.session_state.est_oh[_eohi]

        with st.container(border=True):
            if _editing_oh is not None:
                st.info(f"Editing OH row #{_eohi + 1}: **{_editing_oh.get('description', '')}**")

            _def_code  = _editing_oh.get("oh_code", "") if _editing_oh else ""
            _def_qty   = float(_editing_oh.get("qty", 1)) if _editing_oh else 1.0
            _def_rate  = float(_editing_oh.get("rate", 0)) if _editing_oh else 0.0
            _def_desc  = _editing_oh.get("description", "") if _editing_oh else ""

            _ohk = f"eoh{_eohi}_" if _eohi is not None else "oh_new_"

            o1, o2, o3, o4 = st.columns(4)
            _oh_opts = oh_codes or ["—"]
            _oh_idx  = _oh_opts.index(_def_code) if _def_code in _oh_opts else 0
            oh_sel = o1.selectbox("OH Code", _oh_opts, index=_oh_idx, key=f"{_ohk}sel")
            oh_qty = o2.number_input("Qty / Hours / Area", value=_def_qty, min_value=0.0, step=1.0, key=f"{_ohk}q")
            oh_rov = o3.number_input("Rate Override (0=master)", value=_def_rate, min_value=0.0, key=f"{_ohk}r")
            oh_dov = o4.text_input("Description override (optional)", value=_def_desc, key=f"{_ohk}d")
            oh_inf = oh_master.get(oh_sel, {})
            st.caption(f"Selected: **{oh_inf.get('description', '')}** | Type: {oh_inf.get('oh_type', '')} | UOM: {oh_inf.get('uom', '')} | Rate: ₹{oh_inf.get('rate', 0):,.0f}")

            bcol1, bcol2 = st.columns([3, 1])
            _btn_label = "✅ Update Overhead" if _editing_oh else "➕ Add Overhead"
            if bcol1.button(_btn_label, type="primary", key=f"{_ohk}btn", use_container_width=True):
                rate = oh_rov if oh_rov > 0 else oh_inf.get("rate", 0)
                uom  = oh_inf.get("uom", "")
                desc = oh_dov or oh_inf.get("description", "")
                if uom == "%":
                    base   = sum(p.get("amount", 0) for p in st.session_state.est_parts) + \
                             sum(p.get("amount", 0) for p in st.session_state.est_pipes) + \
                             sum(p.get("amount", 0) for p in st.session_state.est_flanges)
                    amount = base * rate / 100
                else:
                    amount = rate * oh_qty
                new_oh = dict(oh_code=oh_sel, description=desc,
                              oh_type=oh_inf.get("oh_type", ""), uom=uom,
                              qty=oh_qty, rate=rate, amount=round(amount, 2))
                if _editing_oh is not None:
                    st.session_state.est_oh[_eohi] = new_oh
                    st.session_state["edit_oh_idx"] = None
                    st.success(f"✅ Updated OH row #{_eohi + 1}")
                else:
                    st.session_state.est_oh.append(new_oh)
                    st.success(f"✅ Added: {desc}")
                st.rerun()
            if _editing_oh and bcol2.button("✖ Cancel", key=f"{_ohk}cancel", use_container_width=True):
                st.session_state["edit_oh_idx"] = None
                st.rerun()

        if st.session_state.est_oh:
            st.markdown("**Overheads list**")
            hc = st.columns([0.5, 3, 1.5, 0.8, 1, 1.2, 1.5, 0.5, 0.5])
            for col, lbl in zip(hc, ["#", "Description", "Type", "UOM", "Qty",
                                      "Rate", "Amount", "✏️", "🗑️"]):
                col.caption(f"**{lbl}**")

            for idx, o in enumerate(st.session_state.est_oh):
                c0, c1, c2, c3, c4, c5, c6, c7, c8 = st.columns(
                    [0.5, 3, 1.5, 0.8, 1, 1.2, 1.5, 0.5, 0.5])
                c0.write(f"**{idx + 1}**")
                c1.write(o.get("description", ""))
                c2.write(o.get("oh_type", ""))
                c3.write(o.get("uom", ""))
                c4.write(f"{o.get('qty', 1):.1f}")
                c5.write(f"₹{o.get('rate', 0):,.0f}")
                c6.write(f"₹{o.get('amount', 0):,.0f}")
                if c7.button("✏️", key=f"eoh_{idx}", help=f"Edit row {idx + 1}"):
                    st.session_state["edit_oh_idx"] = idx
                    st.rerun()
                if c8.button("🗑️", key=f"doh_{idx}", help=f"Delete row {idx + 1}"):
                    st.session_state.est_oh.pop(idx)
                    st.session_state["edit_oh_idx"] = None
                    st.rerun()

            st.success(f"Total OH: ₹{sum(o['amount'] for o in st.session_state.est_oh):,.0f}")
            if st.button("🗑️ Clear All OH", key="clr_oh"):
                st.session_state.est_oh = []
                st.session_state["edit_oh_idx"] = None
                st.rerun()

        _save_draft_bar("f5")
    # ── F6: SUMMARY & SAVE ─────────────────────────────────────────────────────
    with f6:
        eq_info = EQUIPMENT_TYPES.get(h["equipment_type"], {})
        lo, hi  = eq_info.get("margin_hint", (10, 18))
        st.info(f"Suggested margin for **{h['equipment_type']}**: {lo}–{hi}%  |  Labour: **{eq_info.get('labour_norm', 'Medium')}**")

        qtn_now = h.get("qtn_number","") or ""
        cust_now = h.get("customer_name","") or "no customer"
        rev_now  = h.get("revision","R0")
        stat_now = h.get("status","Draft")
        if qtn_now:
            st.success(f"💾 **{qtn_now}** {rev_now} — {cust_now}  |  Status: {stat_now}  |  {len(st.session_state.est_parts)} parts  |  {len(st.session_state.est_fab)} fab lines")
        else:
            st.warning("⚠️ No Quotation Number set — go to Tab 1️⃣ Header and enter a QTN number before saving.")

        # Row 1: cost-build inputs
        s1, s2, s3, s4, s5, s6 = st.columns(6)
        h["profit_margin_pct"] = s1.number_input("Profit %",      value=float(h["profit_margin_pct"]), min_value=0.0, max_value=60.0, step=0.5)
        h["contingency_pct"]   = s2.number_input("Contingency %", value=float(h["contingency_pct"]),   min_value=0.0, max_value=20.0, step=0.5)
        h["engg_design_amt"]   = s3.number_input("Engg & ASME ₹", value=float(h["engg_design_amt"]),   min_value=0.0, step=1000.0)
        h["packing_amt"]       = s4.number_input("Packing ₹",     value=float(h["packing_amt"]),       min_value=0.0, step=500.0)
        h["freight_amt"]       = s5.number_input("Freight ₹",     value=float(h["freight_amt"]),       min_value=0.0, step=500.0)
        h["gst_pct"]           = s6.number_input("GST %",         value=float(h["gst_pct"]),           min_value=0.0, max_value=28.0, step=0.5)

        # Row 2: discount input (NEW — Option B)
        d1, _d2, _d3 = st.columns([1, 4, 1])
        h["discount_pct"] = d1.number_input(
            "Discount %  (on Ex-Works)",
            value=float(h.get("discount_pct", 0.0)),
            min_value=0.0, max_value=50.0, step=0.5,
            help="Commercial discount offered to customer. Reduces Ex-Works price and the GST charged on it.",
        )

        T = calc_totals(
            st.session_state.est_parts, st.session_state.est_pipes, st.session_state.est_flanges,
            st.session_state.est_struct,
            st.session_state.est_fab, st.session_state.est_bo, st.session_state.est_oh,
            h["profit_margin_pct"], h["contingency_pct"],
            h["packing_amt"], h["freight_amt"], h["gst_pct"], h["engg_design_amt"],
            discount_pct=h.get("discount_pct", 0.0),   # NEW for Option B
        )

        left, right = st.columns([3, 2])
        with left:
            st.markdown("**Cost Breakup**")

            # Build cost rows — Option B structure with Factory/Admin OH split,
            # Discount and Net Realisation lines.
            _cost_rows = [
                ("Plates & Parts",        T["tot_plates"]),
                ("Pipes",                 T["tot_pipes"]),
                ("Flanges",               T["tot_flanges"]),
                ("Structural Steel",      T["tot_struct"]),
                ("▶ Total Raw Material",  T["tot_rm"]),
                ("Fabrication Services",  T["tot_fab"]),
                ("Bought-Out Items",      T["tot_bo"]),
                ("Factory Overhead",      T["tot_factory_oh"]),
                ("Admin Overhead",        T["tot_admin_oh"]),
                ("Engg & ASME Design",    T["engg_design"]),
                ("▶ Total Mfg Cost",      T["tot_mfg"]),
                ("Contingency",           T["cont_amt"]),
                ("Profit",                T["profit_amt"]),
                ("Packing & Freight",     T["packing"] + T["freight"]),
                ("▶ Ex-Works Price",      T["ex_works"]),
            ]
            # Show Discount line only if there is a discount, to keep the table clean
            if T.get("discount_amt", 0) > 0:
                _cost_rows.append(("Discount",          -T["discount_amt"]))
                _cost_rows.append(("▶ Net Realisation", T["net_realisation"]))
            _cost_rows.append(("GST",                   T["gst_amt"]))
            _cost_rows.append(("▶ FOR Price",           T["for_price"]))

            cost_df = pd.DataFrame(_cost_rows, columns=["Component", "Amount (₹)"])
            cost_df["Amount (₹)"] = cost_df["Amount (₹)"].map(lambda x: f"₹{x:,.0f}")
            st.dataframe(cost_df, use_container_width=True, hide_index=True)

            # ── Drill-down: actual line items in each OH bucket ──────────────
            # Shows the estimator exactly what's inside Factory OH and Admin OH,
            # solves the "Additional Overheads is opaque" problem.
            with st.expander(
                f"🔍 Factory Overhead drill-down ({len(T.get('oh_lines_factory', []))} items, "
                f"₹{T.get('tot_factory_oh', 0):,.0f})"
            ):
                _fac = T.get("oh_lines_factory", [])
                if _fac:
                    _df_fac = pd.DataFrame([{
                        "Description": o.get("description", ""),
                        "Type":        o.get("oh_type", ""),
                        "Qty":         o.get("qty", 0),
                        "Rate (₹)":    f"₹{o.get('rate', 0):,.0f}",
                        "Amount (₹)":  f"₹{o.get('amount', 0):,.0f}",
                    } for o in _fac])
                    st.dataframe(_df_fac, use_container_width=True, hide_index=True)
                else:
                    st.caption("_No factory overhead items added. Add in Tab 6️⃣ Bought-Out & OH._")

            with st.expander(
                f"🔍 Admin Overhead drill-down ({len(T.get('oh_lines_admin', []))} items, "
                f"₹{T.get('tot_admin_oh', 0):,.0f})"
            ):
                _adm = T.get("oh_lines_admin", [])
                if _adm:
                    _df_adm = pd.DataFrame([{
                        "Description": o.get("description", ""),
                        "Type":        o.get("oh_type", ""),
                        "Qty":         o.get("qty", 0),
                        "Rate (₹)":    f"₹{o.get('rate', 0):,.0f}",
                        "Amount (₹)":  f"₹{o.get('amount', 0):,.0f}",
                    } for o in _adm])
                    st.dataframe(_df_adm, use_container_width=True, hide_index=True)
                else:
                    st.caption("_No admin overhead items. DOCS/MISC items in Tab 6 land here._")

        with right:
            st.markdown("**Margin Health**")
            for label, val, lo_t, hi_t in [
                ("RM %",      T["rm_pct"],           45, 60),
                ("Fab Svc %", T["fab_pct"],           15, 25),
                ("OH %",      T["oh_pct"],             8, 15),
                ("Profit %",  T["profit_pct_actual"], 12, 20),
            ]:
                icon = "✅" if lo_t <= val <= hi_t else "⚠️"
                st.write(f"{icon} **{label}** {val:.1f}%  _(target {lo_t}–{hi_t}%)_")
            for iss in margin_issues(T):
                st.warning(iss)
            if not margin_issues(T):
                st.success("All margins healthy!")
            st.markdown("**What-If: Ex-Works at Different Margins**")
            st.dataframe(pd.DataFrame([{
                "Margin": f"{m}%",
                "Ex-Works (₹)": f"₹{(T['cbm'] * (1 + m / 100) + T['packing'] + T['freight']):,.0f}",
            } for m in [8, 10, 12, 15, 18, 20]]), hide_index=True, use_container_width=True)

        st.divider()
        # Row 1: cost build-up metrics
        k1, k2, k3, k4, k5, k6 = st.columns(6)
        k1.metric("Raw Material", f"₹{T['tot_rm']:,.0f}")
        k2.metric("Fabrication",  f"₹{T['tot_fab']:,.0f}")
        k3.metric("Total Mfg",    f"₹{T['tot_mfg']:,.0f}")
        k4.metric("Ex-Works",     f"₹{T['ex_works']:,.0f}")
        k5.metric("GST",          f"₹{T['gst_amt']:,.0f}")
        k6.metric("FOR Price",    f"₹{T['for_price']:,.0f}")

        # Row 2: realisation & profit metrics (NEW — Option B)
        # Shows what B&G actually keeps after discount, and the gross profit.
        m1, m2, m3, m4 = st.columns(4)
        m1.metric(
            "Discount",
            f"₹{T.get('discount_amt', 0):,.0f}",
            delta=f"-{T.get('discount_pct', 0):.1f}%" if T.get('discount_amt', 0) > 0 else None,
            delta_color="inverse",
        )
        m2.metric("Net Realisation", f"₹{T.get('net_realisation', T['ex_works']):,.0f}")
        m3.metric(
            "Gross Profit",
            f"₹{T.get('gross_profit', 0):,.0f}",
            delta=f"{T.get('gross_profit_pct', 0):.1f}%",
        )
        m4.metric("Factory OH + Admin OH", f"₹{T.get('tot_factory_oh', 0) + T.get('tot_admin_oh', 0):,.0f}")
        st.divider()

        b1, b2, b3 = st.columns(3)

        if b1.button("💾 Save & Close", type="primary", use_container_width=True, disabled=not h["qtn_number"]):
            if _do_save(reset_after=True):
                st.rerun()

        if b2.button("🔄 Reset / New", use_container_width=True):
            _reset_form(); st.rerun()

        cust_data = next((c for c in clients if c["name"] == h.get("customer_name", "")), {})

        st.divider()
        st.markdown("**📄 Download Customer Quotation**")
        dl_col1, dl_col2 = st.columns(2)
        with dl_col1:
            st.caption("**Standard Quote** — clean single price. Use for most customers.")
            dl_col1.download_button(
                "📄 Download Standard Quote",
                generate_docx(h, cust_data, T, st.session_state.est_fab, show_breakup=False),
                file_name=f"{h.get('qtn_number','QTN')}_{h.get('revision','R0')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True, type="primary",
            )
        with dl_col2:
            st.caption("**With Scope Breakup** — use when customer asks for breakup.")
            dl_col2.download_button(
                "📋 Download Quote with Breakup",
                generate_docx(h, cust_data, T, st.session_state.est_fab, show_breakup=True),
                file_name=f"{h.get('qtn_number','QTN')}_{h.get('revision','R0')}_breakup.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        st.divider()
        st.markdown("**📊 Download Internal Estimation Fact Sheet** _(for cross-checking against Excel)_")
        st.caption("Shows all dimensions, geometry formulas used, weights, rates and cost summary across multiple sheets. Not for customers.")
        if OPENPYXL_OK:
            st.download_button(
                "📊 Download Estimation Fact Sheet (.xlsx)",
                generate_fact_sheet_xlsx(
                    h, st.session_state.est_parts, st.session_state.est_pipes,
                    st.session_state.est_flanges, st.session_state.est_struct,
                    st.session_state.est_fab,
                    st.session_state.est_bo, st.session_state.est_oh,
                    st.session_state.fab_rates, T,
                ),
                file_name=f"{h.get('qtn_number','QTN')}_{h.get('revision','R0')}_FactSheet.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
        else:
            st.warning("📊 Fact Sheet requires **openpyxl**. Add `openpyxl` to requirements.txt and redeploy.")

# ══════════════════════════════════════════════════════════════════════════════
# TAB: QUOTE EDITOR
# ══════════════════════════════════════════════════════════════════════════════
with TAB_QUOTE:
    h    = st.session_state.est_hdr
    _qno = h.get("qtn_number","") or ""
    _cno = h.get("customer_name","") or ""

    if not _qno:
        st.warning("⚠️ No estimation loaded. Go to ➕ New / Edit → Tab 1 and load or create an estimation first.")
    else:
        st.subheader(f"✍️ Quote Editor — {_qno}  |  {_cno}")
        st.caption("Edit every section below. What you see here is exactly what prints in the customer quotation. Download when ready.")

        with st.expander("📋 Section 1 — Offer & Customer Details", expanded=True):
            st.caption("Set in Tab 1 Header. Change fields there to update here.")
            si1, si2, si3 = st.columns(3)
            si1.markdown(f"**QTN:** {h.get('qtn_number','')}  |  **Rev:** {h.get('revision','R0')}")
            si2.markdown(f"**Customer:** {h.get('customer_name','')}")
            si3.markdown(f"**Equipment:** {h.get('equipment_desc','')}")
            si1.markdown(f"**Prepared By:** {h.get('prepared_by','')}  |  **Checked By:** {h.get('checked_by','')}")
            si2.markdown(f"**Status:** {h.get('status','Draft')}")
            si3.markdown(f"**Date:** {date.today().strftime('%d %B %Y')}")

        with st.expander("🔧 Section 2 — Technical Design Basis", expanded=False):
            t2c1, t2c2 = st.columns(2)
            h["surface_finish"]  = t2c1.text_input("Surface Finish", value=h.get("surface_finish","Internal: Ra ≤ 0.8 μm  |  External: Buffed"), key="qe_sf")
            h["design_code"]     = t2c2.text_input("Design Code",     value=h.get("design_code","ASME Sec VIII Div 1"), key="qe_dc")
            h["design_pressure"] = t2c1.text_input("Design Pressure", value=h.get("design_pressure","FV to 4.5 Bar"), key="qe_dp")
            h["design_temp"]     = t2c2.text_input("Design Temperature", value=h.get("design_temp","-50 to 250°C"), key="qe_dt")
            h["jacket_type"]     = t2c1.text_input("Jacket / Heating", value=h.get("jacket_type",""), key="qe_jt")
            h["agitator_type"]   = t2c2.text_input("Agitator / Drive", value=h.get("agitator_type",""), key="qe_at")

        with st.expander("📦 Section 3 — Scope of Supply", expanded=False):
            st.markdown("**Scope items — one per line**")
            h["scope_items"] = st.text_area("Scope items", value=h.get("scope_items",""), height=200, label_visibility="collapsed", key="qe_scope")
            st.markdown("**Exclusions — one per line**")
            h["scope_exclusions"] = st.text_area("Exclusions", value=h.get("scope_exclusions",""), height=140, label_visibility="collapsed", key="qe_excl")

        with st.expander("🔬 Section 4 — Manufacturing & Quality Assurance", expanded=False):
            st.markdown("**Opening paragraph**")
            h["quality_intro"] = st.text_area("Quality intro", value=h.get("quality_intro",""), height=80, label_visibility="collapsed", key="qe_qi")
            st.markdown("**Quality points — one per line**")
            h["quality_points"] = st.text_area("Quality points", value=h.get("quality_points",""), height=300, label_visibility="collapsed", key="qe_qp")

        with st.expander("📋 Section 5 — Documentation Deliverables", expanded=False):
            st.markdown("**Document list — one per line**")
            h["doc_deliverables"] = st.text_area("Documentation", value=h.get("doc_deliverables",""), height=280, label_visibility="collapsed", key="qe_dd")

        with st.expander("💰 Section 6 — Commercial Terms", expanded=False):
            st.markdown("**Payment Terms**")
            h["payment_terms"] = st.text_area("Payment Terms", height=70, label_visibility="collapsed", value=h.get("payment_terms",""), key="qe_pt")
            c1, c2 = st.columns(2)
            h["delivery_weeks"] = c1.text_input("Delivery (weeks)", value=h.get("delivery_weeks","12–16"), key="qe_dw")
            h["delivery_note"]  = c2.text_input("Delivery note",    value=h.get("delivery_note",""),       key="qe_dn")
            h["offer_validity"] = st.text_area("Offer Validity",    height=60, label_visibility="collapsed", value=h.get("offer_validity",""), key="qe_ov")
            st.markdown("**Warranty**")
            h["warranty_clause"] = st.text_area("Warranty", height=80, label_visibility="collapsed", value=h.get("warranty_clause",""), key="qe_wc")
            st.markdown("**Inspection Rights**")
            h["inspection_clause"] = st.text_area("Inspection", height=70, label_visibility="collapsed", value=h.get("inspection_clause",""), key="qe_ic")
            st.markdown("**Price Basis**")
            h["price_basis"] = st.text_area("Price Basis", height=70, label_visibility="collapsed", value=h.get("price_basis",""), key="qe_pb")
            st.markdown("**GST / Taxes Clause**")
            h["gst_clause"] = st.text_area("GST Clause", height=60, label_visibility="collapsed", value=h.get("gst_clause",""), key="qe_gc")
            st.markdown("**Special Notes / Additional Conditions**")
            h["special_notes"] = st.text_area("Special Notes", height=100, label_visibility="collapsed", value=h.get("special_notes",""), key="qe_sn")

        st.divider()

        # AI Assistant
        st.markdown("### 🤖 AI Assistant — Generate & Improve Quotation Content")
        st.caption("Claude reads your equipment parameters and generates professional content. Review and edit before downloading.")

        try:
            import anthropic as _anthropic
            try:
                _claude_key = st.secrets["ANTHROPIC_API_KEY"]
            except Exception:
                _claude_key = ""
            _claude_client = _anthropic.Anthropic(api_key=_claude_key) if _claude_key else None
        except Exception:
            _claude_client = None

        if not _claude_client:
            st.warning("⚠️ Add ANTHROPIC_API_KEY to Streamlit secrets to enable AI features.")
        else:
            _eq_context = f"""
Equipment Type: {h.get('equipment_type', '')}
Description: {h.get('equipment_desc', '')}
Capacity: {h.get('capacity_ltrs', '')} Litres
Shell ID: {h.get('shell_dia_mm', '')} mm | Height: {h.get('shell_ht_mm', '')} mm
MOC Shell: {h.get('moc_shell', 'SS316L')} | MOC Jacket: {h.get('moc_jacket', 'SS304')}
Jacket: {h.get('jacket_type', '')} | Agitator: {h.get('agitator_type', '')}
Design Code: {h.get('design_code', 'ASME Sec VIII Div 1')}
Design Pressure: {h.get('design_pressure', '')} | Temp: {h.get('design_temp', '')}
Customer: {h.get('customer_name', '')}
"""

            _ai_t1, _ai_t2, _ai_t3, _ai_t4 = st.tabs([
                "📦 Generate Scope",
                "🔬 Quality Points",
                "📝 Cover Narrative",
                "🔍 Review Terms",
            ])

            with _ai_t1:
                st.markdown("**Generate scope of supply specific to this equipment**")
                _c1, _c2 = st.columns(2)
                _pharma  = _c1.checkbox("Pharma / cGMP client", value=True, key="ai_pharma")
                _excl    = _c2.checkbox("Include exclusions list", value=True, key="ai_excl_cb")
                if st.button("⚡ Generate Scope", type="primary", key="ai_gen_scope"):
                    with st.spinner("Claude is writing scope items..."):
                        try:
                            _msg = _claude_client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=1000,
                                messages=[{"role": "user", "content": f"""You are a technical writer at B&G Engineering Industries, 
Hyderabad — Indian manufacturer of pharma-grade ASME process equipment.

Equipment:
{_eq_context}

Generate a precise scope of supply for this quotation.
{'This is for a pharma/API/cGMP client.' if _pharma else ''}

Output ONLY a plain list — one item per line, no bullets, no numbers, no headers.
Generate 8-12 items covering: vessel construction, nozzles, jacket/heating,
agitator/drive (if applicable), surface finish, supports, nameplate, testing.
{'Then add a blank line, then write EXCLUSIONS: followed by 5-7 exclusion items one per line.' if _excl else ''}"""}]
                            )
                            _result = _msg.content[0].text
                            if _excl and "EXCLUSIONS:" in _result:
                                _sp = _result.split("EXCLUSIONS:")
                                st.session_state["ai_scope_result"] = _sp[0].strip()
                                st.session_state["ai_excl_result"]  = _sp[1].strip() if len(_sp) > 1 else ""
                            else:
                                st.session_state["ai_scope_result"] = _result.strip()
                            st.success("✅ Generated — review and apply below")
                        except Exception as _e:
                            st.error(f"Claude API error: {_e}")

                if "ai_scope_result" in st.session_state:
                    st.markdown("**Generated Scope — edit if needed:**")
                    _es = st.text_area("Scope", value=st.session_state["ai_scope_result"], height=200, label_visibility="collapsed", key="ai_scope_edit")
                    if "ai_excl_result" in st.session_state:
                        st.markdown("**Generated Exclusions — edit if needed:**")
                        _ee = st.text_area("Excl", value=st.session_state["ai_excl_result"], height=120, label_visibility="collapsed", key="ai_excl_edit")
                    if st.button("✅ Apply to Quotation", type="primary", key="ai_apply_scope"):
                        h["scope_items"] = st.session_state.get("ai_scope_edit", st.session_state["ai_scope_result"])
                        if "ai_excl_result" in st.session_state:
                            h["scope_exclusions"] = st.session_state.get("ai_excl_edit", st.session_state["ai_excl_result"])
                        st.success("✅ Applied. Download quotation to see result.")
                        st.session_state.pop("ai_scope_result", None)
                        st.session_state.pop("ai_excl_result", None)

            with _ai_t2:
                st.markdown("**Generate quality assurance points specific to this equipment**")
                _qa1, _qa2 = st.columns(2)
                _urs = _qa1.checkbox("Pharma URS-specific points", value=True, key="ai_urs")
                _tpi = _qa2.checkbox("TPI / inspection points",   value=True, key="ai_tpi")
                if st.button("⚡ Generate Quality Points", type="primary", key="ai_gen_quality"):
                    with st.spinner("Claude is writing quality points..."):
                        try:
                            _msg = _claude_client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=1000,
                                messages=[{"role": "user", "content": f"""You are a quality engineer at B&G Engineering Industries,
Hyderabad — Indian manufacturer of ASME-coded pharma process equipment.

Equipment:
{_eq_context}

Generate manufacturing and quality assurance points for this quotation.
{'Include pharma URS compliance points.' if _urs else ''}
{'Include TPI and documentation points.' if _tpi else ''}

Output ONLY a plain list — one point per line, no bullets, no numbers, no headers.
Generate 10-14 points. Each must be specific and technical — not generic."""}]
                            )
                            st.session_state["ai_quality_result"] = _msg.content[0].text.strip()
                            st.success("✅ Generated — review and apply below")
                        except Exception as _e:
                            st.error(f"Claude API error: {_e}")

                if "ai_quality_result" in st.session_state:
                    st.markdown("**Generated Quality Points — edit if needed:**")
                    _eq2 = st.text_area("Quality", value=st.session_state["ai_quality_result"], height=280, label_visibility="collapsed", key="ai_quality_edit")
                    if st.button("✅ Apply to Quotation", type="primary", key="ai_apply_quality"):
                        h["quality_points"] = st.session_state.get("ai_quality_edit", st.session_state["ai_quality_result"])
                        st.success("✅ Applied.")
                        st.session_state.pop("ai_quality_result", None)

            with _ai_t3:
                st.markdown("**Write a professional cover letter for this quotation**")
                _tone = st.radio("Tone", ["Formal & Technical", "Warm & Relationship-focused", "Competitive & Value-focused"], horizontal=True, key="ai_tone")
                if st.button("⚡ Write Cover Narrative", type="primary", key="ai_gen_narrative"):
                    with st.spinner("Claude is writing..."):
                        try:
                            _msg = _claude_client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=800,
                                messages=[{"role": "user", "content": f"""You are a senior sales engineer at B&G Engineering Industries,
Hyderabad — leading Indian manufacturer of pharma-grade SS316L process equipment.

You are writing a brief professional cover narrative (3-4 paragraphs) 
to accompany a technical quotation sent to {h.get('customer_name', 'the customer')}.

Equipment: {h.get('equipment_desc', '')} ({h.get('capacity_ltrs', '')} Litres)
Delivery: {h.get('delivery_weeks', '12-16')} weeks
Design Code: {h.get('design_code', 'ASME Sec VIII Div 1')}
MOC: {h.get('moc_shell', 'SS316L')}
Tone: {_tone}

Write 3-4 short paragraphs. Do not mention specific prices. Do not use hollow phrases like 'we are pleased to submit'."""}]
                            )
                            st.session_state["ai_narrative_result"] = _msg.content[0].text.strip()
                            st.success("✅ Generated")
                        except Exception as _e:
                            st.error(f"Claude API error: {_e}")

                if "ai_narrative_result" in st.session_state:
                    st.markdown("**Generated Narrative — edit if needed:**")
                    _en = st.text_area("Narrative", value=st.session_state["ai_narrative_result"], height=250, label_visibility="collapsed", key="ai_narrative_edit")
                    if st.button("✅ Add to Special Notes", type="primary", key="ai_apply_narrative"):
                        _existing = h.get("special_notes", "").strip()
                        _new_note = st.session_state.get("ai_narrative_edit", st.session_state["ai_narrative_result"])
                        h["special_notes"] = (_new_note + "\n\n" + _existing).strip()
                        st.success("✅ Added to Special Notes.")
                        st.session_state.pop("ai_narrative_result", None)

            with _ai_t4:
                st.markdown("**Claude reviews your commercial terms and flags issues**")
                if st.button("🔍 Review My Commercial Terms", type="primary", key="ai_review_terms"):
                    with st.spinner("Claude is reviewing..."):
                        try:
                            _msg = _claude_client.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=600,
                                messages=[{"role": "user", "content": f"""You are a senior commercial manager reviewing a quotation
for Indian pharma process equipment.

Review these commercial terms:
Payment Terms: {h.get('payment_terms', '')}
Delivery: {h.get('delivery_weeks', '')} weeks — {h.get('delivery_note', '')}
Offer Validity: {h.get('offer_validity', '')}
Warranty: {h.get('warranty_clause', '')}
Inspection Rights: {h.get('inspection_clause', '')}
Price Basis: {h.get('price_basis', '')}
GST Clause: {h.get('gst_clause', '')}

Respond with:
1. ✅ What protects B&G well
2. ⚠️ What needs attention or is missing
3. 💡 Specific improvements to suggest

Maximum 250 words."""}]
                            )
                            st.session_state["ai_review_result"] = _msg.content[0].text.strip()
                        except Exception as _e:
                            st.error(f"Claude API error: {_e}")

                if "ai_review_result" in st.session_state:
                    st.markdown("**Claude's Review:**")
                    st.markdown(st.session_state["ai_review_result"])
                    if st.button("🗑️ Clear", key="ai_clear_review"):
                        st.session_state.pop("ai_review_result", None)

        st.divider()
        st.markdown("**📄 Download Final Quotation**")
        st.caption("Edits above are captured in memory. Save Draft to persist, then download.")

        _qe_clients  = load_clients_full()
        _qe_cust     = next((c for c in _qe_clients if c["name"] == h.get("customer_name","")), {})

        _qe_T = calc_totals(
            st.session_state.est_parts, st.session_state.est_pipes, st.session_state.est_flanges,
            st.session_state.est_struct,
            st.session_state.est_fab, st.session_state.est_bo, st.session_state.est_oh,
            float(h.get("profit_margin_pct",10)), float(h.get("contingency_pct",0)),
            float(h.get("packing_amt",5000)), float(h.get("freight_amt",10000)),
            float(h.get("gst_pct",18)), float(h.get("engg_design_amt",25000)),
        )

        qe_dl1, qe_dl2, qe_dl3, qe_dl4 = st.columns(4)
        qe_dl1.download_button(
            "📄 Standard Quote",
            generate_docx(h, _qe_cust, _qe_T, st.session_state.est_fab, show_breakup=False),
            file_name=f"{h.get('qtn_number','QTN')}_{h.get('revision','R0')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, type="primary", key="qe_dl_std",
        )
        qe_dl2.download_button(
            "📋 With Scope Breakup",
            generate_docx(h, _qe_cust, _qe_T, st.session_state.est_fab, show_breakup=True),
            file_name=f"{h.get('qtn_number','QTN')}_{h.get('revision','R0')}_breakup.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key="qe_dl_bk",
        )
        if OPENPYXL_OK:
            qe_dl3.download_button(
                "📊 Fact Sheet (.xlsx)",
                generate_fact_sheet_xlsx(
                    h, st.session_state.est_parts, st.session_state.est_pipes,
                    st.session_state.est_flanges, st.session_state.est_struct,
                    st.session_state.est_fab,
                    st.session_state.est_bo, st.session_state.est_oh,
                    st.session_state.fab_rates, _qe_T,
                ),
                file_name=f"{h.get('qtn_number','QTN')}_{h.get('revision','R0')}_FactSheet.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True, key="qe_dl_fs",
            )
        else:
            qe_dl3.warning("Add `openpyxl` to requirements.txt")
        if qe_dl4.button("💾 Save Draft", use_container_width=True, type="primary", key="qe_save"):
            _do_save(reset_after=False)


# ══════════════════════════════════════════════════════════════════════════════
# TAB: SIMILAR EQUIPMENT
# ══════════════════════════════════════════════════════════════════════════════
with TAB_SIMILAR:
    st.subheader("🔍 Similar Equipment — Price Benchmark")
    st.caption("Search past estimations by type and capacity range to benchmark your current pricing.")
    sc1, sc2, sc3 = st.columns(3)
    s_equip  = sc1.selectbox("Equipment Type", ["All"] + EQUIPMENT_NAMES, key="sim_eq")
    s_cap_lo = sc2.number_input("Capacity from (Ltrs)", value=0.0,     min_value=0.0, key="sim_lo")
    s_cap_hi = sc3.number_input("Capacity to (Ltrs)",   value=99999.0, min_value=0.0, key="sim_hi")
    s_cust   = st.text_input("Customer contains (optional)", key="sim_cu")

    results = []
    for est in load_all_estimations():
        if s_equip != "All" and est.get("equipment_type") != s_equip: continue
        cap = float(est.get("capacity_ltrs") or 0)
        if not (s_cap_lo <= cap <= s_cap_hi): continue
        if s_cust and s_cust.lower() not in (est.get("customer_name", "") or "").lower(): continue
        T = calc_totals(
            json.loads(est.get("parts_json")   or "[]"),
            json.loads(est.get("pipes_json")   or "[]"),
            json.loads(est.get("flanges_json") or "[]"),
            json.loads(est.get("struct_json")  or "[]"),
            json.loads(est.get("fab_json")     or "[]"),
            json.loads(est.get("bo_json")      or "[]"),
            json.loads(est.get("oh_json")      or "[]"),
            float(est.get("profit_margin_pct") or 10), float(est.get("contingency_pct") or 0),
            float(est.get("packing_amt") or 0), float(est.get("freight_amt") or 0),
            float(est.get("gst_pct") or 18), float(est.get("engg_design_amt") or 0),
        )
        results.append({
            "QTN":      est.get("qtn_number", ""), "Customer": est.get("customer_name", ""),
            "Equipment": est.get("equipment_desc", ""), "Cap (L)": cap,
            "Status":   est.get("status", ""), "Ex-Works": f"₹{T['ex_works']:,.0f}",
            "FOR Price": f"₹{T['for_price']:,.0f}", "Margin %": f"{T['profit_pct_actual']:.1f}%",
            "Date":     str(est.get("updated_at", ""))[:10],
        })
    if results:
        st.dataframe(pd.DataFrame(results), use_container_width=True, hide_index=True)
        st.caption(f"{len(results)} estimations found.")
    else:
        st.info("No matching estimations. Adjust filters above.")

# ══════════════════════════════════════════════════════════════════════════════
# TAB: MASTERS
# ══════════════════════════════════════════════════════════════════════════════
with TAB_MASTERS:
    mt1, mt2 = st.tabs(["🔩 RM & BO Master", "⚙️ OH Master"])

    with mt1:
        st.subheader("Raw Material & Bought-Out Master")
        df_rm = pd.DataFrame(sb_fetch("est_rm_master", order="category"))
        if not df_rm.empty:
            # Add display-only serial number (1-based) so rows are easy to refer to
            df_rm.insert(0, "S.No", range(1, len(df_rm) + 1))
            st.dataframe(
                df_rm[["S.No", "ref_code", "description", "category", "material",
                       "spec", "size", "uom", "rate", "unit_wt_kg_per_m", "active"]],
                use_container_width=True, hide_index=True,
            )
            st.caption(f"Total items: {len(df_rm)}")
        with st.expander("➕ Add / Update Item"):
            with st.form("rm_add_form", clear_on_submit=True):
                c1, c2 = st.columns(2)
                rc = c1.text_input("Ref Code (unique)"); desc = c2.text_input("Description")
                c1, c2, c3 = st.columns(3)
                cat = c1.selectbox("Category", ["RM", "BO"]); rmt = c2.text_input("Type"); mat = c3.text_input("Material")
                c1, c2, c3, c4 = st.columns(4)
                spec = c1.text_input("Spec"); sz = c2.text_input("Size")
                uom  = c3.selectbox("UOM", ["Kg", "Nos", "Set", "LS", "Sq.M"])
                rate = c4.number_input("Rate ₹", min_value=0.0)
                uwt  = st.number_input("Unit Wt kg/m (pipes only)", min_value=0.0)
                if st.form_submit_button("Save"):
                    if sb_insert("est_rm_master", dict(ref_code=rc, description=desc, category=cat, rm_type=rmt, material=mat, spec=spec, size=sz, uom=uom, rate=rate, unit_wt_kg_per_m=uwt if uwt > 0 else None, active="Yes")):
                        st.cache_data.clear(); st.success(f"Saved {rc}"); st.rerun()

    with mt2:
        st.subheader("Overhead Master")
        df_oh = pd.DataFrame(sb_fetch("est_oh_master", order="oh_type"))
        if not df_oh.empty:
            # Add display-only serial number (1-based) so rows are easy to refer to
            df_oh.insert(0, "S.No", range(1, len(df_oh) + 1))
            st.dataframe(
                df_oh[["S.No", "oh_code", "description", "oh_type", "uom", "rate", "source"]],
                use_container_width=True, hide_index=True,
            )
            st.caption(f"Total items: {len(df_oh)}")
        with st.expander("➕ Add / Update OH"):
            with st.form("oh_add_form", clear_on_submit=True):
                c1, c2, c3, c4, c5 = st.columns(5)
                oc = c1.text_input("OH Code"); od = c2.text_input("Description")
                ot = c3.selectbox("Type", ["LABOUR", "LABOUR_BUFF", "CONSUMABLES", "TESTING", "DOCS", "PACKING", "TRANSPORT", "MISC", "ELECTRO_POLISH"])
                ou = c4.selectbox("UOM", ["Hr", "Sq.M", "%", "LS"])
                or_ = c5.number_input("Rate", min_value=0.0)
                if st.form_submit_button("Save"):
                    if sb_insert("est_oh_master", dict(oh_code=oc, description=od, oh_type=ot, uom=ou, rate=or_, source="Internal")):
                        st.cache_data.clear(); st.success(f"Saved {oc}"); st.rerun()    

        

                                  
