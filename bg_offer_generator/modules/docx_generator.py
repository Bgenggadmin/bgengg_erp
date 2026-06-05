"""
Offer DOCX Generator — Produces B&G-branded quotation documents.

Uses python-docx for document generation with custom styling
matching B&G brand colors (red/pink), logo, and table formatting.

Structure mirrors the 11-part offer template:
  Cover → TOC → PART I..PART XI

Updated May 2026:
  - Part IV now includes "Overall System Operational Cost" block
  - Part V rewritten with per-unit tables (Stripper / MEE / ATFD)
    matching the offer template (Feed Parameters + 3 system tables)

FIX Jun 2026:
  - _add_general_terms now reads data["general_terms"] instead of
    hardcoded text. ALL-CAPS lines (ending with ':') become bold
    sub-headings; remaining lines are normal paragraphs.
"""
import io
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from bg_offer_generator.utils.brand import (
    BRAND, DOCX_COLORS, FONT_PRIMARY, FONT_HEADING, COMPANY,
    COVER_LETTER_INTRO, OFFER_TOC,
)


# ---------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------
def _set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = DOCX_COLORS["border_gray"]):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _add_paragraph(doc, text: str, bold: bool = False, size: int = 11,
                   color_hex: str = None, alignment=None, space_after: int = 6):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = FONT_PRIMARY
    run.font.size = Pt(size)
    run.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_heading(doc, text: str, level: int = 1):
    h = doc.add_paragraph()
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(text)
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DOCX_COLORS["primary_red"])
        run.font.name = FONT_HEADING
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        run = h.add_run(text)
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DOCX_COLORS["primary_red"])
        run.font.name = FONT_HEADING
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
    else:
        run = h.add_run(text)
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DOCX_COLORS["dark_text"])
        run.font.name = FONT_HEADING
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
    return h


def _add_section_title(doc, part_label: str, part_title: str):
    doc.add_page_break()
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(part_label)
    r1.font.size = Pt(14)
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(DOCX_COLORS["accent_pink"])
    r1.font.name = FONT_HEADING
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(part_title)
    r2.font.size = Pt(22)
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(DOCX_COLORS["primary_red"])
    r2.font.name = FONT_HEADING
    p2.paragraph_format.space_after = Pt(24)
    _add_horizontal_line(doc, color=DOCX_COLORS["primary_red"])


def _add_horizontal_line(doc, color: str = None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), color or DOCX_COLORS["primary_red"])
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _make_table(doc, data: list, header: list = None, col_widths: list = None,
                title_row: str = None):
    """
    Create a table with B&G styling.
    title_row: optional single-cell title spanning the table width
               (used for offer-style block titles like "FEED PARAMETERS")
    """
    n_extra = (1 if title_row else 0) + (1 if header else 0)
    n_rows = len(data) + n_extra
    n_cols = len(data[0]) if data else (len(header) if header else 1)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_idx = 0

    # Title row (merged across all columns)
    if title_row:
        title_cells = table.rows[0].cells
        first = title_cells[0]
        for other in title_cells[1:]:
            first.merge(other)
        _set_cell_bg(first, DOCX_COLORS["primary_red"])
        _set_cell_borders(first)
        first.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        first.text = ""
        p = first.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_row)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.name = FONT_PRIMARY
        row_idx = 1

    # Header row
    if header:
        for col_idx, h_text in enumerate(header):
            cell = table.cell(row_idx, col_idx)
            _set_cell_bg(cell, DOCX_COLORS["table_header"])
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(h_text))
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
            run.font.name = FONT_PRIMARY
        row_idx += 1

    # Data rows
    for r_i, row in enumerate(data):
        alt_bg = (r_i % 2 == 1)
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx + r_i, col_idx)
            if alt_bg:
                _set_cell_bg(cell, DOCX_COLORS["table_row_alt"])
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value) if value is not None else "")
            run.font.size = Pt(10)
            run.font.name = FONT_PRIMARY

    if col_widths:
        for col_idx, width_in in enumerate(col_widths):
            for row in table.rows:
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = Inches(width_in)
    return table


def _fmt_int(value):
    """Format value as comma-separated int when possible, else as-is."""
    try:
        return f"{int(float(value)):,}"
    except (TypeError, ValueError):
        return str(value) if value is not None else "-"


def _fmt_inr_lakh(value_inr):
    """Format ₹ amount in Indian-style: 5,33,25,000."""
    try:
        n = int(float(value_inr))
    except (TypeError, ValueError):
        return str(value_inr)
    s = str(abs(n))
    if len(s) <= 3:
        return f"{'-' if n < 0 else ''}{s}"
    last3 = s[-3:]
    rest = s[:-3]
    groups = []
    while len(rest) > 2:
        groups.insert(0, rest[-2:])
        rest = rest[:-2]
    if rest:
        groups.insert(0, rest)
    return f"{'-' if n < 0 else ''}{','.join(groups)},{last3}"


# ---------------------------------------------------------------------
# Image insertion helper
# ---------------------------------------------------------------------
def _add_picture_flexible(run, source, width_inches):
    if source is None:
        return False
    try:
        if isinstance(source, (bytes, bytearray)):
            run.add_picture(io.BytesIO(source), width=Inches(width_inches))
            return True
        if isinstance(source, io.IOBase) or hasattr(source, "read"):
            run.add_picture(source, width=Inches(width_inches))
            return True
        p = Path(str(source))
        if p.exists():
            run.add_picture(str(p), width=Inches(width_inches))
            return True
    except Exception:
        pass
    return False


# ---------------------------------------------------------------------
# MAIN GENERATOR
# ---------------------------------------------------------------------
def generate_offer_docx(data: dict, logo_path: str = None,
                          tagline_path: str = None,
                          hero_path: str = None) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = FONT_PRIMARY
    style.font.size = Pt(11)

    _add_header_footer(doc, data, logo_path)

    # ============================================
    # COVER PAGE
    # ============================================
    _build_cover_page(doc, data, logo_path, tagline_path, hero_path)

    # ============================================
    # COVER LETTER
    # ============================================
    _build_cover_letter(doc, data)

    # ============================================
    # TABLE OF CONTENTS
    # ============================================
    doc.add_page_break()
    _add_heading(doc, "OFFER INDEX", level=1)
    toc_data = [[toc[0], toc[1]] for toc in OFFER_TOC]
    _make_table(doc, toc_data, header=["PART", "CONTENT"], col_widths=[1.2, 5.0])

    # ============================================
    # PART I: EXECUTIVE SUMMARY
    # ============================================
    _add_section_title(doc, "PART – I", "EXECUTIVE SUMMARY")
    for para in data["executive_summary"].strip().split("\n\n"):
        _add_paragraph(doc, para, size=11)

    # ============================================
    # PART II: PROCESS DESCRIPTION
    # ============================================
    _add_section_title(doc, "PART – II", "PROCESS DESCRIPTION")
    pd_data = data["process_description"]

    _add_heading(doc, "Stripper System", level=2)
    for para in pd_data["stripper"].strip().split("\n\n"):
        _add_paragraph(doc, para)

    _add_heading(doc, "Multi-Effect Evaporator", level=2)
    n_eff = pd_data.get("n_effects", 4)
    for para in pd_data["mee"].format(n_effects=n_eff).strip().split("\n\n"):
        _add_paragraph(doc, para)

    _add_heading(doc, "Agitated Thin Film Dryer (ATFD)", level=2)
    for para in pd_data["atfd"].strip().split("\n\n"):
        _add_paragraph(doc, para)

    # ============================================
    # PART III: PROCESS FLOW DIAGRAM
    # ============================================
    _add_section_title(doc, "PART – III", "PROCESS FLOW DIAGRAM")
    _add_paragraph(doc, "[Process Flow Diagram to be inserted by user]",
                    size=10, color_hex="888888",
                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc,
        "Flow: Effluent Feed → Stripper Column → Multi-Effect Evaporator → ATFD → Dry Solids",
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ============================================
    # PART IV: PLANT ECONOMICS & OPEX
    # ============================================
    _build_part_iv_economics(doc, data)

    # ============================================
    # PART V: TECHNICAL DETAILS & UTILITIES
    # ============================================
    _build_part_v_technical(doc, data)

    # ============================================
    # PART VI: SCOPE OF SUPPLY
    # ============================================
    _add_section_title(doc, "PART – VI", "SCOPE OF SUPPLY")
    _add_heading(doc, "I. Electro-Mechanical Items", level=2)

    _add_heading(doc, "Stripper System", level=3)
    _make_scope_table(doc, data["scope_stripper"])

    _add_heading(doc, "Evaporation System (MEE)", level=3)
    _make_scope_table(doc, data["scope_mee"])

    _add_heading(doc, "ATFD System", level=3)
    _make_scope_table(doc, data["scope_atfd"])

    _add_heading(doc, "II. Instruments & Automation Items", level=2)

    instr_data = data.get("instruments", [])
    inst_rows = [[i.get("item", ""), i.get("qty", ""), i.get("scope", "")]
                 for i in instr_data]
    _make_table(doc, inst_rows, header=["Item Description", "Quantity", "Scope"],
                col_widths=[3.8, 1.2, 1.5])

    _add_paragraph(doc, "", size=4)
    mcc_rows = [
        ["MCC Panel: Non-Compartmental Type, Floor mounting, Ambient temperature"],
        ["IP 54 protection for indoor installation."],
        ["MCC panel with MS CRCA and powder coated. 2 mm Thk for Main doors, frame."],
        ["MCC Panel Mains incoming circuit breaker shall be MCCB type. MCCB Ampere rating suitable for all electrical feeder loads."],
        ["MCC Panel busbar shall be of Electrolytic Aluminium Grade. Earth bus: Aluminium Grade."],
        ["MCC Panel outgoing feeders:\n  • Direct on Line (DOL) starter type up to 15 kW\n  • Star Delta starter from 18.5 kW and above\n  • VFD feeder as per final P&ID and tech details provided by B&G"],
        ["Feeders will be provided with termination with field power and control cable."],
        ["Mains incoming section shall be provided with Energy Meter (kWh), Voltmeter & Ammeter."],
    ]
    _make_table(doc, mcc_rows, col_widths=[6.7], title_row="MCC PANEL DETAILS")

    _add_paragraph(doc, "", size=4)
    plc_rows = [
        ["Plant shall be provided with PLC & SCADA based automation and control system."],
        ["PLC Make: ABB / Siemens / Reputed. Enclosure: MS Cabinet with powder coated."],
        ["Computer: Latest Windows, 21\" screen, 64-bit."],
        ["CPU module. Control system for monitoring and controlling various process parameters and automatic operation of various sections of plant."],
        ["SCADA software with Licensed version."],
        ["On/Off Switches, DO's, I/O's, Relays, MCB's shall be included."],
        ["Analogue/Digital Input-Output modules, power supply module and communication module."],
        ["System will be without redundancy. IP 52 protected."],
        ["To be installed in Non-Flame proof area with proper ventilation, louvers, lightening."],
        ["Note: Parameters monitored/controlled: Feed Flow, Steam Flow, Levels, Steam Pressure, CW in/out, Valve on/off, Temperatures, Pressures etc. B&G reserves right to choose different quantities, type, make, granted that system functionality shall remain same or better."],
    ]
    _make_table(doc, plc_rows, col_widths=[6.7], title_row="PLC AND SCADA SYSTEM DETAILS")

    # ============================================
    # PART VII: BATTERY LIMITS
    # ============================================
    _add_section_title(doc, "PART – VII", "BATTERY LIMITS")
    for i, item in enumerate(data["battery_limits"], 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(item).font.size = Pt(10)

    # ============================================
    # PART VIII: SCOPE MATRIX
    # ============================================
    _add_section_title(doc, "PART – VIII", "SCOPE MATRIX")
    sm_rows = [[i + 1, item["description"],
                 "✓" if item["bg"] else "✗",
                 "✓" if item["client"] else "✗"]
                for i, item in enumerate(data["scope_matrix"])]
    _make_table(doc, sm_rows,
                 header=["S.No", "Description of Work", "B&G", "CLIENT"],
                 col_widths=[0.5, 4.5, 0.75, 0.75])

    # ============================================
    # PART IX: BASIS OF COMMISSIONING
    # ============================================
    _add_section_title(doc, "PART – IX", "BASIS OF COMMISSIONING / TAKE-OVER")
    for b in data.get("commissioning_basis", []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(b).font.size = Pt(10)

    # ============================================
    # PART X: PRICE & COMMERCIAL TERMS
    # ============================================
    _add_section_title(doc, "PART – X", "PRICE & COMMERCIAL TERMS AND CONDITIONS")
    _add_heading(doc, "Price Summary", level=3)
    pr = data["pricing"]
    _add_paragraph(doc,
        f"Our total price for Design, Engineering, Manufacturing, Supply, Installation "
        f"and Commissioning of Plant as per given Technical Specification and Scope of work "
        f"(DAP, {pr['location_dap']}):",
        size=11)

    price_rows = [
        ["1", "Design, Engineering & Supply of Stripper, MEE & ATFD System",
         f"Rs. {pr.get('option1_equipment_price_cr', 0):.2f} Cr"],
        ["2", "Installation & Commissioning",
         f"Rs. {pr.get('option1_install_lakhs', 0):.0f} Lakhs"],
        ["3", "MS Structure",
         f"Rs. {pr.get('option1_ms_structure_lakhs', 0):.0f} Lakhs"],
        ["", "TOTAL PRICE",
         f"Rs. {pr.get('option1_total_cr', 0):.2f} Cr"],
    ]
    _make_table(doc, price_rows,
                 header=["S.N", "Item / Equipment / Service",
                         f"Option 1 — {pr.get('option1_moc', '')}"],
                 col_widths=[0.4, 4.0, 2.4])

    _add_paragraph(doc, "", size=6)
    _add_paragraph(doc,
        "Prices are inclusive of Design, Engineering, Manufacturing, Supply, Installation, "
        "Commissioning, Transportation & P&F. Above prices are EXCLUDING Taxes and Duties, "
        "and Transit Insurance (charged extra at the time of dispatch).",
        size=10)
    _add_paragraph(doc,
        f"Our above-mentioned price validity is {pr['price_validity_days']} days from the date of this offer.",
        bold=True, size=11)

    _add_heading(doc, "Terms of Payment", level=3)
    for term in pr["payment_terms"]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(term).font.size = Pt(10)

    _add_heading(doc, "Delivery Timeline (INCOTERMS 2020)", level=3)
    delivery = pr["delivery_timeline"]
    delivery_rows = [
        [f"Supply (DAP, {pr.get('location_dap', 'Hyderabad')})", delivery.get("supply_option1", "")],
        ["Installation",  delivery.get("installation", "")],
        ["Commissioning", delivery.get("commissioning", "")],
    ]
    _make_table(doc, delivery_rows, header=["Activity", "Timeline"],
                 col_widths=[3.0, 3.8])

    # ============================================
    # PART XI: GENERAL TERMS & CONDITIONS
    # ============================================
    _add_section_title(doc, "PART – XI", "GENERAL TERMS & CONDITIONS")
    _add_general_terms(doc, data)   # ← now receives data; renders data["general_terms"]

    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    _add_paragraph(doc, f"For {COMPANY['name']}", bold=True)
    doc.add_paragraph()
    _add_paragraph(doc, COMPANY["managing_partner"], bold=True)
    _add_paragraph(doc, COMPANY["managing_partner_title"], size=10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------
# PART IV — Plant Economics & OPEX
# ---------------------------------------------------------------------
def _build_part_iv_economics(doc, data: dict):
    _add_section_title(doc, "PART – IV", "ESTIMATED PLANT ECONOMICS & OPEX")

    econ = data["economics"]
    ut = data.get("utilities", {})
    cap = data["cover"].get("capacity_kld", 0)

    _add_heading(doc, "Overall System Operational Cost", level=3)
    op_rows = [
        ["Plant Capacity",         "KLD",      _fmt_int(cap)],
        ["Operating Hours",        "Hrs/Day",  _fmt_int(econ.get("operating_hours_day", 20))],
        ["Annual Operation",       "Days/Year",_fmt_int(econ.get("operating_days_year", 300))],
        ["Total Steam Consumption (Stripper + MEE + ATFD)",
            "Kg/h", _fmt_int(ut.get("total_steam_kgh", 0))],
        ["Total Power Consumption (For all Pump Motors & ATFD Motors, Excl. Cooling tower pumps)",
            "kWh",  _fmt_int(ut.get("total_power_kwh", 0))],
        ["Total Cooling Water (Stripper + MEE + ATFD)",
            "m³/h",
            f"{_fmt_int(ut.get('total_cooling_water_m3h', 0))} "
            f"({_fmt_int(ut.get('total_cooling_water_tr', 0))} TR)"],
        [f"Effluent Treatment Cost ({_fmt_int(cap)} KL)",
            "INR/KL", f"₹{_fmt_int(econ.get('effluent_treatment_cost_inr_kl', 0))} /-"],
        [f"Annual Operational Cost ({_fmt_int(cap)} KL)",
            "INR/Year", f"₹{_fmt_inr_lakh(econ.get('annual_operational_cost_inr', 0))} /-"],
    ]
    _make_table(doc, op_rows, header=["Utility", "UOM", "Value"],
                 col_widths=[4.0, 1.0, 1.7],
                 title_row="OVERALL SYSTEM OPERATIONAL COST")

    _add_paragraph(doc, "", size=6)

    _add_heading(doc, "BG ECOX-ZLD System Advantage", level=3)
    econ_data = [
        ["MEE Steam Consumption",
         f"{_fmt_int(econ['conventional_steam_kgh'])} Kg/h",
         f"{_fmt_int(econ['ecox_steam_kgh'])} Kg/h",
         f"Steam reduction up to ~{econ['steam_reduction_pct']:.1f}%"],
        ["Annual Steam Usage",
         f"{_fmt_int(econ['conventional_annual_steam_tons'])} tons/year",
         f"{_fmt_int(econ['ecox_annual_steam_tons'])} tons/year",
         f"{_fmt_int(econ['annual_steam_savings_tons'])} tons/year (Savings)"],
        ["Operating Cost",
         f"₹{econ['conventional_annual_cost_cr']:.2f} Cr/year",
         f"₹{econ['ecox_annual_cost_cr']:.2f} Cr/year",
         f"~₹{econ['annual_savings_lakhs']:.0f} Lakhs/Year (Savings)"],
    ]
    _make_table(doc, econ_data,
                 header=["Method", "Conventional Evaporation", "BG ECOX-ZLD", "Benefit"],
                 col_widths=[1.5, 1.8, 1.5, 2.0],
                 title_row="BG ECOX-ZLD SYSTEM ADVANTAGE")
    _add_paragraph(doc,
        f"Note: The above figures are based on {econ.get('operating_hours_day', 20):g} hrs/Day Operation, "
        f"{econ.get('operating_days_year', 300)} Days/Year, "
        f"Steam cost of ₹{econ.get('steam_cost_inr_kg', 2):g}/Kg, "
        f"Power cost of ₹{econ.get('power_cost_inr_kwh', 9):g}/kWh, "
        f"Cooling water cost of ₹{econ.get('cooling_water_cost_inr_m3', 90):g}/m³. "
        "Figures are tentative and indicative only.",
        size=9, color_hex="666666")


# ---------------------------------------------------------------------
# PART V — Technical Details & Utilities
# ---------------------------------------------------------------------
def _build_part_v_technical(doc, data: dict):
    _add_section_title(doc, "PART – V", "TECHNICAL DETAILS & UTILITIES")

    fp = data["feed_parameters"]
    feed_rows = [
        ["1",  "Feed / Capacity",            "KLD",      _fmt_int(fp["capacity_kld"])],
        ["2",  "Feed pH",                    "-",        fp.get("feed_ph", "-")],
        ["3",  "Specific Gravity",           "-",        fp.get("specific_gravity", "1.0")],
        ["4",  "Total COD",                  "PPM",      _fmt_int(fp.get("total_cod_ppm", 0))],
        ["5",  "Volatile Organic Solvents",  "PPM",      _fmt_int(fp.get("volatile_organic_solvents_ppm", 0))],
        ["6",  "Total Solids",               "% w/w",    str(fp.get("total_solids_pct", "-"))],
        ["7",  "Suspended Solids",           "PPM",      str(fp.get("suspended_solids_ppm", "-"))],
        ["8",  "Feed Temperature",           "Deg.C",    _fmt_int(fp.get("feed_temp_c", 30))],
        ["9",  "Total Hardness",             "PPM",      str(fp.get("total_hardness_ppm", "-"))],
        ["10", "Silica",                     "PPM",      str(fp.get("silica_ppm", "-"))],
        ["11", "Free Chloride",              "PPM",      str(fp.get("free_chloride_ppm", "-"))],
        ["12", "Feed Nature",                "-",        str(fp.get("feed_nature", "-"))],
    ]
    _make_table(doc, feed_rows,
                 header=["S.No", "Parameter", "UOM", "Value"],
                 col_widths=[0.6, 3.2, 1.0, 1.9],
                 title_row="FEED PARAMETERS")
    _add_paragraph(doc, "Note: The above parameters are considered as per customer suggestion during discussions.",
                   size=9, color_hex="666666")

    s = data["technical_specs"]["stripper"]
    stripper_rows = [
        ["Type",                                         "-",     s.get("type", "-")],
        ["Inlet Feed Rate",                              "Kg/h",  _fmt_int(s.get("feed_kgh", 0))],
        ["Top Distillate Out",                           "Kg/h",  f"{_fmt_int(s.get('distillate_kgh', 0))} ({s.get('distillate_composition', '')})"],
        ["Stripper Bottom Out",                          "Kg/h",  _fmt_int(s.get("bottoms_kgh", 0))],
        ["Reflux Rate",                                  "Kg/h",  _fmt_int(s.get("reflux_kgh", 0))],
        [f"Dry & Saturated Steam ({s.get('steam_pressure', '1.5 Bar-g')})",
                                                          "Kg/h",  _fmt_int(s.get("steam_kgh", 0))],
        ["Power Consumption (415V, 50 Hz, 3 Phase)",     "kWh",   _fmt_int(s.get("power_kwh", 0))],
        [f"Cooling Water ({s.get('cooling_water_temps', 'In/Out: 32 / 38 °C')})",
                                                          "m³/h",  f"{_fmt_int(s.get('cooling_water_m3h', 0))} ({_fmt_int(s.get('cooling_water_tr', 0))} TR)"],
        [f"Compressed Air at {s.get('compressed_air_pressure', '6 Bar-g')}",
                                                          "Nm³/h", str(s.get("compressed_air_nm3h", "-"))],
    ]
    _make_table(doc, stripper_rows,
                 header=["Section", "UOM", "Value"],
                 col_widths=[3.5, 1.0, 2.2],
                 title_row="STRIPPER SYSTEM")

    m = data["technical_specs"]["mee"]
    mee_rows = [
        ["Type",                                          "-",     f"{m.get('type', '4-Effects')}, {m.get('configuration', 'Forced Circulation Type')}"],
        ["Feed Inlet (Stripper Bottom)",                  "Kg/h",  _fmt_int(m.get("feed_kgh", 0))],
        ["Feed Solids",                                   "%",     str(m.get("feed_solids_pct", "-"))],
        ["Water Evaporation Rate",                        "Kg/h",  f"{_fmt_int(m.get('evaporation_kgh', 0))} (Max)"],
        ["MEE Concentrate Out",                           "Kg/h",  f"{_fmt_int(m.get('concentrate_kgh', 0))} (Max)"],
        ["MEE Concentrate Out",                           "%",     _fmt_int(m.get("concentrate_solids_pct", 40))],
        [f"Dry & Saturated Steam ({m.get('steam_pressure', '1.5 Bar-g')})",
                                                           "Kg/h",  _fmt_int(m.get("steam_kgh", 0))],
        ["Steam Economy",                                 "Kg/Kg", f"{m.get('steam_economy', 4.3):.1f}" if isinstance(m.get('steam_economy'), (int, float)) else str(m.get('steam_economy', '-'))],
        ["Power Consumption (415V, 50 Hz, 3 Phase)",      "kWh",   _fmt_int(m.get("power_kwh", 0))],
        [f"Cooling Water ({m.get('cooling_water_temps', 'In/Out: 32 / 38 °C')})",
                                                           "m³/h",  f"{_fmt_int(m.get('cooling_water_m3h', 0))} ({_fmt_int(m.get('cooling_water_tr', 0))} TR)"],
        [f"Compressed Air at {m.get('compressed_air_pressure', '6 Bar-g')}",
                                                           "Nm³/h", str(m.get("compressed_air_nm3h", "-"))],
    ]
    _make_table(doc, mee_rows,
                 header=["Section", "UOM", "Value"],
                 col_widths=[3.5, 1.0, 2.2],
                 title_row="MULTIPLE EFFECT EVAPORATOR SYSTEM")

    a = data["technical_specs"]["atfd"]
    atfd_rows = [
        ["Feed Inlet (MEE Concentrate)",                  "Kg/h",  f"{_fmt_int(a.get('feed_kgh', 0))} (Designed for Max flow)"],
        ["Feed Solids",                                   "%",     _fmt_int(a.get("feed_solids_pct", 40))],
        ["Water Evaporation Rate",                        "Kg/h",  _fmt_int(a.get("evaporation_kgh", 0))],
        ["ATFD Product Out",                              "Kg/h",  _fmt_int(a.get("product_kgh", 0))],
        ["Moisture in ATFD Product",                      "%",     str(a.get("product_moisture_pct", "8-10"))],
        [f"Dry & Saturated Steam ({a.get('steam_pressure', '1.5 Bar-g')})",
                                                           "Kg/h",  _fmt_int(a.get("steam_kgh", 0))],
        ["Power Consumption (415V, 50 Hz, 3 Phase)",      "kWh",   _fmt_int(a.get("power_kwh", 0))],
        [f"Cooling Water ({a.get('cooling_water_temps', 'In/Out: 32 / 38 °C')})",
                                                           "m³/h",  f"{_fmt_int(a.get('cooling_water_m3h', 0))} ({_fmt_int(a.get('cooling_water_tr', 0))} TR)"],
        [f"Compressed Air at {a.get('compressed_air_pressure', '6 Bar-g')}",
                                                           "Nm³/h", str(a.get("compressed_air_nm3h", "-"))],
    ]
    _make_table(doc, atfd_rows,
                 header=["Section", "UOM", "Value"],
                 col_widths=[3.5, 1.0, 2.2],
                 title_row="AGITATED THIN FILM DRYER (ATFD)")

    _add_heading(doc, "Performance Guarantee Parameters", level=3)
    for bullet in data.get("performance_guarantee", []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bullet).font.size = Pt(10)

    _add_heading(doc, "Notes", level=3)
    notes = [
        "System shall be cleaned as suggested by supplier during commissioning process.",
        "Above mentioned energy consumption and performance are based on timely CIP of system.",
        "Any change in feed parameters and utility shall impact the system performance.",
        "Plant performance will depend on regular CIP and maintenance as per plant operation manual.",
        "In case of variation in initial solids of feed on lower side, ATFD feed will be reduced, "
        "thereby vapor generation in ATFD will also be reduced. In such case ATFD vapor which is being "
        "used in evaporator as a heating medium will be reduced and fresh steam consumption will "
        "increase based on inlet solid.",
    ]
    for n in notes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(n).font.size = Pt(10)


# ---------------------------------------------------------------------
# Scope tables, cover page, cover letter, header/footer
# ---------------------------------------------------------------------
def _make_scope_table(doc, scope_items: list):
    rows = [[i.get("equipment", ""), i.get("specification", ""), i.get("qty", ""),
             "✓" if i.get("bg_scope") else "✗",
             "✓" if i.get("buyer_scope") else "✗"]
             for i in scope_items]
    _make_table(doc, rows,
                 header=["Equipment", "Specification", "Qty", "B&G Scope", "Buyer Scope"],
                 col_widths=[1.6, 3.0, 0.6, 0.7, 0.7])
    _add_paragraph(doc, "Note: W = Working, SB = Working Standby, SSB = Store Standby. "
                         "Motors/Instruments per specification above.",
                    size=9, color_hex="666666")


def _build_cover_page(doc, data: dict, logo_path, tagline_path, hero_path):
    if logo_path:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        _add_picture_flexible(run, logo_path, 2.0)

    if tagline_path:
        p_tag = doc.add_paragraph()
        p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_tag.add_run()
        _add_picture_flexible(run, tagline_path, 3.5)

    doc.add_paragraph()
    doc.add_paragraph()

    _add_paragraph(doc, "TECHNO-COMMERCIAL OFFER", bold=True, size=22,
                   color_hex=DOCX_COLORS["primary_red"],
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, 'BG "ECOX-ZLD" SYSTEM', bold=True, size=18,
                   color_hex=DOCX_COLORS["accent_pink"],
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"CAPACITY: {data['cover']['capacity_kld']} KLD",
                   bold=True, size=16,
                   color_hex=DOCX_COLORS["dark_text"],
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    _add_horizontal_line(doc, color=DOCX_COLORS["primary_red"])

    if hero_path:
        p_hero = doc.add_paragraph()
        p_hero.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_hero.add_run()
        _add_picture_flexible(run, hero_path, 4.0)

    doc.add_paragraph()
    cov = data["cover"]
    info_rows = [
        ["Quote Reference", cov["quote_ref"]],
        ["Quotation Date", cov["quote_date"]],
        ["Submitted to", cov["submitted_to"]],
        ["Location", cov["location"]],
        ["Prepared By", cov["prepared_by"]],
        ["Contact Details", cov["contact_details"]],
        ["E-mail", cov["email"]],
    ]
    _make_table(doc, info_rows, col_widths=[2.0, 4.0])


def _build_cover_letter(doc, data: dict):
    doc.add_page_break()
    cov = data["cover"]

    _add_paragraph(doc, "To,", bold=True, size=12)
    _add_paragraph(doc, cov["submitted_to"], bold=True, size=12)
    _add_paragraph(doc, cov["location"], bold=True, size=12)
    doc.add_paragraph()
    _add_paragraph(doc, f"Kind Attn.: {cov['kind_attn']}", bold=True)
    _add_paragraph(doc, f"Subject: {cov['subject']}", bold=True)
    if cov.get("discussion_date"):
        _add_paragraph(doc, f"Reference: As per discussions dated {cov['discussion_date']}",
                       bold=True)
    doc.add_paragraph()

    discussion_date = cov.get("discussion_date") or "recent meetings"
    intro = COVER_LETTER_INTRO.format(
        discussion_date=discussion_date,
        capacity=cov["capacity_kld"]
    )
    for para in intro.split("\n\n"):
        _add_paragraph(doc, para)

    doc.add_paragraph()
    _add_paragraph(doc, f"For, {COMPANY['name']}", bold=True)
    doc.add_paragraph()
    _add_paragraph(doc, COMPANY["managing_partner"], bold=True)
    _add_paragraph(doc, COMPANY["managing_partner_title"], size=10)


def _add_header_footer(doc, data: dict, logo_path: str = None):
    section = doc.sections[0]
    header = section.header
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = h_para.add_run(f"Quote Ref: {data['cover']['quote_ref']}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("888888")

    footer = section.footer
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = f_para.add_run(f"{COMPANY['name']} · {COMPANY['address']}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("888888")


# ---------------------------------------------------------------------
# PART XI — General Terms & Conditions
#
# FIX: Previously this function had hardcoded abbreviated T&C text
# and ignored data["general_terms"] entirely.
#
# Now it reads data["general_terms"] (the full editable text from Tab ⑩)
# and renders it with smart heading detection:
#   - A line that is ALL-CAPS and ends with ':' → bold clause heading (level 3)
#   - Empty lines → spacing only (no blank paragraphs)
#   - Everything else → normal body paragraph at 10pt
# ---------------------------------------------------------------------
def _add_general_terms(doc, data: dict):
    """
    Render Part XI from data["general_terms"].

    Formatting rules applied to each line:
      • ALL-CAPS line ending with ':' → bold sub-heading (_add_heading level 3)
      • Blank line                    → skip (spacing comes from paragraph_format)
      • Any other line                → 10pt body paragraph
    """
    raw_text = (data.get("general_terms") or "").strip()

    if not raw_text:
        _add_paragraph(
            doc,
            "General terms and conditions not specified. Please add them in Tab ⑩ Gen. T&C.",
            size=10, color_hex="888888"
        )
        return

    for line in raw_text.splitlines():
        stripped = line.strip()

        if not stripped:
            # Skip blank lines — paragraph spacing provides visual separation
            continue

        # Detect clause heading: ALL-CAPS (allowing spaces, digits, &, /, comma,
        # apostrophe, parentheses) and ends with ':'
        is_heading = (
            stripped.endswith(":")
            and stripped == stripped.upper()
        )

        if is_heading:
            _add_heading(doc, stripped, level=3)
        else:
            _add_paragraph(doc, stripped, size=10)
